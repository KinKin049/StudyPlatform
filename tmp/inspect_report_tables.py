import sys
from pathlib import Path

from docx import Document


docx_path = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("CET46/StudyPlatform-report-draft-ch1-6-brand-polished.docx")
document = Document(docx_path)

print(f"document={docx_path}")
print(f"tables={len(document.tables)}")
for table_index, table in enumerate(document.tables, 1):
    rows = len(table.rows)
    cols = len(table.columns)
    max_lengths = []
    avg_lengths = []
    for col_index in range(cols):
        values = [
            table.cell(row_index, col_index).text.replace("\n", " ").strip()
            for row_index in range(rows)
        ]
        max_lengths.append(max([len(value) for value in values] or [0]))
        avg_lengths.append(round(sum(len(value) for value in values) / max(len(values), 1), 1))
    headers = " | ".join(
        table.cell(0, col_index).text.replace("\n", " ").strip()[:24]
        for col_index in range(cols)
    )
    grid = table._tbl.tblGrid
    widths = []
    if grid is not None:
        for col in grid.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridCol"):
            widths.append(col.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w"))
    print(f"{table_index}: {rows}x{cols} max={max_lengths} avg={avg_lengths} widths={widths} header={headers}")
