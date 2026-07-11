from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill, Side, Border
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "CET46"
DOCX_PATH = OUT_DIR / "表7-1系统功能模块与页面完成情况.docx"
XLSX_PATH = OUT_DIR / "表7-1系统功能模块与页面完成情况.xlsx"


HEADERS = ["功能模块", "主要功能点", "相关页面/路由", "对应前端文件", "相关后端接口/数据表"]

ROWS = [
    [
        "用户认证与账号管理",
        "完成用户登录、注册、忘记密码、入职引导、角色识别、账号宠物绑定等功能，为后续课程学习、个人画像、教师端和后台管理提供统一身份基础。",
        "/、/login、/register、/forgot-password、/onboarding",
        "studyplatform-vue/src/pages/HomePage.vue；studyplatform-vue/src/pages/AuthLoginPage.vue；studyplatform-vue/src/pages/AuthRegisterPage.vue；studyplatform-vue/src/pages/AuthForgotPasswordPage.vue；studyplatform-vue/src/pages/AuthOnboardingPage.vue；studyplatform-vue/src/api/auth.js",
        "接口：/api/auth/login、/api/auth/register、/api/auth/password-reset/code、/api/auth/password-reset/confirm、/api/auth/onboarding、/api/auth/pet；数据表：users、password_reset_codes、profile_user_profiles",
    ],
    [
        "在线学堂与课程学习",
        "完成在线学堂首页、我的课程、课程分类浏览、课程详情、加入课程、课程评价、作业、考试、精品教材、购物车、教材订单与支付等功能。",
        "/academy/home、/academy/my-class、/academy/my-courses、/academy/assignments、/academy/assignments/:assignmentId、/academy/exams、/academy/exams/:examId、/academy/exams/:examId/take、/academy/open-courses、/academy/open-courses/:id、/academy/general-courses、/academy/general-courses/:id、/academy/micro-majors、/academy/micro-majors/:id、/academy/textbooks、/academy/textbooks/:id、/academy/textbook-cart",
        "studyplatform-vue/src/pages/AcademyPage.vue；studyplatform-vue/src/pages/academy/AcademyHome.vue；AcademyMyClass.vue；AcademyAggregatePage.vue；AcademyOpenCourses.vue；AcademyGeneralCourses.vue；AcademyMicroMajors.vue；AcademyCourseDetail.vue；AcademyAssignmentDetail.vue；AcademyExamIntro.vue；AcademyExamDetail.vue；AcademyTextbooks.vue；AcademyTextbookDetail.vue；AcademyTextbookCart.vue；studyplatform-vue/src/api/academy.js",
        "接口：/api/academy/home、/api/academy/my-courses、/api/academy/assignments、/api/academy/exams、/api/academy/online-open-courses、/api/academy/general-courses、/api/academy/micro-major-courses、/api/academy/textbooks、/api/academy/textbook-cart、/api/academy/textbook-orders、/api/academy/textbook-payments；数据表：learning_content_blocks、online_open_courses、general_courses、micro_major_courses、teacher_published_courses、academy_course_enrollments、academy_course_reviews、academy_assignments、academy_assignment_questions、academy_assignment_submissions、academy_exams、academy_exam_questions、academy_exam_submissions、excellent_textbooks、academy_textbook_details、academy_textbook_cart_items、academy_textbook_orders、academy_textbook_order_items、academy_textbook_payments、academy_textbook_reviews",
    ],
    [
        "题库练习与学习反馈",
        "完成课程题库目录、题目练习、题目详情、错题记录、错题复习、收藏题目、词汇训练和通用题源浏览等功能，并将答题行为沉淀到学习画像和金币结算中。",
        "/academy/question-bank、/academy/question-bank/courses、/academy/question-bank/courses/:courseCode、/academy/question-bank/mistakes、/academy/question-bank/favorites",
        "studyplatform-vue/src/pages/academy/AcademyQuestionBank.vue；AcademyQuestionBankCourses.vue；AcademyQuestionBankCourseDetail.vue；AcademyQuestionBankMistakes.vue；AcademyQuestionBankFavorites.vue；studyplatform-vue/src/api/academy.js",
        "接口：/api/academy/question-bank/course-catalog、/api/academy/question-bank/courses/{code}、/api/academy/question-bank/mistakes、/api/academy/question-bank/mistakes/answers、/api/academy/question-bank/favorites、/api/academy/question-bank/subjects、/api/academy/question-bank/problems；数据表：course_question_bank_categories、course_question_bank_sets、course_question_bank_questions、course_question_bank_mistakes、course_question_bank_favorites、question_bank_subjects、question_bank_tags、question_bank_problems、question_bank_problem_subjects、profile_learning_events、coin_reward_records",
    ],
    [
        "个人主页与学习画像",
        "完成个人资料展示与编辑、头像上传、学习时长统计、今日学习、练习记录、连续学习、学习热力图、金币余额、近期动态、成就徽章和不同账号画像隔离等功能。",
        "/profile",
        "studyplatform-vue/src/pages/ProfilePage.vue；studyplatform-vue/src/api/profile.js",
        "接口：/api/profile/overview、/api/profile/user、/api/profile/avatar、/api/profile/events、/api/profile/learning-time；数据表：users、profile_user_profiles、profile_learning_events、profile_learning_time_records、coin_reward_records、coin_spend_records、game_ladder_jump_records、game_type_warrior_records",
    ],
    [
        "教师端课程管理与教师信箱",
        "完成教师工作台、教师课程发布、课程编辑、课程删除、作业布置、作业题型维护、OJ题目布置、课程评价回复、教师信箱消息查看和已读处理等功能。",
        "/profile、/teacher-mailbox、/academy/assignments、/academy/exams、/academy/open-courses/:id",
        "studyplatform-vue/src/pages/ProfilePage.vue；studyplatform-vue/src/pages/TeacherMailboxPage.vue；studyplatform-vue/src/pages/academy/AcademyAssignmentDetail.vue；AcademyExamDetail.vue；AcademyCourseDetail.vue；studyplatform-vue/src/api/academy.js",
        "接口：/api/academy/teacher/workbench、/api/academy/teacher/workbench/mailbox/read、/api/academy/teacher/assignments、/api/academy/online-open-courses、/api/academy/online-open-courses/teacher/mine、/api/academy/reviews/{reviewId}/reply；数据表：users、teacher_published_courses、online_open_courses、academy_assignments、academy_assignment_questions、academy_assignment_submissions、academy_course_reviews、oj_problems",
    ],
    [
        "AI宠物与学习辅助",
        "完成全局AI宠物浮窗、宠物形象切换、AI问答、页面上下文问答、页面导航意图识别、待办清单、番茄钟和账号宠物隔离等功能。",
        "全局浮窗组件，主要出现在 /、/academy/home、/profile、/exchange、/visualization、/games 等页面",
        "studyplatform-vue/src/components/AiPetWidget.vue；studyplatform-vue/src/data/aiPetShop.js；studyplatform-vue/src/api/aiPet.js；studyplatform-vue/src/pages/ExchangeCenter.vue",
        "接口：/api/ai-pet/chat、/api/auth/pet、/api/rewards/vouchers/items、/api/rewards/vouchers/exchange；数据表：users、voucher_items、user_vouchers、coin_reward_records、coin_spend_records",
    ],
    [
        "OJ在线编程与实验平台",
        "完成实验平台入口、OJ题目列表、题目详情、代码提交、自动评测、测试用例结果、测井仿真、石油气生产仿真、抽油泵仿真、油藏动态仿真、注水开发仿真和压裂酸化仿真等功能。",
        "/lab、/lab/oj、/lab/well-log、/lab/petroleum、/lab/production",
        "studyplatform-vue/src/pages/LabPlatform.vue；studyplatform-vue/src/pages/OjPlatform.vue；studyplatform-vue/src/pages/WellLogSimulation.vue；studyplatform-vue/src/pages/ProductionSimulation.vue；studyplatform-vue/src/pages/petroleum/PetroleumSimulation.vue；PumpIndicatorPanel.vue；ReservoirDynamicsPanel.vue；WaterfloodPanel.vue；WellLogPanel.vue；studyplatform-vue/src/oj/api.js；studyplatform-vue/src/pages/petroleum/api.js",
        "接口：/api/oj/problems、/api/oj/problems/{id}、/api/oj/submissions、/api/oj/submissions/{id}/cases、/api/well-log/template/list、/api/well-log/record/save、/api/well-log/record/page、/api/production/pump/save、/api/production/reservoir/save、/api/production/waterflood/save、/api/production/stimulation/save；数据表：oj_categories、oj_problems、oj_test_cases、oj_submissions、oj_submission_cases、well_log_template、well_log_record、production_pump_record、production_reservoir_record、production_waterflood_record、production_stimulation_record",
    ],
    [
        "可视化学习",
        "完成可视化学习中心、数据结构动画、算法演示、二维函数图像绘制、空间模型引导和三维空间模型展示等功能，用于辅助抽象知识理解。",
        "/visualization、/visualization/data-structure、/visualization/data-structure/:demoId、/visualization/function-2d、/visualization/space-models、/visualization/space-3d",
        "studyplatform-vue/src/pages/visualization/VisualizationHome.vue；DataStructureVisualization.vue；AlgorithmDemoViewer.vue；FunctionGraph2D.vue；SpaceModelGuide.vue；SpaceModel3D.vue；algorithmDemos.js；spaceModelCatalog.js",
        "主要为前端交互与可视化渲染功能，结合 /api/profile/learning-time、/api/profile/events 记录学习行为；数据表：profile_learning_time_records、profile_learning_events、coin_reward_records",
    ],
    [
        "学习游戏",
        "完成学习游戏平台、万题天梯跳、Type Warrior单词训练、题库加载、游戏成绩记录、波次统计、连击统计、游戏金币换算为平台金币等功能。",
        "/games、/games/:gameId",
        "studyplatform-vue/src/pages/games/GamePlatform.vue；LadderJumpGame.vue；TypeWarriorGame.vue；ladder-jump/LadderJumpGamePage.vue；type-warrior/TypeWarriorGamePage.vue；studyplatform-vue/src/api/games.js",
        "接口：/api/games/ladder-jump/question-banks、/api/games/ladder-jump/questions、/api/games/ladder-jump/records、/api/games/type-warrior/records、/api/academy/question-bank/type-warrior/words；数据表：game_ladder_jump_records、game_type_warrior_records、course_question_bank_sets、course_question_bank_questions、profile_learning_events、coin_reward_records",
    ],
    [
        "金币兑换中心与我的卡券",
        "完成金币余额展示、金币来源说明、兑换商品展示、AI宠物形象兑换、学习卡券兑换、已兑换卡券查看、金币扣减、卡券数量维护和账号资产隔离等功能。",
        "/exchange、/exchange/vouchers",
        "studyplatform-vue/src/pages/ExchangeCenter.vue；studyplatform-vue/src/pages/MyVouchers.vue；studyplatform-vue/src/api/vouchers.js；studyplatform-vue/src/composables/useUserVouchers.js；studyplatform-vue/src/data/aiPetShop.js",
        "接口：/api/rewards/vouchers、/api/rewards/vouchers/items、/api/rewards/vouchers/exchange、/api/rewards/vouchers/use、/api/profile/overview；数据表：voucher_items、user_vouchers、coin_reward_records、coin_spend_records、users、profile_user_profiles",
    ],
    [
        "后台管理",
        "完成管理员用户管理、课程管理、课程分类管理、课程评论管理、题库套题管理、题目管理、OJ题目管理、OJ分类管理、兑换商品管理和后台数据维护等功能。",
        "/admin",
        "studyplatform-vue/src/pages/AdminPage.vue；studyplatform-vue/src/api/admin.js",
        "接口：/api/admin/users、/api/admin/courses、/api/admin/course-categories、/api/admin/reviews、/api/admin/question-bank/sets、/api/admin/question-bank/questions、/api/admin/oj/problems、/api/admin/oj/categories、/api/admin/vouchers；数据表：users、online_open_courses、general_courses、micro_major_courses、admin_course_categories、academy_course_reviews、academy_textbook_reviews、course_question_bank_sets、course_question_bank_questions、oj_categories、oj_problems、oj_test_cases、voucher_items",
    ],
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_width(cell, width_cm: float) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def style_docx_cell(cell, *, bold: bool = False, fill: str | None = None) -> None:
    if fill:
        set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1
        for run in paragraph.runs:
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(8)
            run.bold = bold


def create_docx() -> None:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    table = document.add_table(rows=1, cols=len(HEADERS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    widths = [2.4, 5.0, 5.0, 6.5, 7.8]

    for col_idx, header in enumerate(HEADERS):
        cell = table.rows[0].cells[col_idx]
        cell.text = header
        set_cell_width(cell, widths[col_idx])
        style_docx_cell(cell, bold=True, fill="D9EAF7")

    for row in ROWS:
        cells = table.add_row().cells
        for col_idx, value in enumerate(row):
            cells[col_idx].text = value
            set_cell_width(cells[col_idx], widths[col_idx])
            style_docx_cell(cells[col_idx])

    document.save(DOCX_PATH)


def create_xlsx() -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "表7-1"
    sheet.append(HEADERS)
    for row in ROWS:
        sheet.append(row)

    widths = [18, 45, 55, 70, 85]
    header_fill = PatternFill("solid", fgColor="D9EAF7")
    thin = Side(style="thin", color="808080")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for col_idx, width in enumerate(widths, start=1):
        sheet.column_dimensions[get_column_letter(col_idx)].width = width

    for row in sheet.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = border
            cell.font = Font(name="宋体", size=10)
            if cell.row == 1:
                cell.font = Font(name="宋体", size=10, bold=True)
                cell.fill = header_fill

    for row_idx in range(2, sheet.max_row + 1):
        sheet.row_dimensions[row_idx].height = 115
    sheet.freeze_panes = "A2"
    workbook.save(XLSX_PATH)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    create_docx()
    create_xlsx()
    print(DOCX_PATH)
    print(XLSX_PATH)


if __name__ == "__main__":
    main()
