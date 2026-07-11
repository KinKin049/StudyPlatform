import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, Twips


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "CET46" / "报告-latest-附录A模块接口明细页面格式修订.docx"
OUTPUT = ROOT / "CET46" / "报告-latest-全文格式统一修订.docx"

FONT_CN = "宋体"
FONT_EN = "Times New Roman"


def set_run_font(run, size_pt, bold=None):
    run.font.name = FONT_EN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def set_paragraph_format(paragraph, size_pt, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         line_spacing=1.0, first_indent=False):
    pf = paragraph.paragraph_format
    paragraph.alignment = alignment
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = line_spacing
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(21) if first_indent else Pt(0)
    for run in paragraph.runs:
        set_run_font(run, size_pt, bold)


def is_top_heading(text):
    return bool(re.fullmatch(r"\d+\s+.+", text)) or text.startswith("附录")


def is_second_heading(text):
    return bool(re.fullmatch(r"\d+\.\d+\s+.+", text))


def is_third_heading(text):
    return bool(re.fullmatch(r"\d+\.\d+\.\d+\s+.+", text))


def is_figure_caption(text):
    return bool(re.match(r"^图\d+[-－]\d+\s+", text)) or bool(re.match(r"^图X[-－]X\s+", text))


def is_table_caption(text):
    return bool(re.match(r"^表\d+[-－]\d+\s+", text)) or bool(re.match(r"^表\d+\.\d+\s+", text))


def looks_like_cover_or_notice(index):
    return index <= 21


def format_body_paragraphs(document):
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue

        if looks_like_cover_or_notice(index):
            # Preserve cover page and instruction-page special typography.
            continue

        if is_figure_caption(text):
            set_paragraph_format(paragraph, 9, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                 line_spacing=1.0, first_indent=False)
        elif is_table_caption(text):
            set_paragraph_format(paragraph, 9, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                 line_spacing=1.0, first_indent=False)
        elif is_third_heading(text):
            set_paragraph_format(paragraph, 12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                 line_spacing=1.5, first_indent=False)
        elif is_second_heading(text):
            set_paragraph_format(paragraph, 12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                 line_spacing=1.5, first_indent=False)
        elif is_top_heading(text):
            set_paragraph_format(paragraph, 14, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                 line_spacing=1.5, first_indent=False)
        else:
            set_paragraph_format(paragraph, 10.5, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                 line_spacing=1.0, first_indent=True)


def format_table_cell_paragraph(paragraph, header=False):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(0)
    for run in paragraph.runs:
        set_run_font(run, 10.5, bold=True if header else False)


def format_tables(document):
    for table in document.tables:
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    format_table_cell_paragraph(paragraph, header=(row_index == 0))


def normalize_sections(document):
    # A4 in Word uses the same EMU values as python-docx Cm conversion.
    portrait_width = Cm(21)
    portrait_height = Cm(29.7)
    for section in document.sections:
        is_landscape = section.orientation == WD_ORIENTATION.LANDSCAPE or section.page_width > section.page_height
        if is_landscape:
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_width = portrait_height
            section.page_height = portrait_width
            section.top_margin = Cm(3.17)
            section.bottom_margin = Cm(3.17)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)
        else:
            section.orientation = WD_ORIENTATION.PORTRAIT
            section.page_width = portrait_width
            section.page_height = portrait_height
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)


def main():
    document = Document(str(SOURCE))
    normalize_sections(document)
    format_body_paragraphs(document)
    format_tables(document)
    document.save(str(OUTPUT))
    print(f"saved={OUTPUT}")


if __name__ == "__main__":
    main()
