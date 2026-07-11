from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "CET46" / "报告-latest-第九部分经验总结修订.docx"
OUTPUT = ROOT / "CET46" / "报告-latest-附录A模块接口明细页面格式修订.docx"


HEADERS = [
    "序号",
    "模块名称",
    "数据表",
    "功能点",
    "API接口",
    "请求方式",
    "对应页面",
    "前端开发",
    "后端开发",
]


PAGE_COMPONENTS = {
    "/login": "AuthLoginPage.vue",
    "/register": "AuthRegisterPage.vue",
    "/forgot-password": "AuthForgotPasswordPage.vue",
    "/onboarding": "AuthOnboardingPage.vue",
    "/academy/home": "AcademyHome.vue",
    "/academy/my-courses": "AcademyAggregatePage.vue",
    "/academy/open-courses": "AcademyOpenCourses.vue",
    "/academy/open-courses/:id": "AcademyCourseDetail.vue",
    "/academy/general-courses": "AcademyGeneralCourses.vue",
    "/academy/micro-majors": "AcademyMicroMajors.vue",
    "/academy/assignments": "AcademyAggregatePage.vue",
    "/academy/assignments/:assignmentId": "AcademyAssignmentDetail.vue",
    "/academy/exams": "AcademyAggregatePage.vue",
    "/academy/exams/:examId/take": "AcademyExamDetail.vue",
    "/academy/textbooks": "AcademyTextbooks.vue",
    "/academy/textbooks/:id": "AcademyTextbookDetail.vue",
    "/academy/textbook-cart": "AcademyTextbookCart.vue",
    "/academy/question-bank": "AcademyQuestionBank.vue",
    "/academy/question-bank/courses": "AcademyQuestionBankCourses.vue",
    "/academy/question-bank/courses/:courseCode": "AcademyQuestionBankCourseDetail.vue",
    "/academy/question-bank/mistakes": "AcademyQuestionBankMistakes.vue",
    "/academy/question-bank/favorites": "AcademyQuestionBankFavorites.vue",
    "/profile": "ProfilePage.vue",
    "/teacher-mailbox": "TeacherMailboxPage.vue",
    "/exchange": "ExchangeCenter.vue",
    "/exchange/vouchers": "MyVouchers.vue",
    "/lab/oj": "OjPlatform.vue",
    "/lab/well-log": "WellLogSimulation.vue",
    "/lab/production": "PetroleumSimulation.vue",
    "/games/:gameId": "GamePlatform.vue",
    "/admin": "AdminPage.vue",
    "全局AI宠物入口": "AiPetWidget.vue",
}


GROUPS = {
    1: ["1", "用户认证与个人中心模块", "users；password_reset_codes；profile_user_profiles；profile_learning_events；profile_learning_time_records"],
    7: ["2", "学堂课程模块", "learning_content_blocks；online_open_courses；general_courses；micro_major_courses；academy_course_enrollments；academy_course_reviews；teacher_published_courses；admin_course_categories"],
    16: ["3", "作业考试模块", "academy_assignments；academy_assignment_questions；academy_assignment_submissions；academy_exams；academy_exam_questions；academy_exam_submissions；oj_submissions"],
    21: ["4", "题库训练模块", "question_bank_subjects；question_bank_problems；course_question_bank_sets；course_question_bank_questions；course_question_bank_mistakes；course_question_bank_favorites"],
    25: ["5", "在线评测模块", "oj_problems；oj_categories；oj_test_cases；oj_submissions；oj_submission_cases"],
    27: ["6", "教材与兑换模块", "excellent_textbooks；academy_textbook_details；academy_textbook_cart_items；academy_textbook_orders；academy_textbook_reviews；voucher_items；user_vouchers；coin_spend_records"],
}


FUNCTION_ROWS = [
    ["用户注册", "/api/auth/register", "POST", "/register"],
    ["用户登录", "/api/auth/login", "POST", "/login"],
    ["找回密码验证码发送", "/api/auth/password-reset/code", "POST", "/forgot-password"],
    ["密码重置确认", "/api/auth/password-reset/confirm", "POST", "/forgot-password"],
    ["首次登录信息完善", "/api/auth/onboarding", "POST", "/onboarding"],
    ["AI宠物形象绑定", "/api/auth/pet", "PUT", "全局AI宠物入口"],
    ["学堂首页数据聚合", "/api/academy/home", "GET", "/academy/home"],
    ["我的课程聚合查询", "/api/academy/my-courses", "GET", "/academy/my-courses"],
    ["开放课程列表与分类筛选", "/api/academy/online-open-courses；/api/academy/online-open-courses/categories", "GET", "/academy/open-courses"],
    ["开放课程详情展示", "/api/academy/online-open-courses/{id}", "GET", "/academy/open-courses/:id"],
    ["教师发布开放课程", "/api/academy/online-open-courses", "POST", "/profile"],
    ["教师编辑开放课程", "/api/academy/online-open-courses/{id}", "PUT", "/profile"],
    ["教师删除开放课程", "/api/academy/online-open-courses/{id}", "DELETE", "/profile"],
    ["课程报名与取消报名", "/api/academy/online-open-courses/{id}/enroll", "POST/DELETE", "/academy/open-courses/:id"],
    ["课程评价查询、提交与回复", "/api/academy/online-open-courses/{id}/reviews；/api/academy/reviews/{reviewId}/reply", "GET/POST", "/academy/open-courses/:id；/teacher-mailbox"],
    ["作业列表与详情查询", "/api/academy/assignments；/api/academy/assignments/{assignmentCode}", "GET", "/academy/assignments"],
    ["作业草稿保存", "/api/academy/assignments/{assignmentCode}/draft", "POST", "/academy/assignments/:assignmentId"],
    ["作业提交与OJ题提交", "/api/academy/assignments/{assignmentCode}/submit", "POST", "/academy/assignments/:assignmentId"],
    ["考试列表与详情查询", "/api/academy/exams；/api/academy/exams/{examCode}", "GET", "/academy/exams"],
    ["考试开始、草稿保存与提交", "/api/academy/exams/{examCode}/start；/api/academy/exams/{examCode}/draft；/api/academy/exams/{examCode}/submit", "POST", "/academy/exams/:examId/take"],
    ["题库科目与题目列表", "/api/academy/question-bank/subjects；/api/academy/question-bank/problems", "GET", "/academy/question-bank"],
    ["课程题库目录与套题详情", "/api/academy/question-bank/course-catalog；/api/academy/question-bank/courses/{code}", "GET", "/academy/question-bank/courses"],
    ["错题统计、列表与作答记录", "/api/academy/question-bank/mistakes/summary；/api/academy/question-bank/mistakes；/api/academy/question-bank/mistakes/answers", "GET/POST", "/academy/question-bank/mistakes"],
    ["题目收藏、列表与取消收藏", "/api/academy/question-bank/favorites/summary；/api/academy/question-bank/favorites；/api/academy/question-bank/favorites/{questionId}", "GET/POST/DELETE", "/academy/question-bank/favorites"],
    ["OJ题目浏览与测试点查询", "/api/oj/problems；/api/oj/problems/categories；/api/oj/problems/{problemId}/test-cases", "GET", "/lab/oj"],
    ["代码提交判题与结果查询", "/api/oj/submissions；/api/oj/submissions/{id}；/api/oj/submissions/{id}/cases", "POST/GET", "/lab/oj"],
    ["教材列表、分类与详情展示", "/api/academy/textbooks；/api/academy/textbooks/categories；/api/academy/textbooks/{id}", "GET", "/academy/textbooks"],
    ["教材购物车查询、加入、修改与删除", "/api/academy/textbook-cart；/api/academy/textbook-cart/{itemId}", "GET/POST/PUT/DELETE", "/academy/textbook-cart"],
    ["教材订单创建与支付", "/api/academy/textbook-orders；/api/academy/textbook-orders/{orderNo}/pay", "POST", "/academy/textbook-cart"],
    ["金币卡券兑换与使用", "/api/rewards/vouchers；/api/rewards/vouchers/items；/api/rewards/vouchers/exchange；/api/rewards/vouchers/use", "GET/POST", "/exchange；/exchange/vouchers"],
]


INDIVIDUAL_ROWS = [
    ["7", "通识课程模块", "general_courses；academy_course_enrollments；admin_course_categories", "通识课程列表、分类与选课", "/api/academy/general-courses；/api/academy/general-courses/{id}/enroll", "GET/POST/DELETE", "/academy/general-courses"],
    ["8", "微专业模块", "micro_major_courses；academy_course_enrollments；admin_course_categories", "微专业列表、分类与报名", "/api/academy/micro-major-courses；/api/academy/micro-major-courses/{id}/enroll", "GET/POST/DELETE", "/academy/micro-majors"],
    ["9", "AI宠物模块", "users", "宠物对话与学习陪伴", "/api/ai-pet/chat", "POST", "全局AI宠物入口"],
    ["10", "教师工作台模块", "online_open_courses；academy_course_reviews；academy_assignments", "教师工作台概览", "/api/academy/teacher/workbench", "GET", "/profile"],
    ["11", "教师工作台模块", "academy_course_reviews", "教师信箱已读处理", "/api/academy/teacher/workbench/mailbox/read", "POST", "/teacher-mailbox"],
    ["12", "教师工作台模块", "academy_assignments；academy_assignment_questions", "教师布置作业", "/api/academy/teacher/assignments", "POST", "/profile"],
    ["13", "OJ题目管理模块", "oj_problems；oj_categories", "OJ题目创建与编辑", "/api/oj/problems；/api/oj/problems/{id}", "POST/PUT", "/profile"],
    ["14", "OJ题目管理模块", "oj_test_cases", "OJ测试用例维护", "/api/oj/problems/{problemId}/test-cases；/api/oj/problems/{problemId}/test-cases/{testCaseId}", "GET/POST/DELETE", "/profile"],
    ["15", "测井实验模块", "well_log_template", "测井模板列表与详情", "/api/well-log/template/list；/api/well-log/template/{id}", "GET", "/lab/well-log"],
    ["16", "测井实验模块", "well_log_record", "测井实验记录保存", "/api/well-log/record/save", "POST", "/lab/well-log"],
    ["17", "测井实验模块", "well_log_record", "测井记录分页、详情与删除", "/api/well-log/record/page；/api/well-log/record/{id}", "GET/DELETE", "/lab/well-log"],
    ["18", "采油工程实验模块", "production_pump_record", "抽油泵实验记录维护", "/api/production/pump/save；/api/production/pump/page；/api/production/pump/{id}", "POST/GET/DELETE", "/lab/production"],
    ["19", "采油工程实验模块", "production_reservoir_record", "油藏实验记录维护", "/api/production/reservoir/save；/api/production/reservoir/page；/api/production/reservoir/{id}", "POST/GET/DELETE", "/lab/production"],
    ["20", "采油工程实验模块", "production_waterflood_record", "注水开发实验记录维护", "/api/production/waterflood/save；/api/production/waterflood/page；/api/production/waterflood/{id}", "POST/GET/DELETE", "/lab/production"],
    ["21", "采油工程实验模块", "production_stimulation_record", "增产措施实验记录维护", "/api/production/stimulation/save；/api/production/stimulation/page；/api/production/stimulation/{id}", "POST/GET/DELETE", "/lab/production"],
    ["22", "阶梯跳跃游戏模块", "course_question_bank_sets；course_question_bank_questions", "游戏题库与题目获取", "/api/games/ladder-jump/question-banks；/api/games/ladder-jump/questions", "GET", "/games/:gameId"],
    ["23", "阶梯跳跃游戏模块", "game_ladder_jump_records", "游戏成绩记录保存", "/api/games/ladder-jump/records", "POST", "/games/:gameId"],
    ["24", "单词战士游戏模块", "question_bank_problems；game_type_warrior_records", "单词题目获取与成绩记录", "/api/academy/question-bank/type-warrior/words；/api/games/type-warrior/records", "GET/POST", "/games/:gameId"],
    ["25", "系统管理模块", "users", "用户列表、状态修改与删除", "/api/admin/users；/api/admin/users/{userId}", "GET/PUT/DELETE", "/admin"],
    ["26", "系统管理模块", "online_open_courses；general_courses；micro_major_courses；excellent_textbooks", "课程资源查询与新增", "/api/admin/courses", "GET/POST", "/admin"],
    ["27", "系统管理模块", "admin_course_categories；oj_categories", "课程分类与OJ分类维护", "/api/admin/course-categories；/api/admin/oj/categories", "GET/POST/DELETE", "/admin"],
    ["28", "系统管理模块", "online_open_courses；general_courses；micro_major_courses；excellent_textbooks", "课程资源删除", "/api/admin/courses/{resourceType}/{courseId}", "DELETE", "/admin"],
    ["29", "系统管理模块", "academy_course_reviews；academy_textbook_reviews", "评价查询、删除与回复", "/api/admin/reviews；/api/admin/reviews/{reviewType}/{reviewId}/reply", "GET/POST/DELETE", "/admin"],
    ["30", "系统管理模块", "course_question_bank_sets；course_question_bank_categories", "题库套题查询、新增与删除", "/api/admin/question-bank/sets；/api/admin/question-bank/sets/{setCode}", "GET/POST/DELETE", "/admin"],
    ["31", "系统管理模块", "course_question_bank_questions", "题库题目查询、新增、编辑与删除", "/api/admin/question-bank/questions；/api/admin/question-bank/questions/{questionId}", "GET/POST/PUT/DELETE", "/admin"],
    ["32", "系统管理模块", "oj_problems；oj_categories", "OJ题目查询、新增、编辑与删除", "/api/admin/oj/problems；/api/admin/oj/problems/{problemId}", "GET/POST/PUT/DELETE", "/admin"],
    ["33", "系统管理模块", "oj_problems；oj_test_cases", "OJ题目样例校验", "/api/admin/oj/problems/check", "POST", "/admin"],
    ["34", "系统管理模块", "voucher_items", "兑换商品查询、新增与删除", "/api/admin/vouchers；/api/admin/vouchers/{voucherKey}", "GET/POST/DELETE", "/admin"],
    ["35", "题库导入模块", "question_bank_problems；question_bank_tags；question_bank_problem_subjects", "洛谷题目导入", "/api/academy/question-bank/import/luogu", "POST", "/academy/question-bank"],
    ["36", "开放课程分类模块", "admin_course_categories；online_open_courses", "开放课程分类查询", "/api/academy/online-open-courses/categories", "GET", "/academy/open-courses"],
    ["37", "通识课程分类模块", "admin_course_categories；general_courses", "通识课程分类查询", "/api/academy/general-courses/categories", "GET", "/academy/general-courses"],
    ["38", "微专业分类模块", "admin_course_categories；micro_major_courses", "微专业分类查询", "/api/academy/micro-major-courses/categories", "GET", "/academy/micro-majors"],
    ["39", "教材分类模块", "admin_course_categories；excellent_textbooks", "教材分类查询", "/api/academy/textbooks/categories", "GET", "/academy/textbooks"],
    ["40", "教材评价模块", "academy_textbook_reviews；excellent_textbooks", "教材评价提交", "/api/academy/textbooks/{id}/reviews", "POST", "/academy/textbooks/:id"],
    ["41", "开放课程教师模块", "online_open_courses；teacher_published_courses", "教师本人课程查询", "/api/academy/online-open-courses/teacher/mine", "GET", "/profile"],
    ["42", "卡券资产模块", "user_vouchers；voucher_items", "我的卡券查询", "/api/rewards/vouchers", "GET", "/exchange/vouchers"],
    ["43", "个人主页模块", "users；profile_user_profiles", "个人主页概览查询", "/api/profile/overview", "GET", "/profile"],
    ["44", "个人主页模块", "users", "当前用户资料查询", "/api/profile/user", "GET", "/profile"],
    ["45", "个人主页模块", "users；profile_user_profiles", "头像与主页资料更新", "/api/profile/avatar", "POST", "/profile"],
    ["46", "学习记录模块", "profile_learning_events", "学习事件记录", "/api/profile/events", "POST", "/profile"],
    ["47", "学习记录模块", "profile_learning_time_records", "学习时长记录", "/api/profile/learning-time", "POST", "/profile"],
    ["48", "教材订单模块", "academy_textbook_orders；academy_textbook_order_items", "教材订单支付确认", "/api/academy/textbook-orders/{orderNo}/pay", "POST", "/academy/textbook-cart"],
]


def find_appendix_table(document):
    for table in document.tables:
        if len(table.rows) == 73 and len(table.columns) == 9:
            row = [cell.text.strip() for cell in table.rows[0].cells]
            if row == HEADERS:
                return table
    raise RuntimeError("未找到附录A模块及接口明细表。")


def set_cell_text(cell, text):
    tc = cell._tc
    tc_pr = tc.tcPr
    for child in list(tc):
        if child is not tc_pr:
            tc.remove(child)
    paragraph_element = OxmlElement("w:p")
    tc.append(paragraph_element)
    paragraph = cell.paragraphs[0]
    if not text:
        return
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def normalize_page_text(text):
    return "；".join(PAGE_COMPONENTS.get(item.strip(), item.strip()) for item in text.split("；"))


def main():
    document = Document(str(SOURCE))
    table = find_appendix_table(document)

    if len(FUNCTION_ROWS) != 30:
        raise RuntimeError(f"FUNCTION_ROWS数量应为30，实际为{len(FUNCTION_ROWS)}。")
    if len(INDIVIDUAL_ROWS) != 42:
        raise RuntimeError(f"INDIVIDUAL_ROWS数量应为42，实际为{len(INDIVIDUAL_ROWS)}。")

    for word_row_index, function_data in enumerate(FUNCTION_ROWS, start=1):
        cells = table.rows[word_row_index].cells
        if word_row_index in GROUPS:
            for col_index, value in enumerate(GROUPS[word_row_index]):
                set_cell_text(cells[col_index], value)
        for offset, value in enumerate(function_data, start=3):
            if offset == 6:
                value = normalize_page_text(value)
            set_cell_text(cells[offset], value)
        set_cell_text(cells[7], "")
        set_cell_text(cells[8], "")

    for word_row_index, row_data in enumerate(INDIVIDUAL_ROWS, start=31):
        cells = table.rows[word_row_index].cells
        for col_index, value in enumerate(row_data):
            if col_index == 6:
                value = normalize_page_text(value)
            set_cell_text(cells[col_index], value)
        set_cell_text(cells[7], "")
        set_cell_text(cells[8], "")

    document.save(str(OUTPUT))
    print(f"saved={OUTPUT}")


if __name__ == "__main__":
    main()
