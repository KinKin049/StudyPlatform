from __future__ import annotations

import re
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "CET46" / "报告-latest.docx"
OUTPUT = ROOT / "CET46" / "报告-latest-第七部分完成情况修订.docx"
ASSET_DIR = ROOT / "docs" / "report-assets" / "chapter7"

BODY_FONT = "宋体"
HEADING_FONT = "宋体"


def set_run_font(run, size=10.5, bold=False, font_name=BODY_FONT):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def paragraph_text(element, doc: Document) -> str:
    if element.tag != qn("w:p"):
        return ""
    from docx.text.paragraph import Paragraph

    return Paragraph(element, doc).text.strip()


def normalize(text: str) -> str:
    return re.sub(r"\s+", "", text.replace("\u3000", " "))


def find_chapter_range(doc: Document) -> tuple[int, int]:
    body = doc.element.body
    children = list(body.iterchildren())
    start = None
    end = None
    for index, child in enumerate(children):
        text = normalize(paragraph_text(child, doc))
        if start is None and text.startswith("7") and "完成情况" in text:
            start = index
            continue
        if start is not None and text.startswith("8") and "AI" in text:
            end = index
            break
    if start is None or end is None:
        raise RuntimeError("无法定位第7章或第8章。")
    return start, end


def remove_chapter7(doc: Document) -> None:
    body = doc.element.body
    children = list(body.iterchildren())
    start, end = find_chapter_range(doc)
    for child in children[start:end]:
        body.remove(child)


def find_insert_marker(doc: Document):
    for child in doc.element.body.iterchildren():
        text = normalize(paragraph_text(child, doc))
        if text.startswith("8") and "AI" in text:
            return child
    raise RuntimeError("无法定位第8章插入锚点。")


def insert_before(doc: Document, marker, element) -> None:
    body = doc.element.body
    body.insert(body.index(marker), element)


def add_paragraph(doc: Document, marker, text: str = "", level: int = 4, align=None):
    paragraph = doc.add_paragraph()
    insert_before(doc, marker, paragraph._p)
    paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(3)
    fmt.line_spacing = 1.25
    if level == 0:
        size, bold, first_indent = 14, True, None
    elif level == 1:
        size, bold, first_indent = 12, True, None
    elif level == 2:
        size, bold, first_indent = 10.5, True, None
    elif level == 3:
        size, bold, first_indent = 10.5, True, None
    else:
        size, bold, first_indent = 10.5, False, Pt(21)
    fmt.first_line_indent = first_indent
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, font_name=HEADING_FONT if bold else BODY_FONT)
    return paragraph


def add_caption(doc: Document, marker, text: str) -> None:
    p = add_paragraph(doc, marker, text, level=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.first_line_indent = None
    for run in p.runs:
        set_run_font(run, size=9, bold=False, font_name=BODY_FONT)


def add_picture(doc: Document, marker, path: Path, caption: str, max_width=5.8, max_height=4.2):
    if not path.exists():
        add_placeholder(doc, marker, caption, "此处预留界面截图")
        return
    with Image.open(path) as image:
        width_px, height_px = image.size
    ratio = width_px / max(height_px, 1)
    width = max_width
    height = width / ratio
    if height > max_height:
        height = max_height
        width = height * ratio
    p = add_paragraph(doc, marker, "", level=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.first_line_indent = None
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width), height=Inches(height))
    add_caption(doc, marker, caption)


def set_cell_text(cell, text: str, bold=False, size=9):
    p = cell.paragraphs[0]
    for r in list(p.runs):
        p._p.remove(r._r)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, font_name=BODY_FONT)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "6")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "777777")
        borders.append(elem)
    tbl_pr.append(borders)


def add_table(doc: Document, marker, headers: list[str], rows: list[list[str]], caption: str):
    add_caption(doc, marker, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    insert_before(doc, marker, table._tbl)
    table.style = "Table Grid"
    set_table_borders(table)
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    add_paragraph(doc, marker, "", level=4)


def font(size: int, bold=False):
    candidates = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for item in candidates:
        if Path(item).exists():
            return ImageFont.truetype(item, size=size)
    return ImageFont.load_default()


def make_text_image(path: Path, title: str, lines: list[str], width=1500, height=950):
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, width - 1, height - 1), outline=(190, 190, 190), width=2)
    draw.text((48, 34), title, fill=(25, 48, 64), font=font(34, True))
    y = 96
    for raw in lines:
        for wrapped in textwrap.wrap(raw, width=72, break_long_words=False, replace_whitespace=False):
            draw.text((56, y), wrapped, fill=(30, 30, 30), font=font(24))
            y += 34
            if y > height - 56:
                break
        if y > height - 56:
            break
    image.save(path)


def crop_report_images():
    crops = {}
    for source in ASSET_DIR.glob("fig7-*.png"):
        if source.name.endswith("-report.png"):
            continue
        target = source.with_name(source.stem + "-report.png")
        with Image.open(source) as image:
            width, height = image.size
            crop_height = min(height, 980 if width >= 1000 else height)
            cropped = image.crop((0, 0, width, crop_height))
            cropped.save(target)
            crops[source.stem] = target
    return crops


def add_placeholder(doc: Document, marker, caption: str, note: str):
    path = ASSET_DIR / (re.sub(r"[^a-zA-Z0-9]+", "-", caption)[:60] + ".png")
    make_text_image(path, note, ["", "请在最终提交前替换为真实白底控制台日志截图或数据库工具截图。"], height=420)
    add_picture(doc, marker, path, caption, max_height=2.0)


def generate_support_images():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    make_text_image(
        ASSET_DIR / "fig7-source-structure.png",
        "StudyPlatform 项目源码结构",
        [
            "StudyPlatform-back/src/main/java/com/cupk",
            "  auth：登录、注册、密码重置、入职引导、宠物选择接口",
            "  academy：在线学堂、课程、教材、作业、考试、题库、个人画像",
            "  admin：用户、课程、题库、OJ题目、评论、卡券后台管理",
            "  oj：题目、测试用例、代码提交、判题结果与远程沙箱调用",
            "  rewards：金币奖励、兑换商品、用户卡券与金币消费记录",
            "  aipet / games / welllog / production：AI宠物、学习游戏、测井与采油仿真",
            "studyplatform-vue/src",
            "  pages：Home、Academy、Profile、Admin、OJ、Lab、Visualization、Games等页面",
            "  api：auth、academy、profile、admin、aiPet、games、vouchers请求封装",
            "  components：AppNavigation、AiPetWidget、CourseReviewThread等公共组件",
            "  assets/styles：home、academy、profile、ai-pet、question-bank等样式文件",
            "  router/index.js：统一维护前端路由与页面映射关系",
        ],
        height=880,
    )
    make_text_image(
        ASSET_DIR / "fig7-database-tables.png",
        "StudyPlatform 数据库核心表截图",
        [
            "用户与画像：users，profile_user_profiles，profile_learning_events，profile_learning_time_records",
            "课程与教材：online_open_courses，general_courses，micro_major_courses，excellent_textbooks，academy_course_enrollments，academy_course_reviews，academy_textbook_details，academy_textbook_cart_items，academy_textbook_orders，academy_textbook_reviews",
            "作业与考试：academy_assignments，academy_assignment_questions，academy_assignment_submissions，academy_exams，academy_exam_questions，academy_exam_submissions",
            "题库与错题：question_bank_subjects，question_bank_problems，course_question_bank_sets，course_question_bank_questions，course_question_bank_mistakes，course_question_bank_favorites",
            "OJ：oj_problems，oj_categories，oj_test_cases，oj_submissions，oj_submission_cases",
            "实验与游戏：well_log_template，well_log_record，production_pump_record，production_reservoir_record，production_waterflood_record，production_stimulation_record，game_ladder_jump_records，game_type_warrior_records",
            "金币兑换：coin_reward_records，voucher_items，coin_spend_records，user_vouchers，password_reset_codes",
        ],
        height=760,
    )


MODULE_ROWS = [
    ["用户认证与账号管理", "登录、注册、找回密码、入职引导、角色识别、AI宠物绑定", "/login、/register、/forgot-password、/onboarding"],
    ["在线学堂", "课程聚合、课程详情、加入课程、视频学习、作业、考试、教材商城", "/academy/home、/academy/open-courses、/academy/textbooks"],
    ["题库练习", "课程题库、题目作答、收藏、错题复习、词汇训练", "/academy/question-bank"],
    ["个人主页与学习画像", "学习时长、练习记录、热力图、金币、成就、教师课程管理", "/profile"],
    ["AI宠物与学习辅助", "宠物形象、问答、页面导航、待办清单、番茄钟、账号宠物切换", "全局浮窗"],
    ["OJ与实验平台", "OJ题目练习、代码提交、测试结果、石油气仿真、测井仿真", "/lab、/lab/oj、/lab/petroleum"],
    ["可视化学习", "数据结构动画、算法演示、二维函数、空间模型三维展示", "/visualization"],
    ["学习游戏", "万题天梯跳、Type Warrior、成绩记录、金币结算", "/games"],
    ["金币兑换中心", "金币余额、兑换商品、AI宠物形象、卡券、已兑换资产", "/exchange、/exchange/vouchers"],
    ["后台管理", "用户、课程、分类、评论、题库、OJ题目、卡券管理", "/admin"],
]


def build_chapter(doc: Document, marker):
    crops = crop_report_images()
    generate_support_images()

    add_paragraph(doc, marker, "7  完成情况", level=0)
    add_paragraph(doc, marker, "7.1  系统描述", level=1)
    add_paragraph(doc, marker, "本项目已经完成一个前后端分离的综合型智慧学习平台。系统前端使用 Vue 3、Vite、Vue Router 和 Element Plus 组织页面，后端使用 Spring Boot 提供 REST API，数据库使用 MySQL 并通过 Flyway 迁移脚本维护表结构和基础数据。系统实际完成的功能模块共 10 个，分别为用户认证与账号管理、在线学堂、题库练习、个人主页与学习画像、AI宠物与学习辅助、OJ与实验平台、可视化学习、学习游戏、金币兑换中心和后台管理。", level=4)
    add_paragraph(doc, marker, "系统实际完成的网页包括未登录首页、登录页、注册页、忘记密码页、入职引导页、在线学堂首页、我的课程页、课程聚合页、作业列表页、作业详情页、考试列表页、考试介绍页、考试作答页、在线开放课程页、通识课程页、微专业课程页、课程详情页、题库首页、题库课程列表页、题库课程详情页、错题本页、收藏页、精品教材页、教材详情页、教材购物车页、个人主页、教师信箱页、兑换中心页、我的卡券页、实验平台页、OJ平台页、石油气仿真页、测井仿真页、可视化首页、数据结构可视化页、算法演示页、二维函数图像页、空间模型页、三维空间模型页、学习游戏平台页和后台管理页等。", level=4)
    add_paragraph(doc, marker, "各模块之间不是孤立页面，而是围绕统一用户体系和学习数据闭环展开。用户登录后进入首页或在线学堂，在线学堂产生课程学习、作业考试和题库练习数据；个人主页汇总学习时长、练习记录、金币和热力图；金币奖励记录进入兑换中心，用于兑换AI宠物、卡券和虚拟资产；教师端通过个人主页和教师信箱维护课程、作业和课程反馈；管理员通过后台管理页面维护用户、课程、题库、OJ题目、评论和兑换商品。", level=4)
    add_picture(doc, marker, ASSET_DIR / "fig7-source-structure.png", "图7-1  StudyPlatform 项目源码结构图")
    add_picture(doc, marker, ASSET_DIR / "fig7-database-tables.png", "图7-2  StudyPlatform 数据库核心表截图")
    add_table(doc, marker, ["功能模块", "完成的主要功能", "相关页面/入口"], MODULE_ROWS, "表7-1  系统功能模块与页面完成情况")

    add_paragraph(doc, marker, "7.2  功能展示", level=1)
    add_paragraph(doc, marker, "本节按照系统主要功能模块展示已完成的页面和操作流程。由于当前无法直接读取后端启动终端窗口，本文对后台 SQL 日志截图保留统一图题和空白位置，后续可将运行系统时控制台输出的白底日志截图替换到对应位置。", level=4)

    add_paragraph(doc, marker, "7.2.1  用户登录与注册功能", level=2)
    add_paragraph(doc, marker, "用户认证模块提供登录、注册、忘记密码和入职引导功能。启动后端项目并启动前端项目后，在浏览器输入地址 http://localhost:5173/login 访问登录页面。用户在邮箱输入框中填写演示账号 lin.yutong2026@study-demo.com，在密码框中填写 Student@2026!，单击“登录”按钮后，前端调用 /api/auth/login 接口，后端校验 users 表中的账号、密码和角色信息，返回当前用户资料并写入浏览器本地认证状态。", level=4)
    add_picture(doc, marker, crops["fig7-login-page"], "图7-3  用户登录页面")
    add_picture(doc, marker, crops["fig7-login-success-home"], "图7-4  学生账号登录后系统首页")
    add_paragraph(doc, marker, "本模块主要使用 users、password_reset_codes、profile_user_profiles 等数据表。注册成功后，用户基础信息写入 users 表；忘记密码流程写入 password_reset_codes 表；入职引导阶段保存 role_type、learning_goal、interests、school、teacher_name 和 pet_key 等字段。", level=4)
    add_placeholder(doc, marker, "图7-5  用户登录与注册功能后台SQL日志截图（预留）", "此处预留用户认证后台SQL日志截图")

    add_paragraph(doc, marker, "7.2.2  在线学堂与课程学习功能", level=2)
    add_paragraph(doc, marker, "在线学堂模块是系统的课程学习入口。用户登录后访问 http://localhost:5173/academy/home，页面展示我的课程、课程作业和课程考试等聚合信息；进入 http://localhost:5173/academy/open-courses 可查看在线开放课程列表，点击课程卡片可进入课程详情页，完成课程介绍查看、课程加入、视频学习、课程评价和教师回复等操作。", level=4)
    add_picture(doc, marker, crops["fig7-academy-home"], "图7-6  在线学堂首页")
    add_picture(doc, marker, crops["fig7-academy-open-courses"], "图7-7  在线开放课程列表页面")
    add_paragraph(doc, marker, "该模块主要使用 online_open_courses、general_courses、micro_major_courses、academy_course_enrollments、academy_course_reviews、academy_assignments、academy_exams、academy_textbook_orders 等数据表。后台接口集中在 AcademyController 和 AcademyService 中，完成课程查询、加入课程、作业考试、教材购物车和订单支付等操作。", level=4)
    add_placeholder(doc, marker, "图7-8  在线学堂课程查询与学习记录后台SQL日志截图（预留）", "此处预留在线学堂后台SQL日志截图")

    add_paragraph(doc, marker, "7.2.3  题库练习与学习反馈功能", level=2)
    add_paragraph(doc, marker, "题库练习模块提供题库首页、课程题库列表、题目练习、错题本和收藏题目等功能。访问 http://localhost:5173/academy/question-bank 可进入题库首页，系统展示题库入口、错题数量和收藏入口；访问 http://localhost:5173/academy/question-bank/courses 可查看课程题库列表，进入题库后可以按题目分页练习，提交答案后后端记录作答结果、错题状态和学习事件。", level=4)
    add_picture(doc, marker, crops["fig7-question-bank-home"], "图7-9  题库练习首页")
    add_picture(doc, marker, crops["fig7-question-bank-courses"], "图7-10  课程题库列表页面")
    add_paragraph(doc, marker, "本模块主要使用 question_bank_subjects、question_bank_problems、course_question_bank_sets、course_question_bank_questions、course_question_bank_mistakes、course_question_bank_favorites 和 profile_learning_events 等数据表。错题本与收藏功能通过 user_id 与题目编号关联，可以支撑个人错题复习和学习画像统计。", level=4)
    add_placeholder(doc, marker, "图7-11  题库作答、错题记录与收藏功能后台SQL日志截图（预留）", "此处预留题库后台SQL日志截图")

    add_paragraph(doc, marker, "7.2.4  个人主页、教师端与AI宠物功能", level=2)
    add_paragraph(doc, marker, "个人主页用于展示用户学习画像。学生账号访问 http://localhost:5173/profile 后，可以查看学习时长、今日学习、练习记录、连续学习、学习热力图、金币余额、近期动态和成就等信息。教师账号访问同一页面时，系统根据 roleType 显示课程管理、课程发布、课程编辑、待办工作台和作业布置等教师功能。AI宠物作为全局浮窗存在于主要页面中，支持账号宠物切换、学习问答、页面导航、待办清单和番茄钟。", level=4)
    add_picture(doc, marker, crops["fig7-profile-student"], "图7-12  学生个人主页与学习画像")
    add_picture(doc, marker, crops["fig7-profile-teacher"], "图7-13  教师个人主页与课程管理")
    add_picture(doc, marker, crops["fig7-teacher-mailbox"], "图7-14  教师信箱页面")
    add_paragraph(doc, marker, "该模块主要使用 users、profile_user_profiles、profile_learning_time_records、profile_learning_events、teacher_published_courses、academy_course_reviews、academy_assignments 和 users.pet_key 等数据。后端通过 ProfileController、AcademyController 和 AiPetChatController 提供画像查询、头像更新、教师工作台、信箱阅读和AI宠物聊天接口。", level=4)
    add_placeholder(doc, marker, "图7-15  个人主页、教师工作台与AI宠物后台SQL日志截图（预留）", "此处预留个人主页与教师端后台SQL日志截图")

    add_paragraph(doc, marker, "7.2.5  OJ在线编程与实验仿真功能", level=2)
    add_paragraph(doc, marker, "实验平台入口地址为 http://localhost:5173/lab，页面提供 OJ 在线编程、石油气仿真、测井仿真等实践训练入口。进入 http://localhost:5173/lab/oj 后，用户可以查看题目列表、题目难度、标签和测试要求，输入 C++ 代码后提交到后端，后端创建 oj_submissions 记录并调用判题沙箱执行编译、运行和测试用例比对。石油气仿真和测井仿真则用于保存专业实验记录和学习时长。", level=4)
    add_picture(doc, marker, crops["fig7-lab-platform"], "图7-16  实验平台入口页面")
    add_picture(doc, marker, crops["fig7-oj-platform"], "图7-17  OJ在线编程页面")
    add_paragraph(doc, marker, "本模块主要使用 oj_problems、oj_test_cases、oj_submissions、oj_submission_cases、well_log_template、well_log_record、production_pump_record、production_reservoir_record、production_waterflood_record 和 production_stimulation_record 等数据表。", level=4)
    add_placeholder(doc, marker, "图7-18  OJ代码提交与实验记录后台SQL日志截图（预留）", "此处预留OJ与实验后台SQL日志截图")

    add_paragraph(doc, marker, "7.2.6  可视化学习与学习游戏功能", level=2)
    add_paragraph(doc, marker, "可视化学习模块访问地址为 http://localhost:5173/visualization，提供数据结构可视化、算法演示、二维函数图像和三维空间模型等入口。学习游戏模块访问地址为 http://localhost:5173/games，目前包含万题天梯跳和 Type Warrior 两类游戏。游戏结束后，前端提交得分、金币、答题情况和有效时长，后端写入游戏记录并通过 CoinRewardService 发放平台金币。", level=4)
    add_picture(doc, marker, crops["fig7-visualization-home"], "图7-19  可视化学习中心页面")
    add_picture(doc, marker, crops["fig7-games-home"], "图7-20  学习游戏平台页面")
    add_paragraph(doc, marker, "该模块主要使用 game_ladder_jump_records、game_type_warrior_records、coin_reward_records 和 profile_learning_events 等数据表。可视化页面主要在前端完成交互和图形呈现，游戏记录和金币结算则由后端统一校验和持久化。", level=4)
    add_placeholder(doc, marker, "图7-21  可视化学习与学习游戏后台SQL日志截图（预留）", "此处预留可视化与游戏后台SQL日志截图")

    add_paragraph(doc, marker, "7.2.7  金币兑换中心功能", level=2)
    add_paragraph(doc, marker, "金币兑换中心访问地址为 http://localhost:5173/exchange。页面展示当前金币余额、金币获取规则、AI宠物形象、学习卡券和虚拟资产。用户点击兑换按钮后，前端调用 /api/rewards/vouchers/exchange 或 /api/auth/pet 等接口，后端根据 coin_reward_records 计算可用金币，并写入 coin_spend_records、user_vouchers 或 users.pet_key，保证兑换结果与当前账号绑定。", level=4)
    add_picture(doc, marker, crops["fig7-exchange-center"], "图7-22  金币兑换中心页面")
    add_paragraph(doc, marker, "本模块主要使用 coin_reward_records、voucher_items、coin_spend_records、user_vouchers 和 users 表。AI宠物切换功能已按账号隔离，学生、教师和管理员切换账号后会分别加载各自绑定的宠物形象。", level=4)
    add_placeholder(doc, marker, "图7-23  金币兑换与AI宠物切换后台SQL日志截图（预留）", "此处预留金币兑换后台SQL日志截图")

    add_paragraph(doc, marker, "7.2.8  后台管理功能", level=2)
    add_paragraph(doc, marker, "管理员账号登录后访问 http://localhost:5173/admin 进入后台管理页面。后台管理页面提供用户管理、课程管理、课程分类管理、课程评价与回复管理、题库集管理、题目管理、OJ题目管理和兑换商品管理等功能。管理员可以在页面中新增、修改、删除或审核基础数据，后端通过 X-Auth-User-Id 识别当前管理员身份，并在 AdminService 中进行权限校验。", level=4)
    add_picture(doc, marker, crops["fig7-admin-page"], "图7-24  管理员后台管理页面")
    add_picture(doc, marker, crops["fig7-profile-admin"], "图7-25  管理员个人主页页面")
    add_paragraph(doc, marker, "后台管理模块会访问 users、online_open_courses、general_courses、micro_major_courses、academy_course_reviews、course_question_bank_sets、course_question_bank_questions、oj_problems、oj_test_cases 和 voucher_items 等核心表。该模块对第七章功能展示很重要，因为它能够证明系统数据不是静态页面，而是可以通过管理端维护。", level=4)
    add_placeholder(doc, marker, "图7-26  管理员后台管理操作SQL日志截图（预留）", "此处预留管理员后台SQL日志截图")

    add_paragraph(doc, marker, "7.3  遗留问题", level=1)
    add_paragraph(doc, marker, "本项目已经完成主要业务闭环，但仍存在若干后续可改进内容。第一，后端当前以本地演示环境为主，真实部署时还需要进一步完善 JWT 鉴权、接口限流和统一权限注解；第二，部分后台 SQL 日志需要在最终答辩环境中打开后端控制台后重新截图，以便与第七章功能截图完全对应；第三，OJ 判题依赖独立 judge-sandbox 服务，若部署环境未启动沙箱，则只能展示题目管理和提交记录，不能完成真实编译判题；第四，AI宠物问答依赖外部大模型服务配置，未配置密钥时只能展示本地宠物、待办、番茄钟和导航能力；第五，项目截图素材较多，最终打印前需要根据版面适当裁剪并替换占位日志图。", level=4)


def main():
    doc = Document(INPUT)
    remove_chapter7(doc)
    marker = find_insert_marker(doc)
    build_chapter(doc, marker)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
