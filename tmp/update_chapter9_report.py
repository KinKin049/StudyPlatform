from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "CET46" / "报告-latest-第八部分AI辅助修订.docx"
OUTPUT = ROOT / "CET46" / "报告-latest-第九部分经验总结修订.docx"

BODY_FONT = "宋体"
HEADING_FONT = "宋体"


def set_run_font(run, size=10.5, bold=False, font_name=BODY_FONT):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def paragraph_text(element, doc: Document) -> str:
    if element.tag != qn("w:p"):
        return ""
    from docx.text.paragraph import Paragraph

    return Paragraph(element, doc).text.strip()


def normalize(text: str) -> str:
    return re.sub(r"\s+", "", text.replace("\u3000", " "))


def find_chapter_range(doc: Document) -> tuple[int, int]:
    children = list(doc.element.body.iterchildren())
    start = None
    end = None
    for index, child in enumerate(children):
        text = normalize(paragraph_text(child, doc))
        if start is None and text.startswith("9") and "经验总结" in text:
            start = index
            continue
        if start is not None and text.startswith("附录A"):
            end = index
            break
    if start is None or end is None:
        raise RuntimeError("无法定位第9章或附录A。")
    return start, end


def remove_chapter9(doc: Document) -> None:
    body = doc.element.body
    children = list(body.iterchildren())
    start, end = find_chapter_range(doc)
    for child in children[start:end]:
        body.remove(child)


def find_insert_marker(doc: Document):
    for child in doc.element.body.iterchildren():
        if normalize(paragraph_text(child, doc)).startswith("附录A"):
            return child
    raise RuntimeError("无法定位附录A插入锚点。")


def insert_before(doc: Document, marker, element) -> None:
    body = doc.element.body
    body.insert(body.index(marker), element)


def add_paragraph(doc: Document, marker, text: str = "", level: int = 4, align=None):
    paragraph = doc.add_paragraph()
    insert_before(doc, marker, paragraph._p)
    paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(3)
    fmt.line_spacing = 1.25
    if level == 0:
        size, bold, first_indent = 14, True, None
    elif level == 1:
        size, bold, first_indent = 12, True, None
    elif level == 2:
        size, bold, first_indent = 10.5, True, None
    else:
        size, bold, first_indent = 10.5, False, Pt(21)
    fmt.first_line_indent = first_indent
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, font_name=HEADING_FONT if bold else BODY_FONT)
    return paragraph


def build_chapter(doc: Document, marker):
    add_paragraph(doc, marker, "9  经验总结", level=0)
    add_paragraph(
        doc,
        marker,
        "本项目从最初的在线学习平台逐步扩展为包含在线学堂、题库练习、OJ在线编程、虚拟仿真、算法可视化、学习游戏、AI宠物、金币兑换和后台管理的综合型学习系统。项目规模不断扩大，使开发过程不只是完成页面和接口，更需要在需求取舍、数据一致性、功能闭环、测试验证和报告取证之间保持平衡。通过本项目实践，我对前后端分离项目的组织方式、复杂功能的拆分方式以及AI工具的合理使用边界都有了更具体的认识。",
        level=4,
    )

    add_paragraph(doc, marker, "9.1  项目管理方面的经验与教训", level=1)
    add_paragraph(
        doc,
        marker,
        "在项目管理方面，最深的体会是综合型系统必须尽早划清模块边界。本项目包含的功能较多，如果一开始只从“页面效果”出发，很容易不断增加新入口和新功能，却忽略不同模块之间的数据关系。后期在整理报告第五章、第六章和第七章时可以明显感受到，在线学堂、题库、个人主页、金币奖励、AI宠物和后台管理并不是独立功能，而是围绕同一个用户体系和学习行为数据形成闭环。因此，今后开发类似项目时，应先明确核心主线，例如“用户登录—课程学习—练习反馈—学习画像—金币激励—兑换展示”，再安排附加功能的开发顺序。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "第二个经验是演示数据和测试账号应尽早准备。前期开发时更关注功能是否能运行，到了第七部分功能展示阶段才发现，学生端、教师端和管理员端都需要足够真实的数据支撑截图。如果学生账号的热力图、学习时长、金币、课程记录为空，教师端信箱、课程管理、作业布置为空，后台管理没有可维护数据，即使功能已经实现，报告展示效果仍然会显得不完整。本项目后期专门补充了三类展示账号和较完整的演示数据，这提高了报告可信度，也说明测试数据本身就是项目交付的一部分。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "第三个教训是报告素材不能等到最后集中补。第七部分要求功能截图和后台日志相互对应，实际操作时才发现页面截图、数据库截图、源码结构截图和后端日志截图都需要提前规划。尤其是后台日志，如果开发阶段没有统一日志格式，后期就很难快速获得符合要求的白底截图。因此，今后开发时应当边实现功能边保留关键运行证据，包括接口请求、数据库变化、控制台输出和页面状态，减少最后集中整理报告的压力。",
        level=4,
    )

    add_paragraph(doc, marker, "9.2  技术实现方面的经验与教训", level=1)
    add_paragraph(
        doc,
        marker,
        "在技术实现方面，前后端字段一致性非常重要。本项目采用Vue 3和Spring Boot前后端分离架构，页面状态依赖后端接口返回。开发中曾遇到登录字段、用户角色字段、宠物标识、课程资源类型等信息前后端理解不一致的问题。类似问题往往不会导致页面完全打不开，而是表现为某个模块数据不刷新、账号切换后状态不正确或页面展示回退到默认值。AI宠物账号切换问题就是典型例子：用户表中已经存在pet_key字段，但前端组件长期读取全局localStorage，导致切换账号时宠物没有跟随用户变化。这个问题说明，与账号相关的数据应尽量以数据库和后端接口为准，浏览器缓存只能作为辅助状态。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "数据库设计方面，Flyway迁移脚本对多人协作和版本追踪很有帮助。项目中用户、课程、题库、作业考试、OJ、实验记录、游戏记录、金币奖励和兑换商品等表结构逐步增加，如果没有迁移脚本管理，很难判断某个字段或某张表是在什么时候加入的。与此同时，迁移脚本数量增加后也带来新的要求：版本号必须连续清楚，表之间的外键关系和唯一约束要谨慎设计，演示数据脚本不能破坏真实业务表结构。通过第六部分数据库设计的整理，我更加理解“概念结构、逻辑结构、物理表结构”之间应保持一致，而不是只在文档中描述得完整。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "金币奖励和OJ判题等模块让我认识到业务幂等性的重要性。金币不能只依赖前端计算，否则用户重复提交或刷新页面可能造成重复奖励；OJ提交也不能只展示前端结果，而应由后端生成提交记录，再由判题服务返回状态和测试用例结果。本项目通过coin_reward_records中的来源类型和来源键约束奖励记录，通过oj_submissions和oj_submission_cases保存提交与用例结果，使功能更接近真实系统。这个经验说明，涉及积分、成绩、订单和评测结果的模块，必须把关键判断放在后端，并通过数据库约束保证结果可追溯。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "前端实现方面，我体会到复杂页面要重视状态归属。个人主页、教师课程管理、兑换中心和AI宠物都同时依赖当前登录用户、页面路由、后端数据和本地缓存。如果状态来源不清晰，功能越多越容易出现看似细小但影响体验的问题。例如兑换中心既要展示金币余额，又要处理宠物拥有状态、当前宠物和兑换花费；教师端既要展示课程，又要维护课程编辑、作业布置和信箱消息。今后应尽量把与用户身份有关的数据统一封装，减少多个页面重复读取和写入本地缓存。",
        level=4,
    )

    add_paragraph(doc, marker, "9.3  今后项目开发的改进建议", level=1)
    add_paragraph(
        doc,
        marker,
        "第一，后续应进一步完善统一权限体系。本项目当前已经能够区分学生、教师和管理员角色，但在真实部署环境中还需要引入更严格的鉴权机制，例如JWT令牌、角色拦截、接口权限注解和敏感操作审计。这样可以避免仅依赖前端页面入口判断权限，也能提高后台管理和教师端操作的安全性。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "第二，应完善自动化测试和接口测试。项目模块较多，手动测试可以覆盖主要演示流程，但难以及时发现跨模块影响。今后可以为登录注册、课程加入、题库作答、金币结算、兑换商品、教师发布课程、OJ提交和后台管理等核心流程建立测试用例，并在修改接口字段或数据库结构后自动执行。这样能减少后期为了报告展示才发现功能不一致的情况。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "第三，应进一步规范日志和运行取证。本项目后期已经补充业务请求日志，但如果从开发初期就设计好日志格式，例如统一输出模块、操作、路径、用户ID、状态码和耗时，并在关键数据库操作处保留SQL或业务摘要，第七部分功能展示会更加顺利。对于课程设计类项目，日志不仅是调试工具，也是证明系统真实运行的重要材料。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "第四，部署方式可以继续改进。目前项目本地运行需要分别启动前端、后端、数据库和判题沙箱，步骤较多。后续可以考虑使用统一启动脚本或Docker Compose组织前端、后端、MySQL和judge-sandbox服务，使环境初始化、数据库迁移和端口配置更加稳定。这样既方便答辩演示，也方便后续维护和迁移到其他电脑。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "第五，个性化数据应更多后端化。AI宠物、学习热力图、金币兑换、游戏记录和学习时长都与用户长期画像有关，不适合长期依赖浏览器缓存。今后应继续把这类数据沉淀到数据库中，通过后端接口按用户ID读取和更新，保证不同设备、不同账号之间的数据不会混淆。",
        level=4,
    )

    add_paragraph(doc, marker, "9.4  AI工具使用感受", level=1)
    add_paragraph(
        doc,
        marker,
        "通过本项目开发，我认为AI工具更适合作为辅助解释、辅助编程和辅助调试工具，而不是直接替代开发者完成项目。AI能够较快解释错误信息、给出代码写法参考、帮助梳理模块关系，也能在报告写作中提供表达建议。但它的输出并不天然等于正确答案，尤其涉及本项目具体字段、接口路径、数据库表结构和页面状态时，必须回到源码、数据库和运行结果中核对。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "本项目中，AI工具对提高调试效率有一定帮助。例如面对后端启动失败、接口字段不匹配、账号切换状态不同步等问题时，AI可以帮助列出可能原因，使排查方向更清楚。但真正解决问题仍需要自己阅读代码、理解业务流程、修改实现并重新运行验证。这个过程让我认识到，AI可以降低查资料和组织思路的成本，但不能代替对项目整体结构的理解。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "在报告写作方面，AI工具可以帮助调整文字逻辑和表述方式，但也容易生成看似完整却与项目实际不完全一致的内容。因此，本项目报告后期反复强调结合项目文件、页面截图、数据库迁移脚本和后端接口进行核对。我的体会是，合理使用AI并不意味着把成果交给AI完成，而是把AI当作一个可以讨论思路的工具，最终判断、取舍和验证仍然应由开发者完成。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "总体来看，本项目最大的收获不是完成了某一个单独功能，而是经历了从需求分析、数据库设计、前后端开发、调试修复、演示数据准备到报告整理的完整过程。这个过程让我更加理解软件项目的质量来自持续验证和细节管理：页面要能展示，接口要能返回，数据库要能支撑，日志要能证明，文档要与实际一致。今后继续开发类似系统时，我会更早规划核心流程、测试数据和运行证据，使项目实现和项目报告能够同步推进。",
        level=4,
    )


def main():
    doc = Document(INPUT)
    remove_chapter9(doc)
    marker = find_insert_marker(doc)
    build_chapter(doc, marker)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
