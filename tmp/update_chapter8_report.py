from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "CET46" / "报告-latest-第七部分完成情况修订.docx"
OUTPUT = ROOT / "CET46" / "报告-latest-第八部分AI辅助修订.docx"

BODY_FONT = "宋体"
HEADING_FONT = "宋体"


def set_run_font(run, size=10.5, bold=False, font_name=BODY_FONT):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def paragraph_text(element, doc: Document) -> str:
    if element.tag != qn("w:p"):
        return ""
    from docx.text.paragraph import Paragraph

    return Paragraph(element, doc).text.strip()


def normalize(text: str) -> str:
    return re.sub(r"\s+", "", text.replace("\u3000", " "))


def find_chapter_range(doc: Document) -> tuple[int, int]:
    children = list(doc.element.body.iterchildren())
    start = None
    end = None
    for index, child in enumerate(children):
        text = normalize(paragraph_text(child, doc))
        if start is None and text.startswith("8") and "AI" in text:
            start = index
            continue
        if start is not None and text.startswith("9") and "经验总结" in text:
            end = index
            break
    if start is None or end is None:
        raise RuntimeError("无法定位第8章或第9章。")
    return start, end


def remove_chapter8(doc: Document) -> None:
    body = doc.element.body
    children = list(body.iterchildren())
    start, end = find_chapter_range(doc)
    for child in children[start:end]:
        body.remove(child)


def find_insert_marker(doc: Document):
    for child in doc.element.body.iterchildren():
        text = normalize(paragraph_text(child, doc))
        if text.startswith("9") and "经验总结" in text:
            return child
    raise RuntimeError("无法定位第9章插入锚点。")


def insert_before(doc: Document, marker, element) -> None:
    body = doc.element.body
    body.insert(body.index(marker), element)


def add_paragraph(doc: Document, marker, text: str = "", level: int = 4, align=None):
    paragraph = doc.add_paragraph()
    insert_before(doc, marker, paragraph._p)
    paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(3)
    fmt.line_spacing = 1.25
    if level == 0:
        size, bold, first_indent = 14, True, None
    elif level == 1:
        size, bold, first_indent = 12, True, None
    elif level == 2:
        size, bold, first_indent = 10.5, True, None
    else:
        size, bold, first_indent = 10.5, False, Pt(21)
    fmt.first_line_indent = first_indent
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, font_name=HEADING_FONT if bold else BODY_FONT)
    return paragraph


def add_caption(doc: Document, marker, text: str) -> None:
    p = add_paragraph(doc, marker, text, level=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.first_line_indent = None
    for run in p.runs:
        set_run_font(run, size=9, bold=False, font_name=BODY_FONT)


def set_cell_text(cell, text: str, bold=False, size=9):
    p = cell.paragraphs[0]
    for run in list(p.runs):
        p._p.remove(run._r)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, font_name=BODY_FONT)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "777777")
        borders.append(element)
    tbl_pr.append(borders)


def add_table(doc: Document, marker, headers: list[str], rows: list[list[str]], caption: str):
    add_caption(doc, marker, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    insert_before(doc, marker, table._tbl)
    table.style = "Table Grid"
    set_table_borders(table)
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
    add_paragraph(doc, marker, "", level=4)


AI_STAGE_ROWS = [
    ["1", "需求分析", "0%", "需求分析由本人根据课程要求、项目目标和同类学习平台功能调研独立完成，未使用AI工具参与需求确定。"],
    ["2", "数据库设计", "15%", "辅助检查用户、课程、题库、作业考试、OJ提交、实验记录、金币奖励和兑换记录等数据对象之间的关系，提供表结构说明文字的表达建议。"],
    ["3", "编码开发", "15%", "辅助理解Vue组件、Spring Boot接口、前后端请求封装和业务逻辑实现方式；在局部功能开发中参考AI给出的代码示例，再结合项目已有目录、接口和数据结构进行人工修改、整合和运行验证。"],
    ["4", "Bug调试", "25%", "辅助分析后端启动失败、接口字段不匹配、页面数据不刷新、账号切换状态不同步等问题的可能原因，并结合控制台报错和源码定位问题。"],
    ["5", "测试", "15%", "辅助列出登录注册、课程浏览、题库练习、OJ提交、金币兑换、教师课程管理和后台管理等功能的测试关注点。"],
    ["6", "文档撰写", "10%", "辅助优化少量报告文字表达、章节衔接和图表说明，使报告语言更加规范；最终内容均根据项目实际源码、数据库结构和页面截图人工核对。"],
    ["7", "系统部署", "5%", "主要由本人独立完成本地环境启动、前后端联调、数据库迁移和运行验证，AI仅用于解释少量环境报错信息。"],
]


def add_case(doc: Document, marker, title: str, content: str):
    add_paragraph(doc, marker, title, level=2)
    add_paragraph(doc, marker, content, level=4)


def build_chapter(doc: Document, marker):
    add_paragraph(doc, marker, "8  AI辅助", level=0)
    add_paragraph(doc, marker, "本章说明本课程项目开发过程中AI工具的使用情况。整体上，AI工具仅作为辅助学习、辅助编程和辅助排错工具使用，项目需求分析、功能取舍、核心代码整合、系统运行验证、截图采集和最终报告修改均由本人结合项目实际情况完成。", level=4)

    add_paragraph(doc, marker, "8.1  AI工具介绍", level=1)
    add_paragraph(doc, marker, "本项目开发过程中使用的AI工具主要包括ChatGPT/Codex和DeepSeek等文本生成与代码辅助工具。使用方式主要是在遇到技术概念不清楚、代码结构需要理解、报错信息需要分析或报告文字需要调整时，向AI工具输入问题描述、错误信息或局部代码片段，获得解释、排查方向或表达建议。", level=4)
    add_paragraph(doc, marker, "从使用范围看，AI工具没有参与本项目的需求确定，也没有直接替代项目开发。需求分析阶段由本人根据课程设计要求、项目目标和同类学习平台功能进行人工分析。AI工具更多用于后续开发过程中的辅助编程、错误分析、代码理解、测试点整理和少量报告文字润色。对于AI给出的代码示例和解释内容，均结合项目实际目录结构、接口实现、数据库表结构和运行结果进行人工修改与验证。", level=4)
    add_paragraph(doc, marker, "AI工具的优点是能够较快解释技术问题、提供排查思路和帮助整理文字表达，有助于提高资料理解和调试效率；不足是生成内容可能与项目实际代码不完全一致，部分接口名称、字段设计或业务流程需要人工核对。因此，本项目没有直接采用AI输出作为最终结果，而是将其作为参考材料，并通过实际运行和代码检查确认后再使用。", level=4)

    add_paragraph(doc, marker, "8.2  AI工具主要用途", level=1)
    add_paragraph(doc, marker, "1、说明AI参与了哪些环节。", level=4)
    add_table(doc, marker, ["序号", "项目阶段", "AI占比", "具体用途"], AI_STAGE_ROWS, "表8-1  AI工具参与环节说明")

    add_paragraph(doc, marker, "2、如何使用的？具体应用案例如下。", level=4)
    add_case(
        doc,
        marker,
        "（1）代码结构理解",
        "在开发过程中，AI工具主要用于辅助理解Vue组件、前端路由、Spring Boot Controller、Service和Repository之间的调用关系。例如，在分析在线学堂模块时，结合AcademyPage.vue、AcademyController.java和AcademyService.java等文件，请AI解释页面请求如何进入后端接口、后端如何查询数据库并返回前端。最终代码理解和报告表述以项目源码为依据。"
    )
    add_case(
        doc,
        marker,
        "（2）辅助编程",
        "在局部功能开发时，AI工具用于提供代码写法参考，例如前端表单状态管理、接口请求封装、账号状态同步、后端参数校验和错误提示处理等。AI给出的代码不会直接复制到项目中，而是先对照项目原有命名、目录结构、接口字段和数据库字段进行修改，再通过前端构建、后端编译和实际页面操作验证功能是否正确。"
    )
    add_case(
        doc,
        marker,
        "（3）Bug调试",
        "当项目出现后端启动失败、接口请求异常、字段名称不匹配、页面状态未刷新或账号切换后数据未同步等问题时，AI工具用于辅助分析可能原因。例如根据控制台错误信息判断问题可能来自端口占用、数据库迁移、请求体字段、localStorage状态或后端接口校验。最终修复仍通过阅读源码、修改代码和重新运行系统完成。"
    )
    add_case(
        doc,
        marker,
        "（4）测试点整理",
        "在功能测试阶段，AI工具用于辅助列出需要关注的测试点，包括用户登录注册是否成功、课程列表是否能加载、题库作答是否记录错题、OJ题目是否能展示和提交、金币兑换是否扣减余额、教师端课程管理是否能新增和编辑课程、管理员后台是否能维护基础数据等。实际测试结果以浏览器运行页面和后端接口返回为准。"
    )
    add_case(
        doc,
        marker,
        "（5）文档表达优化",
        "在报告撰写阶段，AI工具用于少量文字润色和表达优化，例如帮助调整模块描述、接口说明、表格文字和图题措辞，使报告语言更加规范。第八部分以外的报告内容均结合项目实际源码、数据库结构、运行页面和截图进行人工核对，避免出现与项目实现不一致的描述。"
    )
    add_case(
        doc,
        marker,
        "（6）图表说明辅助",
        "在整理系统结构图、数据库说明图和功能展示图时，AI工具用于辅助检查图题是否清楚、表格表达是否简洁。图表内容本身仍以实际项目文件、Flyway迁移脚本、页面截图和运行结果为依据。"
    )

    add_paragraph(doc, marker, "总体来说，AI工具在本项目中起到辅助解释、辅助编程、辅助调试和辅助表达的作用。项目的需求分析、功能设计取舍、代码整合、数据库验证、运行测试和最终报告审定均由本人完成。", level=4)


def main():
    doc = Document(INPUT)
    remove_chapter8(doc)
    marker = find_insert_marker(doc)
    build_chapter(doc, marker)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
