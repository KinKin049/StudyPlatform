from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
doc_path = ROOT / "CET46" / "报告-latest-第五部分结构修订.docx"
doc = Document(doc_path)

def text(p):
    return "".join(run.text for run in p.runs).strip()

lines = []
in_five = False
for p in doc.paragraphs:
    t = text(p)
    if t.startswith("5") and "系统详细设计" in t:
        in_five = True
    if in_five and t.startswith("6") and "数据库设计" in t:
        break
    if in_five and t:
        if (
            t.startswith("5  ")
            or t.startswith("5.")
            or t.startswith("图5-")
            or t.startswith("基本描述")
            or t.startswith("主要需求")
            or t.startswith("建设目的意义")
            or t.startswith("建设主要内容")
            or t.startswith("模块特点")
        ):
            lines.append(t[:180])

print(f"paragraphs={len(doc.paragraphs)}")
for line in lines:
    print(line)
