from pathlib import Path

from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT / "CET46" / "报告-latest-第五部分结构修订.docx"
OUT = ROOT / "CET46" / "报告-latest-第五部分流程图专业修订.docx"
FLOW_DIR = ROOT / ".codex-tmp" / "professional_flowcharts"
FLOW_DIR.mkdir(parents=True, exist_ok=True)


def load_font(size, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


FONT_TITLE = load_font(42, bold=True)
FONT_LANE = load_font(22, bold=True)
FONT_BOX = load_font(21, bold=True)
FONT_SMALL = load_font(18)
FONT_NOTE = load_font(16)


def text_size(draw, text, font):
    bbox = draw.textbbox((0, 0), text, font=font)
    return bbox[2] - bbox[0], bbox[3] - bbox[1]


def wrap_text(draw, text, font, max_width):
    lines = []
    for raw in text.split("\n"):
        line = ""
        for char in raw:
            trial = line + char
            width, _ = text_size(draw, trial, font)
            if width <= max_width or not line:
                line = trial
            else:
                lines.append(line)
                line = char
        if line:
            lines.append(line)
    return lines


def draw_text_center(draw, box, text, font=FONT_BOX):
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, x2 - x1 - 30)
    line_h = font.size + 7
    total = line_h * len(lines)
    y = y1 + (y2 - y1 - total) / 2
    for line in lines:
        width, _ = text_size(draw, line, font)
        draw.text((x1 + (x2 - x1 - width) / 2, y), line, font=font, fill="#111111")
        y += line_h


def arrow(draw, start, end, label=None, elbow=None):
    points = [start]
    if elbow:
        points.extend(elbow)
    points.append(end)
    draw.line(points, fill="#111111", width=3, joint="curve")
    if len(points) >= 2:
        sx, sy = points[-2]
        ex, ey = points[-1]
    else:
        sx, sy = start
        ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex >= sx else -1
        head = [(ex, ey), (ex - direction * 13, ey - 7), (ex - direction * 13, ey + 7)]
    else:
        direction = 1 if ey >= sy else -1
        head = [(ex, ey), (ex - 7, ey - direction * 13), (ex + 7, ey - direction * 13)]
    draw.polygon(head, fill="#111111")
    if label:
        lx = (points[-2][0] + end[0]) / 2
        ly = (points[-2][1] + end[1]) / 2 - 24
        w, h = text_size(draw, label, FONT_SMALL)
        draw.rectangle((lx - w / 2 - 5, ly - 2, lx + w / 2 + 5, ly + h + 2), fill="#ffffff")
        draw.text((lx - w / 2, ly), label, font=FONT_SMALL, fill="#111111")


def process_box(draw, center, text, w=210, h=76, rounded=True):
    cx, cy = center
    box = (cx - w / 2, cy - h / 2, cx + w / 2, cy + h / 2)
    if rounded:
        draw.rounded_rectangle(box, radius=15, fill="#ffffff", outline="#111111", width=3)
    else:
        draw.rectangle(box, fill="#ffffff", outline="#111111", width=3)
    draw_text_center(draw, box, text)
    return box


def decision_box(draw, center, text, w=190, h=96):
    cx, cy = center
    pts = [(cx, cy - h / 2), (cx + w / 2, cy), (cx, cy + h / 2), (cx - w / 2, cy)]
    draw.polygon(pts, fill="#ffffff", outline="#111111")
    draw.line([pts[0], pts[1], pts[2], pts[3], pts[0]], fill="#111111", width=3)
    draw_text_center(draw, (cx - w / 2 + 12, cy - h / 2 + 8, cx + w / 2 - 12, cy + h / 2 - 8), text)
    return (cx - w / 2, cy - h / 2, cx + w / 2, cy + h / 2)


def terminator(draw, center, text, w=170, h=68):
    return process_box(draw, center, text, w=w, h=h, rounded=True)


def box_anchor(box, side):
    x1, y1, x2, y2 = box
    return {
        "left": (x1, (y1 + y2) / 2),
        "right": (x2, (y1 + y2) / 2),
        "top": ((x1 + x2) / 2, y1),
        "bottom": ((x1 + x2) / 2, y2),
        "center": ((x1 + x2) / 2, (y1 + y2) / 2),
    }[side]


def draw_lanes(draw, lanes, width, top=105, lane_h=152, left=50, right=50):
    for idx, lane in enumerate(lanes):
        y1 = top + idx * lane_h
        y2 = y1 + lane_h - 22
        draw.rounded_rectangle((left, y1, width - right, y2), radius=12, fill="#ffffff", outline="#777777", width=2)
        draw.text((left + 16, y1 + 14), lane, font=FONT_LANE, fill="#111111")
    return [top + i * lane_h + (lane_h - 22) / 2 for i in range(len(lanes))]


def save_chart(filename, title, lanes, draw_content, note):
    width = 1720
    lane_h = 158
    top = 118
    height = top + len(lanes) * lane_h + 92
    img = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(img)
    tw, _ = text_size(draw, title, FONT_TITLE)
    draw.text(((width - tw) / 2, 32), title, font=FONT_TITLE, fill="#000000")
    centers = draw_lanes(draw, lanes, width, top=top, lane_h=lane_h)
    draw_content(draw, centers)
    note_lines = wrap_text(draw, note, FONT_NOTE, width - 260)
    y = height - 64
    for line in note_lines[:2]:
        lw, _ = text_size(draw, line, FONT_NOTE)
        draw.text(((width - lw) / 2, y), line, font=FONT_NOTE, fill="#555555")
        y += 22
    path = FLOW_DIR / filename
    img.save(path, quality=95)
    return path


def chart_academy():
    lanes = ["用户学习入口层", "课程与教材服务层", "业务校验与学习过程层", "数据写入与反馈层"]

    def content(draw, c):
        b1 = terminator(draw, (145, c[0]), "学习者")
        b2 = process_box(draw, (365, c[0]), "选择在线学堂\n课程/教材")
        b3 = process_box(draw, (600, c[0]), "分类/关键词\n检索资源")
        b4 = process_box(draw, (850, c[0]), "查看课程详情\n或教材详情")
        b5 = decision_box(draw, (1120, c[0]), "选择学习\n或购买?")
        b6 = process_box(draw, (1395, c[0]), "发起加入课程\n观看/下单请求")

        s1 = process_box(draw, (365, c[1]), "学习接口\n接收请求")
        s2 = decision_box(draw, (625, c[1]), "资源类型\n是否合法")
        s3 = process_box(draw, (895, c[1]), "查询课程/教材\n分类/评价/视频")
        s4 = process_box(draw, (1175, c[1]), "封装统一\n响应对象")
        s5 = process_box(draw, (1450, c[1]), "前端渲染\n列表与详情")

        p1 = process_box(draw, (390, c[2]), "课程存在性\n校验")
        p2 = decision_box(draw, (670, c[2]), "是否为\n内部资源")
        p3 = process_box(draw, (940, c[2]), "创建报名/订单\n或学习记录")
        p4 = process_box(draw, (1215, c[2]), "视频学习时长\n分段上报")
        p5 = process_box(draw, (1480, c[2]), "作业考试\n草稿/提交")

        d1 = process_box(draw, (310, c[3]), "写入课程报名\n或教材订单")
        d2 = process_box(draw, (575, c[3]), "写入学习时长\n记录表")
        d3 = process_box(draw, (860, c[3]), "金币服务\n计算学习奖励")
        d4 = process_box(draw, (1145, c[3]), "个人画像\n更新统计")
        d5 = terminator(draw, (1425, c[3]), "返回学习结果\n与页面反馈", w=210)

        for a, b in [(b1, b2), (b2, b3), (b3, b4), (b4, b5), (b5, b6)]:
            arrow(draw, box_anchor(a, "right"), box_anchor(b, "left"))
        arrow(draw, box_anchor(b3, "bottom"), box_anchor(s1, "top"), elbow=[(600, c[0] + 68), (365, c[0] + 68)])
        arrow(draw, box_anchor(s1, "right"), box_anchor(s2, "left"))
        arrow(draw, box_anchor(s2, "right"), box_anchor(s3, "left"), "是")
        arrow(draw, box_anchor(s3, "right"), box_anchor(s4, "left"))
        arrow(draw, box_anchor(s4, "right"), box_anchor(s5, "left"))
        arrow(draw, box_anchor(s2, "bottom"), box_anchor(p1, "top"), "否/提交时复核", elbow=[(625, c[1] + 70), (390, c[1] + 70)])
        for a, b in [(p1, p2), (p2, p3), (p3, p4), (p4, p5)]:
            arrow(draw, box_anchor(a, "right"), box_anchor(b, "left"), "是" if a == p2 else None)
        arrow(draw, box_anchor(p2, "bottom"), (670, c[3] - 60), "否")
        arrow(draw, (670, c[3] - 60), box_anchor(d5, "left"), elbow=[(670, c[3] - 60), (1320, c[3] - 60)])
        for a, b in [(p3, d1), (p4, d2), (d2, d3), (d3, d4), (d4, d5)]:
            if a == p3:
                arrow(draw, box_anchor(a, "bottom"), box_anchor(d1, "top"), elbow=[(940, c[2] + 72), (310, c[2] + 72)])
            elif a == p4:
                arrow(draw, box_anchor(a, "bottom"), box_anchor(d2, "top"), elbow=[(1215, c[2] + 72), (575, c[2] + 72)])
            else:
                arrow(draw, box_anchor(a, "right"), box_anchor(b, "left"))

    return save_chart(
        "professional_5_1_academy.png",
        "在线学堂模块详细设计流程图",
        lanes,
        content,
        "说明：课程资源通过资源类型与资源编号定位；视频学习、作业考试、教材订单最终进入个人画像和金币激励数据闭环。",
    )


def chart_question_bank():
    lanes = ["练习入口与题源层", "前端答题交互层", "题库服务与状态判定层", "画像与金币反馈层"]

    def content(draw, c):
        b1 = terminator(draw, (140, c[0]), "学习者")
        b2 = process_box(draw, (360, c[0]), "进入题库首页\n错题/收藏")
        b3 = process_box(draw, (600, c[0]), "选择题库集合\n或课程题库")
        b4 = process_box(draw, (850, c[0]), "加载分页题目\n与掌握状态")
        b5 = process_box(draw, (1110, c[0]), "展示题干\n选项/解析")
        b6 = process_box(draw, (1390, c[0]), "提交答案\n或收藏题目")

        f1 = decision_box(draw, (360, c[1]), "是否有筛选\n关键词/状态")
        f2 = process_box(draw, (620, c[1]), "组合分页\n筛选参数")
        f3 = process_box(draw, (875, c[1]), "调用题库 API")
        f4 = decision_box(draw, (1140, c[1]), "题型是否\n可自动判分")
        f5 = process_box(draw, (1415, c[1]), "前端生成\n作答结果")

        s1 = process_box(draw, (330, c[2]), "题库控制器\n接收请求")
        s2 = process_box(draw, (585, c[2]), "数据仓库查询\n题目/错题/收藏")
        s3 = decision_box(draw, (850, c[2]), "答案是否正确\n或词汇已掌握")
        s4 = process_box(draw, (1115, c[2]), "更新错题状态\n收藏状态")
        s5 = process_box(draw, (1400, c[2]), "写入学习事件\n和学习时长")

        d1 = process_box(draw, (335, c[3]), "画像服务\n沉淀行为")
        d2 = process_box(draw, (600, c[3]), "金币奖励服务\n匹配规则")
        d3 = decision_box(draw, (865, c[3]), "来源键是否\n已奖励")
        d4 = process_box(draw, (1135, c[3]), "写入金币流水\n奖励记录表")
        d5 = terminator(draw, (1420, c[3]), "刷新正确率\n错题数/金币", w=215)

        for a, b in [(b1, b2), (b2, b3), (b3, b4), (b4, b5), (b5, b6)]:
            arrow(draw, box_anchor(a, "right"), box_anchor(b, "left"))
        arrow(draw, box_anchor(b3, "bottom"), box_anchor(f1, "top"), elbow=[(600, c[0] + 70), (360, c[0] + 70)])
        for a, b in [(f1, f2), (f2, f3), (f3, f4), (f4, f5)]:
            arrow(draw, box_anchor(a, "right"), box_anchor(b, "left"), "是" if a == f1 else None)
        arrow(draw, box_anchor(f1, "bottom"), (360, c[2] - 64), "否")
        arrow(draw, (360, c[2] - 64), box_anchor(s1, "top"))
        for a, b in [(f3, s1), (s1, s2), (s2, s3), (s3, s4), (s4, s5)]:
            if a == f3:
                arrow(draw, box_anchor(a, "bottom"), box_anchor(s1, "top"), elbow=[(875, c[1] + 72), (330, c[1] + 72)])
            else:
                arrow(draw, box_anchor(a, "right"), box_anchor(b, "left"), "是" if a == s3 else None)
        arrow(draw, box_anchor(s3, "bottom"), (850, c[3] - 60), "否/仅记录")
        arrow(draw, (850, c[3] - 60), box_anchor(d1, "top"), elbow=[(850, c[3] - 60), (335, c[3] - 60)])
        for a, b in [(s5, d1), (d1, d2), (d2, d3), (d3, d4), (d4, d5)]:
            if a == s5:
                arrow(draw, box_anchor(a, "bottom"), box_anchor(d1, "top"), elbow=[(1400, c[2] + 72), (335, c[2] + 72)])
            else:
                arrow(draw, box_anchor(a, "right"), box_anchor(b, "left"), "否" if a == d3 else None)
        arrow(draw, box_anchor(d3, "bottom"), box_anchor(d5, "left"), "是/跳过", elbow=[(865, c[3] + 62), (1305, c[3] + 62)])

    return save_chart(
        "professional_5_2_question_bank.png",
        "题库练习与金币激励模块详细设计流程图",
        lanes,
        content,
        "说明：题库模块以答题行为和学习时长为输入，错题、收藏、画像统计和金币流水在后端统一落库。",
    )


def chart_profile_admin():
    lanes = ["统一账号入口层", "角色功能分派层", "权限校验与业务处理层", "数据聚合与结果反馈层"]

    def content(draw, c):
        b1 = terminator(draw, (145, c[0]), "用户")
        b2 = process_box(draw, (360, c[0]), "登录/注册\n找回密码")
        b3 = process_box(draw, (595, c[0]), "读取 authUser\n与 roleType")
        b4 = decision_box(draw, (850, c[0]), "角色类型")
        b5 = process_box(draw, (1120, c[0]), "学生/教师/\n管理员入口")
        b6 = process_box(draw, (1400, c[0]), "进入对应页面")

        r1 = process_box(draw, (295, c[1]), "学生：个人主页\n学习画像")
        r2 = process_box(draw, (565, c[1]), "教师：课程发布\n工作台/信箱")
        r3 = process_box(draw, (835, c[1]), "管理员：后台\n用户/课程/题库")
        r4 = process_box(draw, (1110, c[1]), "管理员：OJ\n卡券/评价管理")
        r5 = process_box(draw, (1400, c[1]), "提交查询\n或维护请求")

        s1 = process_box(draw, (295, c[2]), "画像服务\n聚合学习数据")
        s2 = decision_box(draw, (565, c[2]), "是否通过\n教师校验")
        s3 = decision_box(draw, (835, c[2]), "是否通过\n管理员校验")
        s4 = process_box(draw, (1110, c[2]), "校验参数\n规范化表单")
        s5 = process_box(draw, (1400, c[2]), "数据仓库\n写入/查询")

        d1 = process_box(draw, (320, c[3]), "学习时长/事件\n游戏/错题/订单")
        d2 = process_box(draw, (610, c[3]), "教师发布课程\n课程评价信箱")
        d3 = process_box(draw, (905, c[3]), "用户/课程/题库\nOJ/卡券表")
        d4 = process_box(draw, (1195, c[3]), "封装统计指标\n管理结果")
        d5 = terminator(draw, (1465, c[3]), "页面刷新\n展示反馈", w=190)

        for a, b in [(b1, b2), (b2, b3), (b3, b4), (b4, b5), (b5, b6)]:
            arrow(draw, box_anchor(a, "right"), box_anchor(b, "left"))
        arrow(draw, box_anchor(b4, "bottom"), box_anchor(r1, "top"), "学生", elbow=[(850, c[0] + 70), (295, c[0] + 70)])
        arrow(draw, box_anchor(b5, "bottom"), box_anchor(r2, "top"), "教师", elbow=[(1120, c[0] + 70), (565, c[0] + 70)])
        arrow(draw, box_anchor(b6, "bottom"), box_anchor(r3, "top"), "管理员", elbow=[(1400, c[0] + 70), (835, c[0] + 70)])
        for a, b in [(r3, r4), (r4, r5)]:
            arrow(draw, box_anchor(a, "right"), box_anchor(b, "left"))
        arrow(draw, box_anchor(r1, "bottom"), box_anchor(s1, "top"))
        arrow(draw, box_anchor(r2, "bottom"), box_anchor(s2, "top"))
        arrow(draw, box_anchor(r3, "bottom"), box_anchor(s3, "top"))
        arrow(draw, box_anchor(r5, "bottom"), box_anchor(s4, "top"), elbow=[(1400, c[1] + 72), (1110, c[1] + 72)])
        arrow(draw, box_anchor(s2, "right"), box_anchor(s4, "left"), "是")
        arrow(draw, box_anchor(s3, "right"), box_anchor(s4, "left"), "是")
        arrow(draw, box_anchor(s4, "right"), box_anchor(s5, "left"))
        arrow(draw, box_anchor(s2, "bottom"), (565, c[3] - 60), "否/拒绝")
        arrow(draw, box_anchor(s3, "bottom"), (835, c[3] - 60), "否/拒绝")
        for a, b in [(s1, d1), (s5, d3), (d1, d4), (d2, d4), (d3, d4), (d4, d5)]:
            if a in (s1, s5):
                arrow(draw, box_anchor(a, "bottom"), box_anchor(b, "top"))
            else:
                arrow(draw, box_anchor(a, "right"), box_anchor(b, "left"))
        arrow(draw, (565, c[3] - 60), box_anchor(d5, "left"), elbow=[(565, c[3] - 60), (1345, c[3] - 60)])
        arrow(draw, (835, c[3] - 60), box_anchor(d5, "left"), elbow=[(835, c[3] - 50), (1345, c[3] - 50)])

    return save_chart(
        "professional_5_3_profile_admin.png",
        "用户画像、教师端与后台管理模块详细设计流程图",
        lanes,
        content,
        "说明：该模块围绕 roleType 做角色分派，教师和管理员写操作必须先经过后端权限校验。",
    )


def chart_ai_pet():
    lanes = ["页面上下文采集层", "宠物交互意图层", "AI服务处理层", "本地状态与学习反馈层"]

    def content(draw, c):
        b1 = terminator(draw, (145, c[0]), "用户")
        b2 = process_box(draw, (365, c[0]), "打开学习页面")
        b3 = process_box(draw, (610, c[0]), "AiPetWidget\n读取路由")
        b4 = process_box(draw, (870, c[0]), "采集标题/文本\n选中内容")
        b5 = decision_box(draw, (1140, c[0]), "页面是否\n隐藏宠物")
        b6 = process_box(draw, (1410, c[0]), "显示/隐藏\n宠物浮窗")

        i1 = process_box(draw, (305, c[1]), "用户输入\n问题或指令")
        i2 = decision_box(draw, (570, c[1]), "是否本地\n可执行动作")
        i3 = process_box(draw, (835, c[1]), "导航/待办\n番茄钟处理")
        i4 = process_box(draw, (1110, c[1]), "组装聊天请求\nhistory+context")
        i5 = process_box(draw, (1400, c[1]), "调用宠物\n聊天接口")

        s1 = process_box(draw, (320, c[2]), "AI聊天服务\n校验密钥")
        s2 = process_box(draw, (590, c[2]), "限制消息长度\n截取历史")
        s3 = process_box(draw, (860, c[2]), "构建系统提示词\n和页面上下文")
        s4 = process_box(draw, (1135, c[2]), "调用模型接口\n/v1/chat/completions")
        s5 = decision_box(draw, (1410, c[2]), "返回内容\n是否有效")

        d1 = process_box(draw, (320, c[3]), "本地存储\n保存待办/番茄钟")
        d2 = process_box(draw, (600, c[3]), "更新宠物状态\n思考/开心/专注")
        d3 = process_box(draw, (880, c[3]), "展示回答\n或错误提示")
        d4 = process_box(draw, (1160, c[3]), "辅助跳转\n学习入口")
        d5 = terminator(draw, (1430, c[3]), "形成学习陪伴\n与操作反馈", w=215)

        for a, b in [(b1, b2), (b2, b3), (b3, b4), (b4, b5), (b5, b6)]:
            arrow(draw, box_anchor(a, "right"), box_anchor(b, "left"), "否" if a == b5 else None)
        arrow(draw, box_anchor(b5, "bottom"), box_anchor(b6, "bottom"), "是/隐藏", elbow=[(1140, c[0] + 70), (1410, c[0] + 70)])
        arrow(draw, box_anchor(b6, "bottom"), box_anchor(i1, "top"), elbow=[(1410, c[0] + 72), (305, c[0] + 72)])
        for a, b in [(i1, i2), (i2, i3), (i2, i4), (i4, i5)]:
            if a == i2 and b == i3:
                arrow(draw, box_anchor(a, "right"), box_anchor(b, "left"), "是")
            elif a == i2:
                arrow(draw, box_anchor(a, "bottom"), box_anchor(b, "left"), "否", elbow=[(570, c[1] + 66), (760, c[1] + 66)])
            else:
                arrow(draw, box_anchor(a, "right"), box_anchor(b, "left"))
        arrow(draw, box_anchor(i5, "bottom"), box_anchor(s1, "top"), elbow=[(1400, c[1] + 72), (320, c[1] + 72)])
        for a, b in [(s1, s2), (s2, s3), (s3, s4), (s4, s5)]:
            arrow(draw, box_anchor(a, "right"), box_anchor(b, "left"))
        arrow(draw, box_anchor(i3, "bottom"), box_anchor(d1, "top"), elbow=[(835, c[1] + 72), (320, c[1] + 72)])
        arrow(draw, box_anchor(s5, "bottom"), box_anchor(d3, "top"), "是", elbow=[(1410, c[2] + 72), (880, c[2] + 72)])
        arrow(draw, box_anchor(s5, "right"), box_anchor(d3, "right"), "否", elbow=[(1530, c[2]), (1530, c[3])])
        for a, b in [(d1, d2), (d2, d3), (d3, d4), (d4, d5)]:
            arrow(draw, box_anchor(a, "right"), box_anchor(b, "left"))

    return save_chart(
        "professional_5_4_ai_pet.png",
        "AI宠物与学习辅助模块详细设计流程图",
        lanes,
        content,
        "说明：聊天回答依赖当前页面上下文；导航、待办、番茄钟属于前端可执行的本地交互。",
    )


def chart_visual_lab_oj():
    lanes = ["实践学习入口层", "前端交互与渲染层", "后端评测/记录层", "学习数据回流层"]

    def content(draw, c):
        b1 = terminator(draw, (145, c[0]), "学习者")
        b2 = process_box(draw, (365, c[0]), "进入可视化\n或实验平台")
        b3 = decision_box(draw, (610, c[0]), "选择实践\n类型")
        b4 = process_box(draw, (855, c[0]), "可视化学习\n函数/空间")
        b5 = process_box(draw, (1115, c[0]), "石油仿真\n测井/油藏")
        b6 = process_box(draw, (1375, c[0]), "OJ题目\n代码评测")

        f1 = process_box(draw, (855, c[1]), "图表与三维\n交互渲染")
        f2 = process_box(draw, (1115, c[1]), "仿真面板\n参数计算展示")
        f3 = process_box(draw, (1375, c[1]), "OJ页面\n编辑并提交")
        f4 = process_box(draw, (1555, c[1]), "轮询状态\n展示用例结果", w=190)

        s1 = process_box(draw, (360, c[2]), "学习时长上报\n可视化/实验")
        s2 = process_box(draw, (665, c[2]), "提交服务\n保存代码")
        s3 = process_box(draw, (935, c[2]), "判题服务\n读取题目和用例")
        s4 = decision_box(draw, (1200, c[2]), "沙箱地址\n是否配置")
        s5 = process_box(draw, (1470, c[2]), "判题沙箱\n编译运行")

        d1 = process_box(draw, (335, c[3]), "写入学习时长\n记录表")
        d2 = process_box(draw, (640, c[3]), "更新提交状态\n得分/耗时/内存")
        d3 = process_box(draw, (930, c[3]), "写入单用例\n评测结果")
        d4 = process_box(draw, (1210, c[3]), "金币服务\n按时长奖励")
        d5 = terminator(draw, (1480, c[3]), "个人画像\n展示成果", w=190)

        arrow(draw, box_anchor(b1, "right"), box_anchor(b2, "left"))
        arrow(draw, box_anchor(b2, "right"), box_anchor(b3, "left"))
        arrow(draw, box_anchor(b3, "right"), box_anchor(b4, "left"), "可视化")
        arrow(draw, box_anchor(b4, "right"), box_anchor(b5, "left"), "仿真")
        arrow(draw, box_anchor(b5, "right"), box_anchor(b6, "left"), "OJ")

        arrow(draw, box_anchor(b4, "bottom"), box_anchor(f1, "top"))
        arrow(draw, box_anchor(b5, "bottom"), box_anchor(f2, "top"))
        arrow(draw, box_anchor(b6, "bottom"), box_anchor(f3, "top"))
        arrow(draw, box_anchor(f3, "right"), box_anchor(f4, "left"))

        arrow(draw, box_anchor(f1, "bottom"), box_anchor(s1, "top"), elbow=[(855, c[1] + 72), (360, c[1] + 72)])
        arrow(draw, box_anchor(f2, "bottom"), box_anchor(s1, "top"), elbow=[(1115, c[1] + 62), (360, c[1] + 62)])
        arrow(draw, box_anchor(f3, "bottom"), box_anchor(s2, "top"), elbow=[(1375, c[1] + 72), (665, c[1] + 72)])

        arrow(draw, box_anchor(s2, "right"), box_anchor(s3, "left"))
        arrow(draw, box_anchor(s3, "right"), box_anchor(s4, "left"))
        arrow(draw, box_anchor(s4, "right"), box_anchor(s5, "left"), "是")
        arrow(draw, box_anchor(s4, "bottom"), box_anchor(d2, "top"), "否/答案模式", elbow=[(1200, c[2] + 72), (640, c[2] + 72)])

        arrow(draw, box_anchor(s1, "bottom"), box_anchor(d1, "top"))
        arrow(draw, box_anchor(s5, "bottom"), box_anchor(d2, "top"), elbow=[(1470, c[2] + 72), (640, c[2] + 72)])
        arrow(draw, box_anchor(d2, "right"), box_anchor(d3, "left"))
        arrow(draw, box_anchor(d1, "right"), box_anchor(d4, "left"))
        arrow(draw, box_anchor(d3, "right"), box_anchor(d4, "left"))
        arrow(draw, box_anchor(d4, "right"), box_anchor(d5, "left"))

    return save_chart(
        "professional_5_5_visual_lab_oj.png",
        "可视化、实验与OJ模块详细设计流程图",
        lanes,
        content,
        "说明：可视化和仿真主要在前端渲染并记录学习时长；OJ通过后端提交记录和判题沙箱完成评测闭环。",
    )


def chart_games_exchange():
    lanes = ["游戏学习入口层", "成绩采集与结算层", "金币与兑换服务层", "资产发放与展示层"]

    def content(draw, c):
        b1 = terminator(draw, (145, c[0]), "学习者")
        b2 = process_box(draw, (365, c[0]), "进入游戏\n学习平台")
        b3 = decision_box(draw, (625, c[0]), "选择游戏")
        b4 = process_box(draw, (900, c[0]), "万题天梯跳\n题库答题")
        b5 = process_box(draw, (1175, c[0]), "打字战士\n单词输入战斗")
        b6 = process_box(draw, (1450, c[0]), "生成成绩\n和结算面板")

        r1 = process_box(draw, (305, c[1]), "提交游戏\n成绩记录")
        r2 = process_box(draw, (575, c[1]), "游戏记录服务\n非负校验")
        r3 = process_box(draw, (845, c[1]), "写入游戏记录\nladder/type")
        r4 = decision_box(draw, (1120, c[1]), "奖励来源")
        r5 = process_box(draw, (1400, c[1]), "计算平台金币\ncoins/score")

        s1 = process_box(draw, (315, c[2]), "金币奖励服务\n写入流水")
        s2 = process_box(draw, (590, c[2]), "兑换中心读取\n余额与商品")
        s3 = decision_box(draw, (865, c[2]), "金币/库存\n是否充足")
        s4 = process_box(draw, (1135, c[2]), "扣减金币\n减少库存")
        s5 = process_box(draw, (1415, c[2]), "发放卡券\n用户资产表")

        d1 = process_box(draw, (320, c[3]), "个人主页更新\n游戏指标")
        d2 = process_box(draw, (600, c[3]), "兑换中心刷新\n余额/商品状态")
        d3 = process_box(draw, (880, c[3]), "我的卡券\n展示资产")
        d4 = process_box(draw, (1160, c[3]), "教材订单\n可使用优惠券")
        d5 = terminator(draw, (1430, c[3]), "完成激励\n反馈闭环", w=200)

        for a, b in [(b1, b2), (b2, b3), (b3, b4), (b3, b5), (b4, b6), (b5, b6)]:
            if a == b3:
                arrow(draw, box_anchor(a, "right"), box_anchor(b, "left"), "答题" if b == b4 else "打字")
            else:
                arrow(draw, box_anchor(a, "right"), box_anchor(b, "left"))
        arrow(draw, box_anchor(b6, "bottom"), box_anchor(r1, "top"), elbow=[(1450, c[0] + 72), (305, c[0] + 72)])
        for a, b in [(r1, r2), (r2, r3), (r3, r4), (r4, r5)]:
            arrow(draw, box_anchor(a, "right"), box_anchor(b, "left"), "天梯/Type" if a == r4 else None)
        arrow(draw, box_anchor(r5, "bottom"), box_anchor(s1, "top"), elbow=[(1400, c[1] + 72), (315, c[1] + 72)])
        for a, b in [(s1, s2), (s2, s3), (s3, s4), (s4, s5)]:
            arrow(draw, box_anchor(a, "right"), box_anchor(b, "left"), "是" if a == s3 else None)
        arrow(draw, box_anchor(s3, "bottom"), box_anchor(d2, "top"), "否/提示失败", elbow=[(865, c[2] + 72), (600, c[2] + 72)])
        arrow(draw, box_anchor(s1, "bottom"), box_anchor(d1, "top"))
        arrow(draw, box_anchor(s5, "bottom"), box_anchor(d3, "top"), elbow=[(1415, c[2] + 72), (880, c[2] + 72)])
        for a, b in [(d1, d2), (d2, d3), (d3, d4), (d4, d5)]:
            arrow(draw, box_anchor(a, "right"), box_anchor(b, "left"))

    return save_chart(
        "professional_5_6_games_exchange.png",
        "学习游戏与兑换中心模块详细设计流程图",
        lanes,
        content,
        "说明：游戏成绩经后端校验后进入金币流水；兑换中心再把金币转化为卡券、装扮、主题等用户资产。",
    )


chart_paths = {
    "图5-1  在线学堂模块流程图": chart_academy(),
    "图5-2  题库练习与金币激励模块流程图": chart_question_bank(),
    "图5-3  用户画像、教师端与后台管理模块流程图": chart_profile_admin(),
    "图5-4  AI宠物与学习辅助模块流程图": chart_ai_pet(),
    "图5-5  可视化、实验与OJ模块流程图": chart_visual_lab_oj(),
    "图5-6  学习游戏与兑换中心模块流程图": chart_games_exchange(),
}


def paragraph_text(paragraph):
    return "".join(run.text for run in paragraph.runs).strip()


def replace_previous_image(paragraph, image_path):
    paragraph.clear()
    paragraph.alignment = 1
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.7))


doc = Document(BASE)
replaced = 0
for idx, paragraph in enumerate(doc.paragraphs):
    text = paragraph_text(paragraph)
    if text in chart_paths:
        for prev_idx in range(idx - 1, max(-1, idx - 6), -1):
            prev = doc.paragraphs[prev_idx]
            if prev._p.xpath(".//w:drawing"):
                replace_previous_image(prev, chart_paths[text])
                replaced += 1
                break

doc.save(OUT)
print(OUT)
print(f"replaced={replaced}")
for title, path in chart_paths.items():
    print(f"{title}: {path}")
