from copy import deepcopy
from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / ".codex-tmp"))
from inspect_database_schema import parse_schema  # noqa: E402

SRC = ROOT / "CET46" / "报告-latest-第五部分流程图无遮挡修订.docx"
OUT = ROOT / "CET46" / "报告-latest-第六部分数据库结构修订.docx"


TABLE_ZH = {
    "users": "用户信息表",
    "learning_content_blocks": "学习内容板块表",
    "online_open_courses": "在线开放课程表",
    "general_courses": "通识课程表",
    "micro_major_courses": "微专业课程表",
    "excellent_textbooks": "精品教材表",
    "teacher_published_courses": "教师发布课程表",
    "academy_course_enrollments": "课程加入记录表",
    "academy_course_reviews": "课程评价表",
    "admin_course_categories": "后台课程分类表",
    "course_question_bank_categories": "课程题库分类表",
    "course_question_bank_sets": "课程题库套题表",
    "course_question_bank_questions": "课程题库题目表",
    "course_question_bank_mistakes": "课程题库错题表",
    "course_question_bank_favorites": "课程题库收藏表",
    "question_bank_subjects": "通用题源学科表",
    "question_bank_tags": "通用题源标签表",
    "question_bank_problems": "通用题源题目表",
    "question_bank_problem_subjects": "题源题目学科关联表",
    "academy_assignments": "课程作业表",
    "academy_assignment_questions": "作业题目表",
    "academy_assignment_submissions": "作业提交表",
    "academy_exams": "课程考试表",
    "academy_exam_questions": "考试题目表",
    "academy_exam_submissions": "考试提交表",
    "oj_categories": "OJ分类表",
    "oj_problems": "OJ题目表",
    "oj_test_cases": "OJ测试用例表",
    "oj_submissions": "OJ提交表",
    "oj_submission_cases": "OJ单用例结果表",
    "academy_textbook_details": "教材详情表",
    "academy_textbook_cart_items": "教材购物车表",
    "academy_textbook_orders": "教材订单表",
    "academy_textbook_order_items": "教材订单明细表",
    "academy_textbook_reviews": "教材评价表",
    "profile_user_profiles": "用户画像表",
    "profile_learning_events": "学习事件表",
    "profile_learning_time_records": "学习时长记录表",
    "password_reset_codes": "找回密码验证码表",
    "coin_reward_records": "金币奖励记录表",
    "coin_spend_records": "金币支出记录表",
    "user_vouchers": "用户卡券表",
    "voucher_items": "兑换卡券商品表",
    "game_ladder_jump_records": "万题天梯跳记录表",
    "game_type_warrior_records": "Type Warrior记录表",
    "well_log_template": "测井模板表",
    "well_log_record": "测井解释记录表",
    "production_pump_record": "抽油泵仿真记录表",
    "production_reservoir_record": "油藏动态仿真记录表",
    "production_waterflood_record": "注水开发仿真记录表",
    "production_stimulation_record": "压裂酸化仿真记录表",
}

GROUPS = [
    ("用户认证与学习画像", ["users", "profile_user_profiles", "profile_learning_events", "profile_learning_time_records", "password_reset_codes"]),
    ("课程资源与教师发布", ["learning_content_blocks", "online_open_courses", "general_courses", "micro_major_courses", "teacher_published_courses", "academy_course_enrollments", "academy_course_reviews", "admin_course_categories"]),
    ("题库、作业与考试", ["course_question_bank_categories", "course_question_bank_sets", "course_question_bank_questions", "course_question_bank_mistakes", "course_question_bank_favorites", "question_bank_subjects", "question_bank_tags", "question_bank_problems", "question_bank_problem_subjects", "academy_assignments", "academy_assignment_questions", "academy_assignment_submissions", "academy_exams", "academy_exam_questions", "academy_exam_submissions"]),
    ("OJ在线评测", ["oj_categories", "oj_problems", "oj_test_cases", "oj_submissions", "oj_submission_cases"]),
    ("教材商城与评价", ["excellent_textbooks", "academy_textbook_details", "academy_textbook_cart_items", "academy_textbook_orders", "academy_textbook_order_items", "academy_textbook_reviews"]),
    ("金币奖励、卡券与游戏", ["coin_reward_records", "coin_spend_records", "voucher_items", "user_vouchers", "game_ladder_jump_records", "game_type_warrior_records"]),
    ("测井与采油仿真实验", ["well_log_template", "well_log_record", "production_pump_record", "production_reservoir_record", "production_waterflood_record", "production_stimulation_record"]),
]

LOGICAL_REFS = {
    ("profile_user_profiles", "user_id"): "users",
    ("academy_course_enrollments", "user_id"): "users",
    ("academy_course_enrollments", "course_id"): "online_open_courses/general_courses/micro_major_courses",
    ("academy_course_reviews", "user_id"): "users",
    ("academy_course_reviews", "reply_user_id"): "users",
    ("academy_course_reviews", "parent_review_id"): "academy_course_reviews",
    ("academy_course_reviews", "course_id"): "online_open_courses/general_courses/micro_major_courses",
    ("academy_assignments", "course_id"): "online_open_courses/general_courses/micro_major_courses",
    ("academy_assignment_submissions", "user_id"): "users",
    ("academy_exams", "course_id"): "online_open_courses/general_courses/micro_major_courses",
    ("academy_exam_submissions", "user_id"): "users",
    ("academy_exam_questions", "oj_problem_id"): "oj_problems",
    ("academy_assignment_questions", "oj_problem_id"): "oj_problems",
    ("academy_textbook_details", "textbook_id"): "excellent_textbooks",
    ("academy_textbook_cart_items", "textbook_id"): "excellent_textbooks",
    ("academy_textbook_order_items", "textbook_id"): "excellent_textbooks",
    ("academy_textbook_reviews", "textbook_id"): "excellent_textbooks",
    ("password_reset_codes", "email"): "users",
}

FIELD_ZH = {
    "id": "编号",
    "user_id": "用户编号",
    "username": "用户名",
    "password_hash": "密码散列",
    "nickname": "昵称",
    "role": "角色",
    "avatar_url": "头像地址",
    "email": "邮箱",
    "role_type": "角色类型",
    "learning_goal": "学习目标",
    "interests_json": "兴趣列表",
    "school": "学校",
    "teacher_name": "教师姓名",
    "pet_key": "宠物标识",
    "agreement_accepted": "是否同意协议",
    "onboarding_completed": "是否完成入驻",
    "enabled": "启用状态",
    "created_at": "创建时间",
    "updated_at": "更新时间",
    "external_course_id": "外部课程编号",
    "course_name": "课程名称",
    "category": "分类",
    "school_name": "学校名称",
    "cover_url": "封面地址",
    "cover_file_path": "封面本地路径",
    "start_time": "开课时间",
    "participant_count": "参与人数",
    "course_comment": "课程评论",
    "course_description": "课程描述",
    "source_url": "来源地址",
    "source_synced_at": "来源同步时间",
    "certified": "认证状态",
    "certification_label": "认证标签",
    "resource_type": "资源类型",
    "course_id": "课程业务编号",
    "publisher_user_id": "发布教师编号",
    "semester_plan": "学期计划",
    "overview": "课程概述",
    "detail": "课程详情",
    "video_file_path": "视频文件路径",
    "rating": "评分",
    "content": "内容",
    "reply_content": "回复内容",
    "reply_user_id": "回复用户编号",
    "reply_user_name": "回复用户名",
    "reply_user_role_type": "回复用户角色",
    "replied_at": "回复时间",
    "parent_review_id": "父级评价编号",
    "teacher_read_at": "教师阅读时间",
    "category_code": "分类编码",
    "category_name": "分类名称",
    "description": "说明",
    "sort_order": "排序号",
    "set_code": "套题编码",
    "title": "标题",
    "subtitle": "副标题",
    "cover_file_path": "封面本地路径",
    "difficulty_label": "难度标签",
    "status_label": "状态标签",
    "source_name": "来源名称",
    "question_count": "题目数量",
    "question_type": "题目类型",
    "question_text": "题干",
    "option_items": "选项列表",
    "answer": "答案",
    "analysis": "解析",
    "question_order": "题目序号",
    "selected_answer": "所选答案",
    "correct_answer": "正确答案",
    "wrong_count": "错误次数",
    "consecutive_correct_count": "连续正确次数",
    "mastered": "是否掌握",
    "first_wrong_at": "首次错误时间",
    "last_wrong_at": "最近错误时间",
    "last_reviewed_at": "最近复习时间",
    "subject_code": "学科编码",
    "subject_name": "学科名称",
    "source": "来源",
    "external_problem_id": "外部题目编号",
    "difficulty": "难度",
    "tag_ids": "标签编号列表",
    "tag_names": "标签名称列表",
    "problem_statement": "题目描述",
    "input_description": "输入描述",
    "output_description": "输出描述",
    "hints": "提示",
    "submit_count": "提交次数",
    "accepted_count": "通过次数",
    "imported_at": "导入时间",
    "assignment_code": "作业编码",
    "course_resource_type": "课程资源类型",
    "assignment_title": "作业标题",
    "assignment_status": "作业状态",
    "deadline_at": "截止时间",
    "attempts_limit": "尝试次数限制",
    "duration_minutes": "限时分钟数",
    "total_score": "总分",
    "assignment_description": "作业说明",
    "assignment_id": "作业编号",
    "question_label": "题目标签",
    "question_title": "题目标题",
    "question_options": "题目选项",
    "placeholder_text": "输入提示",
    "score": "分值",
    "answer_explanation": "答案解释",
    "auto_gradable": "是否自动评分",
    "oj_problem_id": "OJ题目编号",
    "requires_teacher_review": "是否需教师批阅",
    "submission_status": "提交状态",
    "answer_payload": "作答数据",
    "teacher_feedback": "教师反馈",
    "submitted_at": "提交时间",
    "exam_code": "考试编码",
    "course_title": "课程标题",
    "exam_title": "考试标题",
    "exam_status": "考试状态",
    "starts_at": "开始时间",
    "exam_description": "考试说明",
    "exam_id": "考试编号",
    "started_at": "开始作答时间",
    "slug": "访问标识",
    "time_limit_ms": "时间限制毫秒",
    "memory_limit_mb": "内存限制MB",
    "tags": "标签",
    "status": "状态",
    "created_by": "创建者编号",
    "problem_id": "题目编号",
    "input_data": "输入数据",
    "expected_output": "期望输出",
    "is_sample": "是否样例",
    "weight": "权重",
    "language": "编程语言",
    "source_code": "源代码",
    "judge_status": "评测状态",
    "time_used_ms": "耗时毫秒",
    "memory_used_kb": "内存KB",
    "judge_message": "评测消息",
    "judged_at": "评测时间",
    "submission_id": "提交编号",
    "test_case_id": "测试用例编号",
    "external_textbook_id": "外部教材编号",
    "textbook_name": "教材名称",
    "chief_editor": "主编",
    "publisher": "出版社",
    "publish_date": "出版日期",
    "isbn": "ISBN",
    "recommendation": "推荐语",
    "original_price": "原价",
    "discount_price": "折扣价",
    "reader_count": "阅读人数",
    "summary": "内容简介",
    "catalog_text": "目录文本",
    "review_text": "评论文本",
    "collected_at": "采集时间",
    "textbook_id": "教材编号",
    "quantity": "数量",
    "order_no": "订单号",
    "total_amount": "订单总额",
    "order_status": "订单状态",
    "original_amount": "原始金额",
    "discount_amount": "优惠金额",
    "voucher_discount_amount": "卡券优惠金额",
    "coin_discount_amount": "金币抵扣金额",
    "voucher_payload": "卡券快照",
    "unit_price": "单价",
    "display_name": "显示名称",
    "student_no": "用户标识",
    "role_label": "角色标签",
    "bio": "个人简介",
    "location": "所在地",
    "avatar_path": "头像路径",
    "admin_coin_adjustment": "后台金币调整值",
    "admin_data_note": "后台数据备注",
    "event_type": "事件类型",
    "set_id": "套题编号",
    "vocabulary_mastery": "词汇掌握状态",
    "target_type": "目标类型",
    "target_id": "目标编号",
    "target_title": "目标标题",
    "duration_seconds": "学习时长秒数",
    "module_type": "模块类型",
    "code_hash": "验证码散列",
    "expires_at": "过期时间",
    "used": "是否已使用",
    "attempts": "尝试次数",
    "source_type": "来源类型",
    "source_key": "来源标识",
    "reason": "原因",
    "amount": "数量",
    "related_record_id": "关联记录编号",
    "balance_after": "变动后余额",
    "voucher_key": "卡券标识",
    "voucher_type": "卡券类型",
    "name": "名称",
    "price": "价格",
    "stock_quantity": "库存数量",
    "unlimited_stock": "是否不限库存",
    "discount_type": "优惠类型",
    "threshold_amount": "门槛金额",
    "discount_rate": "折扣率",
    "max_discount_amount": "最大优惠金额",
    "valid_from": "生效时间",
    "valid_until": "失效时间",
    "bank_code": "题库编码",
    "coins_awarded": "获得金币数",
    "correct_count": "正确数",
    "wrong_count": "错误数",
    "duration_sec": "持续秒数",
    "wave_reached": "到达波次",
    "waves_completed": "完成波次数",
    "max_combo": "最大连击",
    "words_solved": "解决单词数",
    "kills": "击杀数",
    "letters_typed": "输入字母数",
    "active_seconds": "有效输入秒数",
    "template_name": "模板名称",
    "depth_points": "深度数组",
    "gamma_baseline": "自然伽马基线",
    "notes": "备注",
    "porosity": "孔隙度",
    "oil_saturation": "含油饱和度",
    "interpretation": "解释报告",
    "stroke": "冲程",
    "frequency": "冲次",
    "pump_diameter": "泵径",
    "condition_type": "工况类型",
    "dynamometer_card": "示功图数据",
    "pressure": "地层压力",
    "permeability": "渗透率",
    "water_saturation": "含水饱和度",
    "viscosity": "黏度",
    "oil_rate": "日产油量",
    "water_rate": "日产水量",
    "injection_rate": "注入速度",
    "response_days": "见效天数",
    "water_breakthrough_days": "见水天数",
    "peak_oil_rate": "峰值产油量",
    "production_curve": "产量曲线",
    "operation_type": "施工类型",
    "sand_volume": "加砂量",
    "injection_volume": "排量",
    "acid_volume": "酸液量",
    "fracture_length": "裂缝长度",
    "production_multiplier": "增产倍数",
}

FIELD_ZH.update({
    "attempt_count": "尝试次数",
    "block_code": "板块编码",
    "block_name": "板块名称",
    "category_id": "分类编号",
    "comments_text": "评论文本",
    "completed_wave_count": "完成波次数",
    "correct_streak": "连续正确次数",
    "course_detail": "课程详情",
    "course_overview": "课程概述",
    "crawled_at": "采集时间",
    "create_time": "创建时间",
    "daily_oil": "日产油量",
    "daily_water": "日产水量",
    "depth_array": "深度数组",
    "displacement": "排量",
    "effect_day": "见效天数",
    "effective_typing_seconds": "有效输入秒数",
    "explanation": "解析",
    "external_tag_id": "外部标签编号",
    "formation_pressure": "地层压力",
    "gr_base": "自然伽马基线",
    "handle": "用户标识",
    "hint": "提示",
    "indicator_chart_data": "示功图数据",
    "is_correct": "是否正确",
    "memory_limit_kb": "内存限制KB",
    "message": "消息",
    "options_json": "选项数据",
    "order_id": "订单编号",
    "parent_external_tag_id": "父级外部标签编号",
    "peak_oil": "峰值产油量",
    "question_bank_code": "题库编码",
    "question_id": "题目编号",
    "reached_wave": "到达波次",
    "reference_id": "关联记录编号",
    "remark": "备注",
    "report_json": "解释报告",
    "route_path": "路由路径",
    "sample": "是否样例",
    "samples": "样例",
    "solved_word_count": "解决单词数",
    "source_refs": "来源引用",
    "standard_code": "标准代码",
    "stem": "题干",
    "stimulation_ratio": "增产倍数",
    "storage_folder": "存储目录",
    "stroke_times": "冲次",
    "subject_id": "学科编号",
    "tag_name": "标签名称",
    "tag_type": "标签类型",
    "target_code": "目标编码",
    "total_accepted": "通过总数",
    "total_coins": "获得金币总数",
    "total_kill_count": "击杀总数",
    "total_submit": "提交总数",
    "type": "施工类型",
    "typed_letter_count": "输入字母数",
    "update_time": "更新时间",
    "user_name": "用户名称",
    "vocabulary_status": "词汇状态",
    "voucher_consumed": "卡券消耗数",
    "voucher_name": "卡券名称",
    "water_breakthrough_day": "见水天数",
    "work_condition": "工况类型",
})


def xml_text(elem):
    return "".join(node.text or "" for node in elem.xpath(".//*[local-name()='t']")).strip()


def has_drawing(elem):
    return bool(elem.xpath(".//*[local-name()='drawing']"))


def find_section_bounds(doc):
    children = list(doc.element.body)
    start = None
    end = None
    for idx, child in enumerate(children):
        text = xml_text(child)
        if start is None and text.startswith("6") and "数据库设计" in text:
            start = idx
            continue
        if start is not None and idx > start and text.startswith("7"):
            end = idx
            break
    if start is None or end is None:
        raise RuntimeError("未找到第六部分或第七部分边界")
    return start, end


def set_cell_text(cell, text, bold=False, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text or ""))
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for m in ("top", "left", "bottom", "right"):
                node = tc_mar.find(qn(f"w:{m}"))
                if node is None:
                    node = OxmlElement(f"w:{m}")
                    tc_mar.append(node)
                node.set(qn("w:w"), "60")
                node.set(qn("w:type"), "dxa")


def set_paragraph_font(p, size=10.5, bold=False):
    for run in p.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(size)
        run.bold = bold


def add_paragraph_at(doc, body, pos_ref, text="", size=10.5, bold=False, align=None):
    p = doc.add_paragraph(text)
    if align is not None:
        p.alignment = align
    set_paragraph_font(p, size=size, bold=bold)
    elem = p._p
    body.remove(elem)
    body.insert(pos_ref[0], elem)
    pos_ref[0] += 1
    return p


def add_existing_at(body, pos_ref, elem):
    body.insert(pos_ref[0], deepcopy(elem))
    pos_ref[0] += 1


def add_table_at(doc, body, pos_ref, rows, cols):
    table = doc.add_table(rows=rows, cols=cols)
    elem = table._tbl
    body.remove(elem)
    body.insert(pos_ref[0], elem)
    pos_ref[0] += 1
    style_table(table)
    return table


def field_label(name):
    if name in FIELD_ZH:
        return FIELD_ZH[name]
    return name.replace("_", " ")


def table_fields_for_model(table, cols):
    labels = [field_label(c["name"]) for c in cols.values()]
    return f"{TABLE_ZH.get(table, table)}（{ '，'.join(labels) }）"


def field_note(table, col):
    name = col["name"]
    label = field_label(name)
    if name in {"duration_seconds", "duration_sec", "active_seconds"}:
        return f"{label}，单位为秒"
    if name in {"duration_minutes"}:
        return f"{label}，单位为分钟"
    if name in {"time_limit_ms", "time_used_ms"}:
        return f"{label}，单位为毫秒"
    if name in {"memory_limit_mb"}:
        return f"{label}，单位为MB"
    if name in {"memory_used_kb"}:
        return f"{label}，单位为KB"
    if name in {"rating"}:
        return "评分，取值通常为1至5"
    if name in {"enabled", "certified", "agreement_accepted", "onboarding_completed", "used", "mastered", "is_sample", "auto_gradable", "requires_teacher_review", "unlimited_stock"}:
        return f"{label}，0表示否，1表示是"
    if name in {"role", "role_type"}:
        return f"{label}，取值区分学生、教师和管理员"
    if name in {"resource_type", "course_resource_type"}:
        return "课程资源类型，区分在线开放课程、通识课程、微专业课程等"
    if name in {"amount", "price", "total_amount", "original_amount", "discount_amount", "voucher_discount_amount", "coin_discount_amount", "unit_price", "original_price", "discount_price", "threshold_amount", "max_discount_amount"}:
        return f"{label}，金额或金币数值"
    if name.endswith("_at") or name in {"created_at", "updated_at"}:
        return f"{label}，记录业务发生或维护时间"
    return label


def key_text(table, col):
    flags = []
    if col["pk"]:
        flags.append("主键")
    ref = col.get("ref") or LOGICAL_REFS.get((table, col["name"]), "")
    if col.get("fk") or ref:
        flags.append("外键")
    return "/".join(flags), ref


def add_schema_table(doc, body, pos_ref, table_no, table_name, cols):
    zh = TABLE_ZH.get(table_name, table_name)
    add_paragraph_at(doc, body, pos_ref, f"表6.{table_no}  {zh}", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    table = add_table_at(doc, body, pos_ref, len(cols) + 3, 6)
    row = table.rows[0]
    row.cells[0].merge(row.cells[2])
    row.cells[3].merge(row.cells[5])
    set_cell_text(row.cells[0], "表名（中文）", bold=True, size=9)
    set_cell_text(row.cells[3], "表名（英文）", bold=True, size=9)
    for idx in (0, 3):
        shade_cell(row.cells[idx])

    row = table.rows[1]
    row.cells[0].merge(row.cells[2])
    row.cells[3].merge(row.cells[5])
    set_cell_text(row.cells[0], zh, size=9)
    set_cell_text(row.cells[3], table_name, size=9)

    headers = ["字段名", "类型", "长度", "主键/外键", "参照表", "说明"]
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(2, idx), header, bold=True, size=8.5)
        shade_cell(table.cell(2, idx), "EAF2F8")

    for r_idx, col in enumerate(cols.values(), start=3):
        key, ref = key_text(table_name, col)
        values = [col["name"], col["type"], col["length"], key, ref, field_note(table_name, col)]
        for c_idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.LEFT if c_idx in {0, 4, 5} else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(table.cell(r_idx, c_idx), value, size=8, align=align)
    add_paragraph_at(doc, body, pos_ref, "", size=10.5)


def collect_er_images(children, start, end):
    images = []
    for child in children[start:end]:
        if has_drawing(child):
            clone = deepcopy(child)
            for text_node in clone.xpath(".//*[local-name()='t']"):
                text_node.text = ""
            images.append(clone)
    return images


def rewrite():
    doc = Document(SRC)
    body = doc.element.body
    children = list(body)
    start, end = find_section_bounds(doc)
    er_images = collect_er_images(children, start, end)

    for child in children[start:end]:
        body.remove(child)

    pos_ref = [start]
    schema = parse_schema()
    schema.pop("auth_users", None)

    add_paragraph_at(doc, body, pos_ref, "6  数据库设计", size=14, bold=True)
    add_paragraph_at(doc, body, pos_ref, "6.1  概念结构设计", size=12, bold=True)
    add_paragraph_at(
        doc,
        body,
        pos_ref,
        "本系统数据库的概念结构以学习平台真实业务对象为基础进行抽象，用户视图围绕“用户、课程、题库、作业考试、OJ评测、教材交易、金币卡券、学习游戏、测井与采油仿真”等实体展开。概念层不直接关心数据表拆分细节，而是描述现实世界中对象的属性、记录形式和联系：用户可以登录系统、完善画像、加入课程、完成题库练习、提交作业考试、购买教材、兑换卡券、参与游戏并形成学习轨迹；教师可以发布课程、查看课程评价和处理教学反馈；管理员可以维护用户、课程分类、题库、OJ题目和兑换资源。",
    )
    add_paragraph_at(
        doc,
        body,
        pos_ref,
        "从数据项和值域看，系统将编号类数据统一设计为BIGINT主键或业务编码；名称、标题、学校、教师姓名等文本类数据使用VARCHAR或TEXT；题目选项、用户兴趣、作答内容、卡券快照和仿真曲线等半结构化数据使用JSON；创建时间、提交时间、截止时间、评测时间等时间类数据使用DATETIME或TIMESTAMP；评分、数量、金币、价格和实验参数等数值类数据使用INT、BIGINT、DECIMAL、DOUBLE等类型，并通过服务层和部分CHECK约束限制取值范围，例如评分通常为1至5，数量和金额不应为负，布尔状态字段以0/1表示。",
    )
    add_paragraph_at(
        doc,
        body,
        pos_ref,
        "概念结构设计采用“统一用户中心+分域业务实体+行为记录沉淀”的思路。用户与用户画像是一对一关系；用户与学习事件、学习时长、课程加入、评价、题库错题、收藏、OJ提交、教材订单、游戏成绩、金币流水和仿真实验记录是一对多关系；课程题库分类与题库套题、题库套题与题目、作业/考试与题目、OJ题目与测试用例、教材订单与订单明细均为一对多关系；通用题源题目与学科之间通过中间表形成多对多关系。在线开放课程、通识课程和微专业课程保留独立实体，通过resource_type/course_id形成逻辑关联，便于不同来源课程扩展字段而不破坏统一交互流程。",
    )
    for title, tables in GROUPS:
        actual = [TABLE_ZH.get(t, t) for t in tables if t in schema]
        add_paragraph_at(doc, body, pos_ref, f"（{title}）主要实体包括：{'、'.join(actual)}。", size=10.5)

    add_paragraph_at(doc, body, pos_ref, "根据上述概念结构，第六部分保留原报告中已经绘制完成的陈氏ER图。ER图从系统总览、用户画像、在线学堂、题库作业考试、OJ评测、实验仿真与学习游戏等角度展示实体、属性与联系。", size=10.5)
    captions = [
        "图6-1  数据库总览陈氏ER图",
        "图6-2  用户认证与学习画像陈氏ER图",
        "图6-3  在线学堂与精品教材陈氏ER图",
        "图6-4  题库、作业与考试陈氏ER图",
        "图6-5  OJ在线编程陈氏ER图",
        "图6-6  实验仿真与学习游戏陈氏ER图",
    ]
    for idx, image in enumerate(er_images[:6]):
        add_existing_at(body, pos_ref, image)
        add_paragraph_at(doc, body, pos_ref, captions[idx], size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_paragraph_at(doc, body, pos_ref, "6.2  逻辑结构设计", size=12, bold=True)
    add_paragraph_at(
        doc,
        body,
        pos_ref,
        "数据库管理员视图在概念结构基础上将业务实体重新组织为全局逻辑结构。系统使用users作为统一用户主表，其他个体化业务表通过user_id与用户建立关系；课程资源保持三张来源表，课程加入、评价、作业、考试通过resource_type与course_id关联到具体课程；题库、作业考试、OJ、教材交易、金币卡券和仿真实验均按业务域拆表，既保证单表职责清晰，也便于后端Controller、Service和Repository按模块查询维护。",
    )
    add_paragraph_at(doc, body, pos_ref, "本系统的数据模型表示如下：", size=10.5)
    for title, tables in GROUPS:
        add_paragraph_at(doc, body, pos_ref, f"{title}：", size=10.5, bold=True)
        for table in tables:
            if table in schema:
                add_paragraph_at(doc, body, pos_ref, table_fields_for_model(table, schema[table]), size=10)
    add_paragraph_at(
        doc,
        body,
        pos_ref,
        "主要表间关系为：users.id关联画像、学习事件、学习时长、错题收藏、OJ提交、教材订单、金币流水、卡券、游戏成绩和仿真实验记录；course_question_bank_categories.id关联course_question_bank_sets.category_id，course_question_bank_sets.id关联course_question_bank_questions.set_id；academy_assignments.id关联作业题目和作业提交，academy_exams.id关联考试题目和考试提交；oj_problems.id关联测试用例和代码提交，oj_submissions.id关联单用例评测结果；academy_textbook_orders.id关联教材订单明细；question_bank_problem_subjects以复合主键实现通用题源题目与学科的多对多关系。",
    )
    add_paragraph_at(doc, body, pos_ref, "各核心业务表结构如下。", size=10.5)

    table_no = 1
    for _, tables in GROUPS:
        for table_name in tables:
            if table_name in schema:
                add_schema_table(doc, body, pos_ref, table_no, table_name, schema[table_name])
                table_no += 1

    doc.save(OUT)
    print(f"OUT={OUT}")
    print(f"schema_tables={len(schema)}")
    print(f"word_tables_inserted={table_no - 1}")
    print(f"er_images_preserved={min(len(er_images), 6)}")


if __name__ == "__main__":
    rewrite()
