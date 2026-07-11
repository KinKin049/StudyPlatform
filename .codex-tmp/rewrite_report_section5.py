from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "CET46" / "报告-latest.docx"
OUT = ROOT / "CET46" / "报告-latest-第五部分结构修订.docx"
FLOW_DIR = ROOT / ".codex-tmp" / "report_flowcharts"
FLOW_DIR.mkdir(parents=True, exist_ok=True)


def xml_text(elem):
    return "".join(node.text or "" for node in elem.xpath(".//w:t")).strip()


def has_drawing(elem):
    return bool(elem.xpath(".//w:drawing"))


def next_text(children, index, span=3):
    parts = []
    for child in children[index + 1 : index + 1 + span]:
        parts.append(xml_text(child))
    return "".join(parts)


def find_section_bounds(doc):
    body = doc.element.body
    children = list(body)
    start = None
    end = None
    for idx, child in enumerate(children):
        text = xml_text(child)
        if start is None and text.startswith("5") and "系统详细设计" in text:
            start = idx
            continue
        if start is not None and idx > start and text.startswith("6") and "数据库设计" in text:
            end = idx
            break
    if start is None or end is None:
        raise RuntimeError("未找到第五章或第六章边界")
    return start, end


def collect_interface_images(doc, start, end):
    body = doc.element.body
    children = list(body)
    groups = {f"5.{idx}": [] for idx in range(1, 7)}
    current = None
    for idx in range(start, end):
        child = children[idx]
        text = xml_text(child)
        for section in groups:
            if text.startswith(section):
                current = section
                break
        if current and has_drawing(child):
            caption_nearby = next_text(children, idx)
            if "流程图" in caption_nearby:
                continue
            groups[current].append(deepcopy(child))
    return groups


def load_font(size, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def wrap_text(draw, text, font, max_width):
    lines = []
    for raw in text.split("\n"):
        line = ""
        for char in raw:
            trial = line + char
            bbox = draw.textbbox((0, 0), trial, font=font)
            if bbox[2] - bbox[0] <= max_width or not line:
                line = trial
            else:
                lines.append(line)
                line = char
        if line:
            lines.append(line)
    return lines


def draw_centered_text(draw, box, text, font, fill):
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, x2 - x1 - 34)
    line_h = font.size + 7
    total_h = line_h * len(lines)
    y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        w = bbox[2] - bbox[0]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=font, fill=fill)
        y += line_h


def draw_arrow(draw, start, end, fill):
    draw.line([start, end], fill=fill, width=4)
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex >= sx else -1
        pts = [(ex, ey), (ex - 14 * direction, ey - 8), (ex - 14 * direction, ey + 8)]
    else:
        direction = 1 if ey >= sy else -1
        pts = [(ex, ey), (ex - 8, ey - 14 * direction), (ex + 8, ey - 14 * direction)]
    draw.polygon(pts, fill=fill)


def make_flowchart(key, title, nodes):
    width = 1500
    height = 780
    img = Image.new("RGB", (width, height), "#f7fbff")
    draw = ImageDraw.Draw(img)
    title_font = load_font(42, bold=True)
    node_font = load_font(30, bold=True)
    small_font = load_font(22)
    draw.rounded_rectangle((20, 20, width - 20, height - 20), radius=28, fill="#ffffff", outline="#b6d7ef", width=3)
    draw.text((60, 48), title, font=title_font, fill="#17324d")
    draw.text((60, 108), "系统处理逻辑流程", font=small_font, fill="#5f7182")
    box_w = 360
    box_h = 118
    gap_x = 145
    row1_y = 180
    row2_y = 455
    boxes = []
    for i, node in enumerate(nodes):
        row = 0 if i < 3 else 1
        col = i if i < 3 else 5 - i
        x = 70 + col * (box_w + gap_x)
        y = row1_y if row == 0 else row2_y
        boxes.append((x, y, x + box_w, y + box_h, node))
    fills = ["#e8f6ff", "#eef9f0", "#fff7e6", "#f1edff", "#ffeef4", "#edf7f6"]
    outlines = ["#5ba9d6", "#67b877", "#e0a340", "#9b7ad9", "#d9719f", "#49a99e"]
    for i, (x1, y1, x2, y2, node) in enumerate(boxes):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=20, fill=fills[i], outline=outlines[i], width=3)
        draw_centered_text(draw, (x1, y1, x2, y2), node, node_font, "#15324a")
    for i in range(len(boxes) - 1):
        x1, y1, x2, y2, _ = boxes[i]
        nx1, ny1, nx2, ny2, _ = boxes[i + 1]
        if i == 2:
            draw_arrow(draw, ((x1 + x2) / 2, y2 + 4), ((nx1 + nx2) / 2, ny1 - 4), "#52708a")
        else:
            draw_arrow(draw, (x2 + 8, (y1 + y2) / 2), (nx1 - 8, (ny1 + ny2) / 2), "#52708a")
    path = FLOW_DIR / f"{key}.png"
    img.save(path)
    return path


flowcharts = {
    "5.1": make_flowchart(
        "section5_1_academy",
        "图5-1 在线学堂模块流程图",
        [
            "用户进入在线学堂",
            "按资源类型、分类、关键词检索",
            "前端调用 Academy API",
            "后端查询课程、教材、作业考试数据",
            "加入课程、学习视频或购买教材",
            "写入报名、订单、学习时长与金币记录",
        ],
    ),
    "5.2": make_flowchart(
        "section5_2_question_bank",
        "图5-2 题库练习与金币激励模块流程图",
        [
            "用户进入题库、错题或收藏",
            "加载目录、题目、掌握状态",
            "提交答案或切换收藏",
            "后端更新错题、收藏、学习事件",
            "金币服务按规则结算奖励",
            "个人主页汇总正确率、时长和金币",
        ],
    ),
    "5.3": make_flowchart(
        "section5_3_profile_admin",
        "图5-3 用户画像、教师端与后台管理模块流程图",
        [
            "用户登录并识别角色",
            "学生维护资料、查看学习画像",
            "教师发布课程、查看工作台",
            "管理员维护用户、课程、题库、OJ和卡券",
            "后端校验角色权限并写库",
            "前端刷新画像、信箱和管理结果",
        ],
    ),
    "5.4": make_flowchart(
        "section5_4_ai_pet",
        "图5-4 AI宠物与学习辅助模块流程图",
        [
            "全局宠物组件读取当前页面上下文",
            "用户发起聊天、导航、待办或番茄钟操作",
            "聊天请求进入 /api/ai-pet/chat",
            "服务端组装提示词并调用模型",
            "前端展示回复或执行本地交互",
            "通过陪伴、提醒和入口跳转辅助学习",
        ],
    ),
    "5.5": make_flowchart(
        "section5_5_visual_lab_oj",
        "图5-5 可视化、实验与OJ模块流程图",
        [
            "用户进入可视化中心或实验平台",
            "选择算法、函数、空间模型、石油仿真或OJ",
            "前端渲染 ECharts、Three.js 或实验面板",
            "OJ提交进入后端判题服务",
            "判题服务调用 judge-sandbox 或保存结果",
            "学习时长与评测结果回流个人画像",
        ],
    ),
    "5.6": make_flowchart(
        "section5_6_games_exchange",
        "图5-6 学习游戏与兑换中心模块流程图",
        [
            "用户进入游戏学习平台",
            "完成万题天梯跳或 TypeWarrior",
            "前端提交游戏成绩记录",
            "后端校验并触发金币奖励",
            "兑换中心读取金币余额和商品",
            "校验库存后扣减金币并发放卡券",
        ],
    ),
}

flowchart_titles = {
    "5.1": "图5-1  在线学堂模块流程图",
    "5.2": "图5-2  题库练习与金币激励模块流程图",
    "5.3": "图5-3  用户画像、教师端与后台管理模块流程图",
    "5.4": "图5-4  AI宠物与学习辅助模块流程图",
    "5.5": "图5-5  可视化、实验与OJ模块流程图",
    "5.6": "图5-6  学习游戏与兑换中心模块流程图",
}


sections = [
    {
        "key": "5.1",
        "title": "5.1  在线学堂模块",
        "description": [
            "基本描述：在线学堂模块是平台课程学习与教材服务的综合入口，前端由 AcademyPage.vue、AcademyHome.vue、AcademyOpenCourses.vue、AcademyGeneralCourses.vue、AcademyMicroMajors.vue、AcademyCourseDetail.vue、AcademyTextbooks.vue、AcademyTextbookDetail.vue、AcademyTextbookCart.vue、AcademyAggregatePage.vue、AcademyAssignmentDetail.vue、AcademyExamIntro.vue 和 AcademyExamDetail.vue 等页面组成，后端由 AcademyController、AcademyService、AcademyRepository 及作业、考试相关 Repository 共同支撑。",
            "主要需求：解决学习资源来源分散、课程详情展示不统一、作业考试入口分散、教材交易缺少闭环的问题，使学生能够从一个在线学堂入口完成课程浏览、课程加入、视频学习、作业考试、教材查看、购物车结算和课程评价。",
            "建设目的意义：通过“课程资源—学习行为—个人画像—金币激励”的链路，把静态课程展示转化为可记录、可反馈、可激励的学习过程，为后续教师课程发布、学习推荐和平台运营提供数据基础。",
            "建设主要内容：系统按 resource 参数区分 online-open-courses、general-courses、micro-major-courses 和 textbooks，列表页复用 useAcademyList.js 完成分类、关键词和分页处理；课程详情页统一展示封面、教师、学校、课程说明、评价和视频资源；教材模块支持详情、目录、购物车、订单、支付、优惠券和评论；聚合页按 courses、assignments、exams 三种 variant 展示我的课程、作业和考试。",
            "模块特点：该模块采用前后端分离设计，资源类型与资源编号解耦；教师上传的封面和视频保存到 storage/teacher_courses 目录并由统一资源地址解析；视频学习时长通过 useVideoLearningTimeTracker 上报；作业考试页面通过路由 meta.hidePet 隐藏宠物浮窗，保证严肃答题场景的专注性。",
        ],
        "flow": "在线学堂模块的流程从用户选择资源开始。前端根据课程类型、分类和关键词请求 Academy API，后端分别查询在线开放课程、通识课程、微专业课程、教材、报名关系、作业考试和订单等数据，并返回统一结构给前端展示。用户继续执行加入课程、观看视频、提交作业、参加考试、加入购物车或支付教材订单时，系统将报名关系、学习时长、作业考试提交、订单状态和金币奖励写入数据库，形成课程学习闭环。其实现流程如图5-1所示。",
        "interface": "本模块界面设计主要包括在线学堂首页、课程列表、课程详情、我的课程、作业详情、考试作答、教材详情和教材购物车等页面。界面以资源卡片、详情信息区、学习操作区和结果反馈区为核心，既保证课程浏览效率，也保证作业考试和教材交易流程的完整性。",
    },
    {
        "key": "5.2",
        "title": "5.2  题库练习与金币激励模块",
        "description": [
            "基本描述：题库练习与金币激励模块负责题库目录、题目练习、错题复盘、收藏复习、学习事件记录和金币奖励结算。前端页面包括 AcademyQuestionBank.vue、AcademyQuestionBankCourses.vue、AcademyQuestionBankCourseDetail.vue、AcademyQuestionBankMistakes.vue 和 AcademyQuestionBankFavorites.vue，后端主要由 QuestionBankController、QuestionBankService、QuestionBankRepository、ProfileService 和 CoinRewardService 实现。",
            "主要需求：满足学生按科目、课程题库、错题状态和收藏内容进行持续练习的需求，同时解决练习行为缺少统计和激励的问题。",
            "建设目的意义：通过把答题结果、词汇掌握、错题状态和学习时长统一写入用户画像，并由后端统一结算金币，增强学生持续练习的动力，避免前端直接决定奖励金额造成数据不一致。",
            "建设主要内容：模块提供题库科目、题库课程目录、题目分页、题目详情、错题摘要、错题筛选、答题结果记录、收藏摘要、收藏列表、收藏切换、TypeWarrior 词库和洛谷题库导入等接口；题库数据来源包括课程题库、CET4/CET6 词汇、思政类题库、计算机等级题库和外部题源导入。",
            "模块特点：金币规则集中在 CoinRewardService 中实现，学习时长按10分钟为一个结算单元，题库、错题、收藏、OJ、作业和考试等学习模块可共享结算逻辑；coin_reward_records 使用 source_type 和 source_key 区分奖励来源，支持幂等去重与余额汇总。",
        ],
        "flow": "题库练习流程以“加载题目—提交行为—更新状态—结算奖励”为主线。用户进入题库后，前端加载题库目录、题目分页、错题或收藏状态；用户作答后，前端一方面调用题库答题接口更新错题和掌握状态，另一方面调用个人画像接口记录学习事件和学习时长；后端 ProfileService 写入事件后触发 CoinRewardService，按答题正确性、词汇掌握状态和学习时长统一生成金币记录。其实现流程如图5-2所示。",
        "interface": "本模块界面设计主要包括题库首页、课程题库目录页、题目练习页、错题本页和收藏题目页。界面需要突出题目内容、选项反馈、解析展示、错题状态、收藏状态、分页检索和练习统计，使用户能够在普通练习和复盘练习之间顺畅切换。",
    },
    {
        "key": "5.3",
        "title": "5.3  用户画像、教师端与后台管理模块",
        "description": [
            "基本描述：用户画像、教师端与后台管理模块承担账号角色识别、个人学习数据聚合、教师课程维护和管理员运营配置功能。前端主要页面包括 ProfilePage.vue、TeacherMailboxPage.vue、AdminPage.vue、AuthLoginPage.vue、AuthRegisterPage.vue、AuthForgotPasswordPage.vue 和 AuthOnboardingPage.vue，后端由 Auth、Profile、Academy、Admin、Rewards 等包共同支撑。",
            "主要需求：满足学生查看学习成果和维护资料、教师发布课程和查看课程反馈、管理员管理平台用户与内容的需求，解决不同角色使用同一平台时权限边界不清的问题。",
            "建设目的意义：通过统一账号体系和 roleType 字段区分学生、教师、管理员，使学习端、教师端和管理端在同一数据库和后端服务下协作，同时为平台运营、课程审核和学习数据分析提供可靠入口。",
            "建设主要内容：个人主页展示学习时长、今日学习、练习记录、连续学习、答题正确率、词汇掌握、错题统计、游戏成绩、金币余额、近期动态、成就徽章和教材订单；教师账号可在个人主页发布在线开放课程、上传课程封面和视频，并通过教师工作台和信箱查看课程反馈；管理员可维护用户、课程分类、课程、评价回复、题库集合、题目、OJ题目、OJ分类、卡券和兑换商品。",
            "模块特点：教师发布课程前由后端 ensureTeacher 校验教师身份，课程文件保存到用户目录；管理员操作前由 AdminService 校验管理员账号；个人画像数据由 ProfileService 从学习事件、学习时长、游戏记录、题库错题和金币奖励中聚合，避免各页面重复统计。",
        ],
        "flow": "本模块以角色识别为入口。用户登录后，系统根据账号信息和 roleType 展示相应功能：学生侧读取个人画像和学习统计，教师侧进入课程发布、已发布课程和信箱反馈，管理员侧进入用户、课程、题库、OJ和卡券管理。所有教师和管理员写操作先经过后端权限校验，再由对应 Repository 写入数据库，前端根据返回结果刷新界面。其实现流程如图5-3所示。",
        "interface": "本模块界面设计主要包括登录注册页面、账号引导页面、个人主页、教师课程发布区域、教师信箱和管理员后台。界面应体现角色差异：学生侧突出学习成果，教师侧突出课程维护和反馈处理，管理员侧突出表格管理、筛选、编辑和删除确认。",
    },
    {
        "key": "5.4",
        "title": "5.4  AI宠物与学习辅助模块",
        "description": [
            "基本描述：AI宠物与学习辅助模块由全局 AiPetWidget.vue 组件和后端 AiPetChatService 组成，不是独立页面路由，而是覆盖在主要学习页面上的陪伴式辅助层。",
            "主要需求：满足用户在学习过程中进行即时问答、页面理解、入口导航、待办记录和番茄专注的需求，降低用户在多个功能模块之间切换和查找入口的成本。",
            "建设目的意义：AI宠物把学习问答、页面上下文解释、学习提醒和轻量任务管理整合到同一交互对象中，使平台从单纯工具集合变成具有陪伴和反馈能力的学习环境。",
            "建设主要内容：前端宠物组件支持拖拽、展开聊天、读取页面可见文本、识别路由、处理选中文本、保存待办和番茄钟状态；chatWithAiPet 调用 /api/ai-pet/chat，后端限制历史消息、用户消息和页面上下文长度，组装系统提示词后调用配置的 AI 中转服务。",
            "模块特点：宠物组件会根据 route meta.hidePet 在作业和考试等严肃答题页面隐藏；本地待办和番茄钟使用 localStorage 保存，聊天服务使用 pageContext 增强回答的页面相关性，并在提示词中约束模型不得虚构已完成的操作。",
        ],
        "flow": "AI宠物流程以页面上下文采集和用户意图处理为核心。前端组件读取当前路由、标题、可见文本、选中文本和表单摘要；用户提出问题时，组件把消息、历史记录和页面上下文提交给后端；AiPetChatService 生成系统提示词和用户提示词后调用模型接口，并把回复返回前端。对于导航、待办和番茄钟等可在前端本地完成的操作，组件直接执行并更新本地状态。其实现流程如图5-4所示。",
        "interface": "本模块界面设计主要包括悬浮宠物、聊天面板、待办清单、番茄钟控制区和快捷导航反馈。界面应保持轻量，不遮挡主要学习内容，同时通过动画状态表达待机、思考、开心、专注和休息等学习陪伴状态。",
    },
    {
        "key": "5.5",
        "title": "5.5  可视化、实验与OJ模块",
        "description": [
            "基本描述：可视化、实验与OJ模块承担算法结构演示、函数图像绘制、三维空间模型、石油专业仿真和在线编程评测等实践学习功能。前端页面包括 VisualizationHome.vue、DataStructureVisualization.vue、AlgorithmDemoViewer.vue、FunctionGraph2D.vue、SpaceModelGuide.vue、SpaceModel3D.vue、LabPlatform.vue、PetroleumSimulation.vue、WellLogSimulation.vue 和 OjPlatform.vue。",
            "主要需求：满足学生通过图形化、交互式和编程评测方式理解抽象知识的需求，同时为专业实验和编程训练提供独立入口。",
            "建设目的意义：可视化和实验模块将课程学习从文本阅读扩展到交互操作；OJ模块将编程题面、代码提交、沙箱评测和结果回传串联起来，使平台具备“学习—练习—验证”的实践闭环。",
            "建设主要内容：函数图像实验室使用 ECharts 根据表达式、x轴范围和采样点实时绘制曲线；三维空间模型页面使用 Three.js 展示空间模型；石油气仿真平台组织测井曲线、油藏动态、注水开发和抽油机功图等面板；OJ页面加载题目、展示题面、接收 C++ 或答案模式提交，并通过 /api/oj/submissions 查询判题结果和测试用例详情。",
            "模块特点：可视化和石油仿真均通过 useLearningTimeTracker 上报 visualization 或 petroleum 学习时长；OJ后端将提交先保存到 oj_submissions，再通过 OjJudgeService 调用 JudgeSandboxClient，配置 oj.sandbox-url 后可连接 judge-sandbox 独立服务完成真实编译运行。",
        ],
        "flow": "本模块根据用户选择分为可视化实验、石油仿真和OJ评测三条流程。可视化页面在前端完成图形渲染并记录学习时长；石油仿真页面通过选项卡组织不同实验场景，同样上报实验时长；OJ页面在用户提交代码后生成提交记录，后端异步调用判题服务，判题服务读取题目和测试用例，调用远程沙箱或答案模式得到结果，再写回提交状态、得分、耗时、内存和单用例结果。其实现流程如图5-5所示。",
        "interface": "本模块界面设计主要包括可视化中心、二维函数图像实验室、三维空间模型页、实验平台入口、石油气仿真面板和OJ评测页面。界面需要突出交互参数、图形展示区域、代码编辑区、提交按钮、判题状态和测试用例结果。",
    },
    {
        "key": "5.6",
        "title": "5.6  学习游戏与兑换中心模块",
        "description": [
            "基本描述：学习游戏与兑换中心模块负责将题库练习和英语词汇训练包装为游戏化体验，并把金币余额转化为可兑换的宠物形象、装扮、主题、徽章和优惠卡券。前端主要由 GamePlatform.vue、LadderJumpGame.vue、TypeWarriorGamePage.vue、ExchangeCenter.vue 和 MyVouchers.vue 组成，后端由 Games 与 Rewards 包实现。",
            "主要需求：满足用户在轻量游戏场景中进行知识练习、获得即时反馈和使用学习金币兑换权益的需求，提升学习过程的趣味性与持续性。",
            "建设目的意义：游戏模块把答题、拼写和反应训练转化为可量化的成绩记录；兑换中心把金币奖励转化为可见资产，使学习成果具有持续反馈和运营价值。",
            "建设主要内容：万题天梯跳读取题库并保存总金币、答对数、答错数和游戏时长；TypeWarrior 保存得分、波次、连击、拼对单词数、击杀数、输入字母数和有效输入时长；GameRecordService 对所有数值做非负校验后写入数据库，并调用 CoinRewardService 发放游戏金币；兑换中心读取个人金币余额、可兑换商品和用户卡券，兑换时由 VoucherService 校验余额与库存，扣减金币并发放卡券。",
            "模块特点：游戏奖励与学习奖励共用 CoinRewardService 和 coin_reward_records 表，万题天梯跳按游戏内金币发放，TypeWarrior 按 score/100 折算；兑换商品使用 voucher_items 和 user_vouchers 保存库存、价格、类型和用户持有数量，可与教材订单优惠券联动。",
        ],
        "flow": "学习游戏流程从游戏入口开始，用户完成万题天梯跳或TypeWarrior后，前端提交成绩，后端校验记录并触发金币奖励。兑换流程从读取个人画像金币余额开始，用户选择商品后，VoucherService 校验商品是否可用、库存是否充足、金币是否足够，成功后扣减金币并写入用户卡券。其实现流程如图5-6所示。",
        "interface": "本模块界面设计主要包括游戏学习平台、万题天梯跳游戏界面、TypeWarrior战斗界面、结果结算弹窗、兑换中心商品列表和我的卡券页面。界面应突出游戏反馈、成绩结算、金币余额、商品价格、库存状态和兑换结果。",
    },
]


def add_paragraph_before(doc, anchor, text="", bold=False, size=None, align_center=False):
    para = doc.add_paragraph()
    if text:
        run = para.add_run(text)
        run.bold = bold
        if size:
            run.font.size = Pt(size)
    if align_center:
        para.alignment = 1
    anchor.addprevious(para._p)
    return para


def add_picture_before(doc, anchor, image_path, width=6.4):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    para.alignment = 1
    anchor.addprevious(para._p)
    return para


def insert_old_images(anchor, images):
    for image_elem in images:
        anchor.addprevious(deepcopy(image_elem))


doc = Document(SRC)
start, end = find_section_bounds(doc)
image_groups = collect_interface_images(doc, start, end)
body = doc.element.body
children = list(body)
anchor = children[end]

for child in children[start:end]:
    body.remove(child)

add_paragraph_before(doc, anchor, "5  系统详细设计", bold=True, size=15, align_center=True)

for section in sections:
    add_paragraph_before(doc, anchor, section["title"], bold=True, size=14)
    add_paragraph_before(doc, anchor, f"{section['key']}.1  模块描述", bold=True, size=12)
    for item in section["description"]:
        add_paragraph_before(doc, anchor, item)
    add_paragraph_before(doc, anchor, f"{section['key']}.2  流程描述", bold=True, size=12)
    add_paragraph_before(doc, anchor, section["flow"])
    add_picture_before(doc, anchor, flowcharts[section["key"]])
    add_paragraph_before(doc, anchor, flowchart_titles[section["key"]], align_center=True)
    add_paragraph_before(doc, anchor, f"{section['key']}.3  界面设计", bold=True, size=12)
    add_paragraph_before(doc, anchor, section["interface"])
    insert_old_images(anchor, image_groups.get(section["key"], []))
    add_paragraph_before(doc, anchor, "")

doc.save(OUT)
print(OUT)
for key, images in image_groups.items():
    print(f"{key} preserved_images={len(images)}")
