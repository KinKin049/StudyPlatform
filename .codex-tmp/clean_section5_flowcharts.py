from pathlib import Path

from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT / "CET46" / "报告-latest-第五部分流程图专业修订.docx"
OUT = ROOT / "CET46" / "报告-latest-第五部分流程图无遮挡修订.docx"
FLOW_DIR = ROOT / ".codex-tmp" / "clean_flowcharts"
FLOW_DIR.mkdir(parents=True, exist_ok=True)


def font(size, bold=False):
    paths = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for path in paths:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


TITLE_FONT = font(44, True)
LANE_FONT = font(24, True)
BOX_FONT = font(22, True)
LABEL_FONT = font(18)

W = 2200
LEFT = 55
RIGHT = 55
TITLE_Y = 36
LANE_X1 = 55
LANE_LABEL_W = 245
CONTENT_X1 = LANE_X1 + LANE_LABEL_W + 42
CONTENT_X2 = W - RIGHT - 70
GUTTER_X = W - RIGHT - 38
LANE_H = 230
LANE_GAP = 28
TOP = 135
BOX_W = 235
BOX_H = 82
DEC_W = 235
DEC_H = 112


def size(draw, text, fnt):
    box = draw.textbbox((0, 0), text, font=fnt)
    return box[2] - box[0], box[3] - box[1]


def wrap(draw, text, fnt, max_w):
    lines = []
    for raw in text.split("\n"):
        line = ""
        for ch in raw:
            trial = line + ch
            if size(draw, trial, fnt)[0] <= max_w or not line:
                line = trial
            else:
                lines.append(line)
                line = ch
        if line:
            lines.append(line)
    return lines


def centered_text(draw, box, text, fnt=BOX_FONT):
    x1, y1, x2, y2 = box
    lines = wrap(draw, text, fnt, x2 - x1 - 28)
    lh = fnt.size + 8
    total_h = lh * len(lines)
    y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        tw, _ = size(draw, line, fnt)
        draw.text((x1 + (x2 - x1 - tw) / 2, y), line, fill="#111111", font=fnt)
        y += lh


def anchor(box, side):
    x1, y1, x2, y2 = box
    return {
        "left": (x1, (y1 + y2) / 2),
        "right": (x2, (y1 + y2) / 2),
        "top": ((x1 + x2) / 2, y1),
        "bottom": ((x1 + x2) / 2, y2),
        "center": ((x1 + x2) / 2, (y1 + y2) / 2),
    }[side]


def arrow_head(draw, p1, p2):
    sx, sy = p1
    ex, ey = p2
    if abs(ex - sx) >= abs(ey - sy):
        d = 1 if ex >= sx else -1
        pts = [(ex, ey), (ex - d * 14, ey - 7), (ex - d * 14, ey + 7)]
    else:
        d = 1 if ey >= sy else -1
        pts = [(ex, ey), (ex - 7, ey - d * 14), (ex + 7, ey - d * 14)]
    draw.polygon(pts, fill="#111111")


def poly_arrow(draw, pts, label=None, label_pos=0.5):
    draw.line(pts, fill="#111111", width=3, joint="curve")
    arrow_head(draw, pts[-2], pts[-1])
    if label:
        seg_idx = max(0, min(len(pts) - 2, int((len(pts) - 1) * label_pos)))
        x1, y1 = pts[seg_idx]
        x2, y2 = pts[seg_idx + 1]
        lx = (x1 + x2) / 2
        ly = (y1 + y2) / 2
        if abs(x2 - x1) >= abs(y2 - y1):
            ly -= 32
        else:
            lx += 18
        tw, th = size(draw, label, LABEL_FONT)
        draw.rectangle((lx - tw / 2 - 7, ly - 4, lx + tw / 2 + 7, ly + th + 4), fill="#ffffff")
        draw.text((lx - tw / 2, ly), label, fill="#111111", font=LABEL_FONT)


def node(draw, center, text, kind="process", w=BOX_W, h=BOX_H):
    cx, cy = center
    if kind == "decision":
        w = max(w, DEC_W)
        h = max(h, DEC_H)
        pts = [(cx, cy - h / 2), (cx + w / 2, cy), (cx, cy + h / 2), (cx - w / 2, cy)]
        draw.polygon(pts, fill="#ffffff", outline="#111111")
        draw.line([pts[0], pts[1], pts[2], pts[3], pts[0]], fill="#111111", width=3)
        box = (cx - w / 2 + 18, cy - h / 2 + 12, cx + w / 2 - 18, cy + h / 2 - 12)
        centered_text(draw, box, text)
        return (cx - w / 2, cy - h / 2, cx + w / 2, cy + h / 2)
    box = (cx - w / 2, cy - h / 2, cx + w / 2, cy + h / 2)
    if kind == "terminator":
        draw.rounded_rectangle(box, radius=32, fill="#ffffff", outline="#111111", width=3)
    else:
        draw.rounded_rectangle(box, radius=14, fill="#ffffff", outline="#111111", width=3)
    centered_text(draw, box, text)
    return box


def lane_bounds(row):
    y1 = TOP + row * (LANE_H + LANE_GAP)
    y2 = y1 + LANE_H
    return y1, y2


def lane_center(row):
    y1, y2 = lane_bounds(row)
    return (y1 + y2) / 2


def draw_frame(draw, lanes):
    for i, lane in enumerate(lanes):
        y1, y2 = lane_bounds(i)
        draw.rounded_rectangle((LEFT, y1, W - RIGHT, y2), radius=12, fill="#ffffff", outline="#777777", width=2)
        draw.line((LEFT + LANE_LABEL_W, y1, LEFT + LANE_LABEL_W, y2), fill="#999999", width=2)
        centered_text(draw, (LEFT + 8, y1 + 8, LEFT + LANE_LABEL_W - 8, y2 - 8), lane, LANE_FONT)


def row_nodes(draw, row, specs, xs=None):
    cy = lane_center(row)
    count = len(specs)
    if xs is None:
        start = CONTENT_X1 + BOX_W / 2
        end = CONTENT_X2 - BOX_W / 2
        step = 0 if count == 1 else (end - start) / (count - 1)
        xs = [start + i * step for i in range(count)]
    boxes = []
    for x, spec in zip(xs, specs):
        if isinstance(spec, str):
            text, kind, w, h = spec, "process", BOX_W, BOX_H
        else:
            text = spec.get("text", "")
            kind = spec.get("kind", "process")
            w = spec.get("w", BOX_W)
            h = spec.get("h", BOX_H)
        boxes.append(node(draw, (x, cy), text, kind=kind, w=w, h=h))
    return boxes


def connect_row(draw, boxes, labels=None):
    labels = labels or [None] * (len(boxes) - 1)
    for left, right, label in zip(boxes, boxes[1:], labels):
        start = anchor(left, "right")
        end = anchor(right, "left")
        poly_arrow(draw, [(start[0] + 5, start[1]), (end[0] - 5, end[1])], label)


def connect_down(draw, source, target, label=None):
    sx, sy = anchor(source, "bottom")
    tx, ty = anchor(target, "top")
    mid_y = sy + 32
    poly_arrow(draw, [(sx, sy + 5), (sx, mid_y), (tx, mid_y), (tx, ty - 5)], label)


def connect_row_to_next_start(draw, source, target, label=None):
    sx, sy = anchor(source, "right")
    tx, ty = anchor(target, "top")
    path_y = lane_bounds(round((ty - TOP) / (LANE_H + LANE_GAP)))[0] - 26
    poly_arrow(draw, [(sx + 5, sy), (GUTTER_X, sy), (GUTTER_X, path_y), (tx, path_y), (tx, ty - 5)], label)


def save_chart(name, title, lanes, build):
    height = TOP + len(lanes) * LANE_H + (len(lanes) - 1) * LANE_GAP + 58
    img = Image.new("RGB", (W, height), "#ffffff")
    draw = ImageDraw.Draw(img)
    tw, _ = size(draw, title, TITLE_FONT)
    draw.text(((W - tw) / 2, TITLE_Y), title, font=TITLE_FONT, fill="#000000")
    draw_frame(draw, lanes)
    build(draw)
    path = FLOW_DIR / name
    img.save(path, quality=95)
    return path


def chart_5_1():
    lanes = ["用户学习入口层", "课程与教材服务层", "业务校验与学习过程层", "数据写入与反馈层"]

    def build(draw):
        r0 = row_nodes(draw, 0, [
            {"text": "学习者", "kind": "terminator", "w": 180},
            "进入在线学堂",
            "选择课程\n或教材资源",
            {"text": "检索分类\n关键词", "kind": "process"},
            "查看详情",
            {"text": "发起学习\n或购买", "kind": "process"},
        ])
        connect_row(draw, r0)
        r1 = row_nodes(draw, 1, [
            "学习接口\n接收请求",
            {"text": "资源类型\n是否合法", "kind": "decision"},
            "查询课程\n教材/评价/视频",
            "封装统一\n响应对象",
            "前端渲染\n列表与详情",
        ])
        connect_row_to_next_start(draw, r0[-1], r1[0])
        connect_row(draw, r1, [None, "是", None, None])
        r2 = row_nodes(draw, 2, [
            "课程存在性\n校验",
            {"text": "是否内部\n平台资源", "kind": "decision"},
            "创建报名\n订单/提交",
            "视频学习\n分段上报",
            "作业考试\n草稿/提交",
        ])
        connect_down(draw, r1[1], r2[0], "否/复核")
        connect_row(draw, r2, [None, "是", None, None])
        r3 = row_nodes(draw, 3, [
            "写入报名\n或教材订单",
            "写入学习时长\n记录表",
            "金币服务\n计算奖励",
            "个人画像\n更新统计",
            {"text": "页面反馈\n学习结果", "kind": "terminator", "w": 210},
        ])
        connect_down(draw, r2[2], r3[0])
        connect_down(draw, r2[3], r3[1])
        connect_row(draw, r3)

    return save_chart("clean_5_1_academy.png", "在线学堂模块详细设计流程图", lanes, build)


def chart_5_2():
    lanes = ["练习入口与题源层", "前端答题交互层", "题库服务与状态判定层", "画像与金币反馈层"]

    def build(draw):
        r0 = row_nodes(draw, 0, [
            {"text": "学习者", "kind": "terminator", "w": 180},
            "进入题库\n错题/收藏",
            "选择题库集合\n或课程题库",
            "加载分页题目\n与掌握状态",
            "展示题干\n选项/解析",
            "提交答案\n或收藏题目",
        ])
        connect_row(draw, r0)
        r1 = row_nodes(draw, 1, [
            {"text": "是否存在\n筛选条件", "kind": "decision"},
            "组合分页\n筛选参数",
            "调用题库接口",
            {"text": "题型是否\n自动判分", "kind": "decision"},
            "前端生成\n作答结果",
        ])
        connect_down(draw, r0[2], r1[0])
        connect_row(draw, r1, ["是", None, None, "是"])
        r2 = row_nodes(draw, 2, [
            "题库控制器\n接收请求",
            "数据仓库查询\n题目/错题/收藏",
            {"text": "答案正确\n或词汇掌握", "kind": "decision"},
            "更新错题\n收藏状态",
            "写入学习事件\n和学习时长",
        ])
        connect_down(draw, r1[2], r2[0])
        connect_row(draw, r2, [None, None, "是", None])
        r3 = row_nodes(draw, 3, [
            "画像服务\n沉淀行为",
            "金币奖励服务\n匹配规则",
            {"text": "来源键是否\n已奖励", "kind": "decision"},
            "写入金币流水\n奖励记录表",
            {"text": "刷新正确率\n错题数/金币", "kind": "terminator", "w": 220},
        ])
        connect_down(draw, r2[-1], r3[0])
        connect_row(draw, r3, [None, None, "否", None])
        connect_down(draw, r2[2], r3[0], "否/仅记录")
        connect_down(draw, r3[2], r3[-1], "是/跳过")

    return save_chart("clean_5_2_question_bank.png", "题库练习与金币激励模块详细设计流程图", lanes, build)


def chart_5_3():
    lanes = ["统一账号入口层", "角色功能分派层", "权限校验与业务处理层", "数据聚合与结果反馈层"]

    def build(draw):
        r0 = row_nodes(draw, 0, [
            {"text": "用户", "kind": "terminator", "w": 170},
            "登录/注册\n找回密码",
            "读取账号\n角色类型",
            {"text": "判断角色\n类型", "kind": "decision"},
            "进入对应\n功能页面",
        ])
        connect_row(draw, r0)
        r1 = row_nodes(draw, 1, [
            "学生：个人主页\n学习画像",
            "教师：课程发布\n工作台/信箱",
            "管理员：后台\n用户/课程/题库",
            "管理员：OJ\n卡券/评价",
            "提交查询\n或维护请求",
        ])
        connect_down(draw, r0[3], r1[0], "学生")
        connect_down(draw, r0[3], r1[1], "教师")
        connect_down(draw, r0[3], r1[2], "管理员")
        connect_row(draw, r1[2:])
        r2 = row_nodes(draw, 2, [
            "画像服务\n聚合学习数据",
            {"text": "是否通过\n教师校验", "kind": "decision"},
            {"text": "是否通过\n管理员校验", "kind": "decision"},
            "校验参数\n规范化表单",
            "数据仓库\n写入/查询",
        ])
        connect_down(draw, r1[0], r2[0])
        connect_down(draw, r1[1], r2[1])
        connect_down(draw, r1[2], r2[2])
        connect_row(draw, r2[1:], ["是", "是", None])
        r3 = row_nodes(draw, 3, [
            "学习事件/时长\n游戏/错题/订单",
            "教师发布课程\n课程评价信箱",
            "用户/课程/题库\nOJ/卡券表",
            "封装统计指标\n管理结果",
            {"text": "页面刷新\n展示反馈", "kind": "terminator", "w": 200},
        ])
        connect_down(draw, r2[0], r3[0])
        connect_down(draw, r2[4], r3[2])
        connect_row(draw, r3)

    return save_chart("clean_5_3_profile_admin.png", "用户画像、教师端与后台管理模块详细设计流程图", lanes, build)


def chart_5_4():
    lanes = ["页面上下文采集层", "宠物交互意图层", "AI服务处理层", "本地状态与学习反馈层"]

    def build(draw):
        r0 = row_nodes(draw, 0, [
            {"text": "用户", "kind": "terminator", "w": 170},
            "打开学习页面",
            "宠物组件\n读取路由",
            "采集标题/文本\n选中内容",
            {"text": "页面是否\n隐藏宠物", "kind": "decision"},
            "显示或隐藏\n宠物浮窗",
        ])
        connect_row(draw, r0, [None, None, None, None, "否"])
        r1 = row_nodes(draw, 1, [
            "用户输入\n问题或指令",
            {"text": "是否本地\n可执行动作", "kind": "decision"},
            "导航/待办\n番茄钟处理",
            "组装聊天请求\n历史+上下文",
            "调用宠物\n聊天接口",
        ])
        connect_down(draw, r0[-1], r1[0])
        connect_row(draw, r1, [None, "是", None, None])
        r2 = row_nodes(draw, 2, [
            "AI聊天服务\n校验密钥",
            "限制长度\n截取历史",
            "构建系统提示\n页面上下文",
            "调用模型接口",
            {"text": "返回内容\n是否有效", "kind": "decision"},
        ])
        connect_down(draw, r1[-1], r2[0])
        connect_row(draw, r2)
        r3 = row_nodes(draw, 3, [
            "本地存储\n保存任务/计时",
            "更新宠物状态\n思考/专注",
            "展示回答\n或错误提示",
            "辅助跳转\n学习入口",
            {"text": "形成学习陪伴\n与操作反馈", "kind": "terminator", "w": 235},
        ])
        connect_down(draw, r1[2], r3[0])
        connect_down(draw, r2[-1], r3[2], "是")
        connect_row(draw, r3)

    return save_chart("clean_5_4_ai_pet.png", "AI宠物与学习辅助模块详细设计流程图", lanes, build)


def chart_5_5():
    lanes = ["实践学习入口层", "前端交互与渲染层", "后端评测/记录层", "学习数据回流层"]

    def build(draw):
        r0 = row_nodes(draw, 0, [
            {"text": "学习者", "kind": "terminator", "w": 180},
            "进入可视化\n或实验平台",
            {"text": "选择实践\n类型", "kind": "decision"},
            "可视化学习\n函数/空间",
            "石油仿真\n测井/油藏",
            "OJ题目\n代码评测",
        ])
        connect_row(draw, r0, [None, None, "可视化", "仿真", "OJ"])
        r1 = row_nodes(draw, 1, [
            "图表与三维\n交互渲染",
            "仿真面板\n参数计算展示",
            "OJ页面\n编辑并提交",
            "轮询状态\n展示用例结果",
        ], xs=[520, 840, 1160, 1480])
        connect_down(draw, r0[3], r1[0])
        connect_down(draw, r0[4], r1[1])
        connect_down(draw, r0[5], r1[2])
        connect_row(draw, r1[2:])
        r2 = row_nodes(draw, 2, [
            "学习时长上报\n可视化/实验",
            "提交服务\n保存代码",
            "判题服务\n读取题目和用例",
            {"text": "沙箱地址\n是否配置", "kind": "decision"},
            "判题沙箱\n编译运行",
        ])
        connect_down(draw, r1[0], r2[0])
        connect_down(draw, r1[1], r2[0])
        connect_down(draw, r1[2], r2[1])
        connect_row(draw, r2[1:], [None, None, "是"])
        r3 = row_nodes(draw, 3, [
            "写入学习时长\n记录表",
            "更新提交状态\n得分/耗时/内存",
            "写入单用例\n评测结果",
            "金币服务\n按时长奖励",
            {"text": "个人画像\n展示成果", "kind": "terminator", "w": 200},
        ])
        connect_down(draw, r2[0], r3[0])
        connect_down(draw, r2[3], r3[1], "否/答案模式")
        connect_down(draw, r2[4], r3[1])
        connect_row(draw, r3)

    return save_chart("clean_5_5_visual_lab_oj.png", "可视化、实验与OJ模块详细设计流程图", lanes, build)


def chart_5_6():
    lanes = ["游戏学习入口层", "成绩采集与结算层", "金币与兑换服务层", "资产发放与展示层"]

    def build(draw):
        r0 = row_nodes(draw, 0, [
            {"text": "学习者", "kind": "terminator", "w": 180},
            "进入游戏\n学习平台",
            {"text": "选择游戏", "kind": "decision"},
            "万题天梯跳\n题库答题",
            "打字战士\n单词输入战斗",
            "生成成绩\n结算面板",
        ])
        connect_row(draw, r0, [None, None, "答题", "打字", None])
        r1 = row_nodes(draw, 1, [
            "提交游戏\n成绩记录",
            "游戏记录服务\n非负校验",
            "写入游戏记录\n成绩表",
            {"text": "奖励来源\n类型判断", "kind": "decision"},
            "计算平台金币\n题数/分数",
        ])
        connect_down(draw, r0[-1], r1[0])
        connect_row(draw, r1)
        r2 = row_nodes(draw, 2, [
            "金币奖励服务\n写入流水",
            "兑换中心读取\n余额与商品",
            {"text": "金币/库存\n是否充足", "kind": "decision"},
            "扣减金币\n减少库存",
            "发放卡券\n用户资产表",
        ])
        connect_down(draw, r1[-1], r2[0])
        connect_row(draw, r2, [None, None, "是", None])
        r3 = row_nodes(draw, 3, [
            "个人主页更新\n游戏指标",
            "兑换中心刷新\n余额/商品状态",
            "我的卡券\n展示资产",
            "教材订单\n使用优惠券",
            {"text": "完成激励\n反馈闭环", "kind": "terminator", "w": 210},
        ])
        connect_down(draw, r2[0], r3[0])
        connect_down(draw, r2[2], r3[1], "否/提示失败")
        connect_down(draw, r2[-1], r3[2])
        connect_row(draw, r3)

    return save_chart("clean_5_6_games_exchange.png", "学习游戏与兑换中心模块详细设计流程图", lanes, build)


charts = {
    "图5-1  在线学堂模块流程图": chart_5_1(),
    "图5-2  题库练习与金币激励模块流程图": chart_5_2(),
    "图5-3  用户画像、教师端与后台管理模块流程图": chart_5_3(),
    "图5-4  AI宠物与学习辅助模块流程图": chart_5_4(),
    "图5-5  可视化、实验与OJ模块流程图": chart_5_5(),
    "图5-6  学习游戏与兑换中心模块流程图": chart_5_6(),
}


def para_text(p):
    return "".join(run.text for run in p.runs).strip()


def replace_picture_paragraph(paragraph, image_path):
    paragraph.clear()
    paragraph.alignment = 1
    paragraph.add_run().add_picture(str(image_path), width=Inches(6.8))


doc = Document(BASE)
replaced = 0
for i, p in enumerate(doc.paragraphs):
    caption = para_text(p)
    if caption in charts:
        for j in range(i - 1, max(-1, i - 8), -1):
            prev = doc.paragraphs[j]
            if prev._p.xpath(".//w:drawing"):
                replace_picture_paragraph(prev, charts[caption])
                replaced += 1
                break

doc.save(OUT)
print(OUT)
print(f"replaced={replaced}")
for caption, path in charts.items():
    print(caption, path)
