from __future__ import annotations

import shutil
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


INPUT = Path("CET46/StudyPlatform-report-draft-ch1-6-brand-polished.docx")
OUTPUT = Path("CET46/StudyPlatform-report-draft-ch1-6-brand-polished-tables.docx")


def visual_len(text: str) -> float:
    total = 0.0
    for char in text:
        if char.isspace():
            total += 0.2
        elif unicodedata.east_asian_width(char) in {"F", "W", "A"}:
            total += 1.0
        elif char.isalnum():
            total += 0.55
        else:
            total += 0.45
    return total


def set_cell_text_font(cell, size_pt: float, bold: bool = False, header: bool = False) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.08
        for run in paragraph.runs:
            run.font.size = Pt(size_pt)
            run.font.bold = bold
            run.font.name = "SimSun" if not header else "SimHei"
            run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体" if not header else "黑体")
            run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Times New Roman")
            run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Times New Roman")


def set_cell_margins(cell, top: int = 90, bottom: int = 90, left: int = 120, right: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in {"top": top, "bottom": bottom, "left": left, "right": right}.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def border_node(name: str, value: str, size: int = 4, color: str = "D9D9D9") -> OxmlElement:
    node = OxmlElement(f"w:{name}")
    node.set(qn("w:val"), value)
    if value != "nil":
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)
    return node


def set_cell_borders(cell, top: tuple[str, int, str], bottom: tuple[str, int, str]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is not None:
        tc_pr.remove(borders)
    borders = OxmlElement("w:tcBorders")
    borders.append(border_node("top", *top))
    borders.append(border_node("left", "nil"))
    borders.append(border_node("bottom", *bottom))
    borders.append(border_node("right", "nil"))
    borders.append(border_node("insideH", "nil"))
    borders.append(border_node("insideV", "nil"))
    tc_pr.append(borders)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], table_width: int) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(table_width))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for col_index, cell in enumerate(row.cells):
            if col_index < len(widths):
                set_cell_width(cell, widths[col_index])


def usable_width_dxa(document: Document) -> int:
    section = document.sections[0]
    width_emu = section.page_width - section.left_margin - section.right_margin
    return int(width_emu / 635)


def column_stats(table) -> list[dict[str, float]]:
    stats = []
    for col_index in range(len(table.columns)):
        values = []
        for row_index in range(len(table.rows)):
            text = table.cell(row_index, col_index).text.replace("\n", " ").strip()
            values.append(visual_len(text))
        body = values[1:] if len(values) > 1 else values
        stats.append(
            {
                "max": max(values or [1.0]),
                "avg": sum(body or values or [1.0]) / max(len(body or values), 1),
                "header": values[0] if values else 1.0,
            }
        )
    return stats


def table_width_for(stats: list[dict[str, float]], usable_width: int) -> int:
    cols = len(stats)
    max_len = max(stat["max"] for stat in stats)
    if cols <= 4 and max_len <= 18:
        return min(int(usable_width * 0.72), 6600)
    if cols <= 4 and max_len <= 35:
        return int(usable_width * 0.84)
    if cols >= 7:
        return int(usable_width * 0.99)
    return int(usable_width * 0.94)


def distribute_widths(stats: list[dict[str, float]], table_width: int) -> list[int]:
    cols = len(stats)
    min_width = 580 if cols >= 8 else 700 if cols >= 6 else 850
    max_width = int(table_width * 0.48) if cols <= 4 else int(table_width * 0.36)
    weights = []
    for stat in stats:
        weight = max(3.5, stat["header"] * 0.9, stat["avg"] * 0.8 + stat["max"] * 0.55)
        if stat["max"] <= 8:
            weight *= 0.78
        if stat["max"] >= 45:
            weight *= 1.42
        if stat["max"] >= 90:
            weight *= 1.72
        weights.append(weight)
    widths = [max(min_width, int(table_width * weight / sum(weights))) for weight in weights]
    while sum(widths) > table_width:
        adjustable = [index for index, width in enumerate(widths) if width > min_width]
        if not adjustable:
            break
        excess = sum(widths) - table_width
        for index in sorted(adjustable, key=lambda idx: widths[idx], reverse=True):
            cut = min(widths[index] - min_width, max(1, excess // len(adjustable) + 1))
            widths[index] -= cut
            excess -= cut
            if excess <= 0:
                break
    if sum(widths) < table_width:
        spare = table_width - sum(widths)
        expandable = sorted(range(cols), key=lambda idx: stats[idx]["max"], reverse=True)
        for index in expandable:
            room = max(0, max_width - widths[index])
            add = min(room, spare)
            widths[index] += add
            spare -= add
            if spare <= 0:
                break
        if spare > 0:
            widths[-1] += spare
    return widths


def is_short_column(stats: dict[str, float], header_text: str) -> bool:
    compact_keywords = ["序号", "编号", "角色", "类型", "状态", "频度", "得分", "阶段", "模块", "难度", "占比", "文档", "验收"]
    return stats["max"] <= 14 or any(keyword in header_text for keyword in compact_keywords)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def polish_table(table, usable_width: int) -> None:
    if not table.rows or not table.columns:
        return
    stats = column_stats(table)
    table_width = table_width_for(stats, usable_width)
    widths = distribute_widths(stats, table_width)
    set_table_geometry(table, widths, table_width)
    repeat_header(table.rows[0])
    body_size = 8.5 if len(widths) >= 8 else 9.2 if len(widths) >= 6 else 10.0
    header_size = body_size
    last_row_index = len(table.rows) - 1
    headers = [table.cell(0, col_index).text.strip() for col_index in range(len(widths))]
    alignments = [
        WD_ALIGN_PARAGRAPH.CENTER if is_short_column(stats[col_index], headers[col_index]) else WD_ALIGN_PARAGRAPH.LEFT
        for col_index in range(len(widths))
    ]
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            if col_index >= len(widths):
                continue
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_text_font(cell, header_size, bold=True, header=True)
                set_cell_shading(cell, "F2F2F2")
                set_cell_borders(cell, ("single", 12, "000000"), ("single", 8, "000000"))
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                set_cell_text_font(cell, body_size, bold=False, header=False)
                set_cell_shading(cell, "FFFFFF")
                bottom = ("single", 12, "000000") if row_index == last_row_index else ("single", 3, "E5E7EB")
                set_cell_borders(cell, ("nil", 0, "FFFFFF"), bottom)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = alignments[col_index]


def add_spacing_around_tables(document: Document) -> None:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("表") and len(text) <= 40:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(4)
            for run in paragraph.runs:
                run.font.name = "SimHei"
                run.font.size = Pt(10.5)
                run.font.bold = True
                run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")


def main() -> None:
    if not INPUT.exists():
        raise FileNotFoundError(INPUT)
    shutil.copyfile(INPUT, OUTPUT)
    document = Document(OUTPUT)
    width = usable_width_dxa(document)
    for table in document.tables:
        polish_table(table, width)
    add_spacing_around_tables(document)
    document.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    main()
