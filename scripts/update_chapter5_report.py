from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "CET46" / "StudyPlatform-report-draft-ch1-6-brand-polished-tables.docx"
OUTPUT = ROOT / "CET46" / "StudyPlatform-report-draft-ch1-6-brand-polished-tables-chapter5-revised.docx"
FIGURE = ROOT / "docs" / "report-assets" / "studyplatform-module-detail-design.png"

BODY_FONT = "宋体"
HEADING_FONT = "黑体"


def set_run_font(run, size: float = 10.5, bold: bool = False, font_name: str = BODY_FONT) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def paragraph_text(element, doc: Document) -> str:
    if element.tag != qn("w:p"):
        return ""
    from docx.text.paragraph import Paragraph

    return Paragraph(element, doc).text.strip()


def normalize(text: str) -> str:
    return re.sub(r"\s+", "", text)


def find_chapter_range(doc: Document) -> tuple[int, int]:
    body = doc.element.body
    children = list(body.iterchildren())
    start = None
    end = None
    for index, child in enumerate(children):
        text = normalize(paragraph_text(child, doc))
        if start is None and text.startswith("5") and "系统详细设计" in text:
            start = index
            continue
        if start is not None and text.startswith("6") and "数据库设计" in text:
            end = index
            break
    if start is None or end is None:
        raise RuntimeError("无法定位第5章或第6章标题。")
    return start, end


def remove_chapter5(doc: Document) -> None:
    body = doc.element.body
    children = list(body.iterchildren())
    start, end = find_chapter_range(doc)
    for child in children[start:end]:
        body.remove(child)


def find_insert_marker(doc: Document):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p") and normalize(paragraph_text(child, doc)).startswith("6") and "数据库设计" in normalize(paragraph_text(child, doc)):
            return child
    raise RuntimeError("无法定位第6章插入锚点。")


def insert_before(doc: Document, marker, element) -> None:
    body = doc.element.body
    body.insert(body.index(marker), element)


def add_paragraph(doc: Document, marker, text: str = "", level: int = 0, align=None):
    paragraph = doc.add_paragraph()
    insert_before(doc, marker, paragraph._p)
    if align is not None:
        paragraph.alignment = align
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(5 if level == 0 else 4)
    fmt.line_spacing = 1.25
    if level == 0:
        fmt.first_line_indent = None
        size, bold, font = 15, True, HEADING_FONT
    elif level == 1:
        fmt.first_line_indent = None
        size, bold, font = 13, True, HEADING_FONT
    elif level == 2:
        fmt.first_line_indent = None
        size, bold, font = 11.5, True, HEADING_FONT
    elif level == 3:
        fmt.first_line_indent = None
        size, bold, font = 10.5, True, HEADING_FONT
    else:
        fmt.first_line_indent = Pt(21)
        size, bold, font = 10.5, False, BODY_FONT
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, font_name=font)
    return paragraph


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "666666")
        borders.append(element)
    tbl_pr.append(borders)


def set_table_width(table, width_twips: int) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    if grid is not None:
        table._tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    table._tbl.insert(1, grid)


def add_grid_col(table, width_twips: int) -> None:
    grid = table._tbl.tblGrid
    col = OxmlElement("w:gridCol")
    col.set(qn("w:w"), str(width_twips))
    grid.append(col)


def set_cell_text(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT, size: float = 9.5) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, font_name=HEADING_FONT if bold else BODY_FONT)


def mark_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_caption(doc: Document, marker, text: str) -> None:
    paragraph = add_paragraph(doc, marker, text, level=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run, size=9.5, bold=False, font_name=BODY_FONT)


def add_table(doc: Document, marker, caption: str, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    add_caption(doc, marker, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    insert_before(doc, marker, table._tbl)
    table.autofit = False
    table.allow_autofit = False
    total_twips = int(sum(widths) * 1440)
    set_table_width(table, total_twips)
    for width in widths:
        add_grid_col(table, int(width * 1440))
    header_row = table.rows[0]
    mark_repeat_header(header_row)
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        set_cell_width(cell, int(widths[index] * 1440))
        set_cell_shading(cell, "F2F2F2")
        set_cell_text(cell, header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9.5)
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            set_cell_width(cell, int(widths[index] * 1440))
            set_cell_text(cell, value, align=WD_ALIGN_PARAGRAPH.LEFT, size=9.2)
    set_table_borders(table)
    add_paragraph(doc, marker, "", level=4)


def add_figure(doc: Document, marker) -> None:
    paragraph = doc.add_paragraph()
    insert_before(doc, marker, paragraph._p)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run()
    run.add_picture(str(FIGURE), width=Inches(6.5))
    add_caption(doc, marker, "图5-1  金币激励与学习模块详细设计流程图")


def copy_document_core_properties(src: Document, dst: Document) -> None:
    # python-docx preserves core properties when opening and saving the same document.
    # This hook is intentionally empty; it documents why no metadata rewrite is needed.
    _ = (src, dst)


def build_chapter5(doc: Document) -> None:
    marker = find_insert_marker(doc)

    add_paragraph(doc, marker, "5  系统详细设计", level=0)

    add_paragraph(doc, marker, "5.1  在线学堂模块", level=1)
    add_paragraph(doc, marker, "5.1.1  模块描述", level=2)
    add_paragraph(
        doc,
        marker,
        "在线学堂模块是 StudyPlatform 承载课程学习、作业考试和精品教材服务的核心业务模块，负责把不同来源的课程资源组织为统一的学习入口，并把学习行为继续沉淀到用户画像和金币激励中心。前端由在线学堂外壳页（AcademyPage.vue）、在线学堂首页（AcademyHome.vue）、我的班级页（AcademyMyClass.vue）、我的课程/作业/考试聚合页（AcademyAggregatePage.vue）、作业详情页（AcademyAssignmentDetail.vue）、考试介绍页（AcademyExamIntro.vue）、考试作答页（AcademyExamDetail.vue）、在线开放课程列表页（AcademyOpenCourses.vue）、通识课程列表页（AcademyGeneralCourses.vue）、微专业课程列表页（AcademyMicroMajors.vue）、课程详情与视频播放页（AcademyCourseDetail.vue）、精品教材列表页（AcademyTextbooks.vue）、教材详情与购买页（AcademyTextbookDetail.vue）以及教材购物车与结算页（AcademyTextbookCart.vue）组成。后端主要由 AcademyController、AcademyService、AcademyAssignmentService 和 AcademyExamService 提供 REST 接口与业务处理。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "模块设计采用“资源类型 + 资源编号”的抽象方式，将在线开放课程、通识课程和微专业课程分别映射为 online-open-courses、general-courses 和 micro-major-courses。前端路由把 resource、listPath、moduleTitle 和 courseId 传入课程详情页，使同一套详情、加入课程、评价和视频学习逻辑可以复用于三类课程，同时又保留不同课程来源在分类、学校、教师、封面和来源链接上的差异。",
        level=4,
    )
    add_paragraph(doc, marker, "5.1.2  课程资源聚合与详情展示", level=2)
    add_paragraph(
        doc,
        marker,
        "课程资源聚合由 useAcademyList.js 统一完成。该组合函数同时请求分类接口和资源列表接口，按“全部/分类”按钮、关键词和分页参数生成前端可展示的课程卡片。在线开放课程、通识课程、微专业课程和精品教材均复用该机制，只是搜索字段分别对应课程名称、教师、学校、分类或教材名称、主编、出版社、ISBN 等业务字段。后端根据资源类型分别查询 online_open_courses、general_courses、micro_major_courses 和 excellent_textbooks 等表，并返回封面、来源链接、参与人数、课程介绍等展示信息。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "课程详情页（AcademyCourseDetail.vue）根据路由传入的 resource 和 courseId 调用 fetchAcademyCourse(resource, id)，展示课程封面、教师、学校、开课时间、课程说明、评价列表和视频资源。封面与视频地址通过 resolveResourceUrl 统一处理，既兼容外部采集链接，也兼容教师上传后保存到本地 storage/teacher_courses 目录的资源文件。课程评价通过 /reviews 接口独立加载和提交，避免课程基础信息与用户互动数据相互耦合。",
        level=4,
    )
    add_paragraph(doc, marker, "5.1.3  课程加入、视频学习与作业考试", level=2)
    add_paragraph(
        doc,
        marker,
        "课程加入流程以“平台内部资源”为准，而不是直接依赖外部网站页面地址。系统在课程采集/导入阶段从中国大学MOOC、超星、雨课堂、智慧树等来源保留 external_course_id、source_url、cover_url 和 cover_file_path 等字段；用户点击“立即参加”时，前端调用 enrollAcademyCourse(resource, id, { userId })，后端先通过 ensureCourseExists 校验对应资源类型下课程存在，再写入 academy_course_enrollments 关系记录。这样既能追溯课程来源，也能保证加入关系属于平台自身数据闭环。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "视频学习由 useVideoLearningTimeTracker 封装。该逻辑只在视频处于 play/playing 状态且页面可见时累计学习时长，在 pause、ended、waiting、页面隐藏或离开页面时通过 /api/profile/learning-time 上报 moduleType=video、targetCode 和 durationSeconds。后端 ProfileService 保存学习时长记录后交由 CoinRewardService 按 10 分钟一个结算单元计算金币，避免前端直接决定奖励数量。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "作业与考试入口由在线学堂首页和聚合页共同承载。我的课程/作业/考试聚合页（AcademyAggregatePage.vue）通过 variant 区分三类数据；作业详情页支持草稿保存和最终提交；考试介绍页负责展示考试说明，考试作答页负责开始考试、保存草稿和提交试卷。考试与作业详情路由配置 hidePet 元数据，答题过程中隐藏 AI 宠物浮窗，减少对严肃作答场景的干扰。",
        level=4,
    )
    add_paragraph(doc, marker, "5.1.4  精品教材浏览与教材订单", level=2)
    add_paragraph(
        doc,
        marker,
        "精品教材列表页（AcademyTextbooks.vue）复用 useAcademyList('textbooks') 实现分类筛选、关键词检索和分页展示。教材基础数据在导入阶段保留 external_textbook_id、source_url、cover_url、cover_file_path、ISBN、出版社和出版日期等字段，来源主要对应中国大学MOOC精品教材页面；教材详情扩展信息由 academy_textbook_details 表保存推荐语、原价、折扣价、阅读人数、内容概述、目录和评论文本，使列表数据与商店化详情数据分层维护。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "教材详情与购买页（AcademyTextbookDetail.vue）支持教材介绍、目录展开、购物车、立即购买、模拟支付和购买后评论。购物车页（AcademyTextbookCart.vue）支持勾选教材、调整数量、删除条目、计算优惠/运费/应付金额并创建订单。后端通过 academy_textbook_cart_items、academy_textbook_orders 和 academy_textbook_order_items 保存购物车与订单明细，支付接口将订单状态更新为“已支付”，并据此控制用户是否可以发布教材评价。",
        level=4,
    )
    add_table(
        doc,
        marker,
        "表5-1  在线学堂核心功能与技术设计",
        ["功能域", "页面/接口", "技术设计要点"],
        [
            ["课程资源聚合", "在线开放课程列表页（AcademyOpenCourses.vue）、通识课程列表页（AcademyGeneralCourses.vue）、微专业课程列表页（AcademyMicroMajors.vue）；/api/academy/{resource}", "统一使用 useAcademyList 组织分类、关键词与分页；resourceType 区分不同课程表，保留 external_course_id 和 source_url 追溯采集来源。"],
            ["课程详情展示", "课程详情与视频播放页（AcademyCourseDetail.vue）；/api/academy/{resource}/{id}", "同一详情页复用封面、教师、学校、介绍、评价和视频播放区域；resolveResourceUrl 兼容外部链接与本地上传文件。"],
            ["课程加入", "立即参加按钮；POST /api/academy/{resource}/{id}/enroll", "后端先校验课程存在，再写入 academy_course_enrollments；加入关系只记录平台内部 resourceType 与 courseId。"],
            ["视频学习", "HTML5 video + useVideoLearningTimeTracker；POST /api/profile/learning-time", "只在播放且页面可见时计时，离开页面或暂停时上报；后端按时长统一结算金币。"],
            ["教材订单", "精品教材列表页、教材详情与购买页、教材购物车与结算页；/textbook-cart、/textbook-orders", "教材列表、详情、购物车、订单、支付和购买后评论分层实现，订单状态控制教材评价权限。"],
        ],
        [1.25, 2.35, 2.90],
    )

    add_paragraph(doc, marker, "5.2  题库练习与金币激励模块", level=1)
    add_paragraph(doc, marker, "5.2.1  模块描述", level=2)
    add_paragraph(
        doc,
        marker,
        "题库练习模块负责将课程题库、英语词汇题库、错题复习和收藏题目组织为统一练习流程。前端由题库首页（AcademyQuestionBank.vue）、课程题库列表页（AcademyQuestionBankCourses.vue）、课程题库练习页（AcademyQuestionBankCourseDetail.vue）、错题本页（AcademyQuestionBankMistakes.vue）和收藏题目页（AcademyQuestionBankFavorites.vue）组成。后端由 QuestionBankController 和 QuestionBankService 提供题库目录、题目分页、答题记录、错题统计、收藏状态和 Type Warrior 单词池接口。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "题库首页负责展示错题数、收藏数和练习入口；课程题库列表页负责按课程或题库分类组织入口；课程题库练习页负责题目分页、关键词检索、选择题/词汇题作答、收藏切换和学习时长记录；错题本页负责按题库、状态和关键词筛选错题，并在连续答对后更新掌握状态；收藏题目页负责保存重点题目并支持快速复习。",
        level=4,
    )
    add_paragraph(doc, marker, "5.2.2  练习流程与数据闭环", level=2)
    add_paragraph(
        doc,
        marker,
        "普通题目练习时，前端在用户提交答案后同时执行两类上报：一是调用 recordQuestionBankAnswer 写入错题/掌握状态，二是调用 recordProfileLearningEvent 写入学习行为事件。对于选择题，系统比对用户选项与正确答案后生成 isCorrect；对于词汇题，系统根据“认识/不认识”等状态写入 vocabularyStatus。题库页面还通过 useLearningTimeTracker 记录 moduleType=question_bank、mistake 或 favorite 的在线时长，使“答题行为”和“复习时长”都能进入用户画像。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "题库数据来源包括平台内置课程题库、英语词汇题库以及洛谷题目导入接口。/api/academy/question-bank/import/luogu 支持按页数和数量导入外部题目，导入后统一进入题库目录和题目分页接口。这样做可以把外部题目抓取、平台题目展示和用户练习记录解耦，便于后续扩展更多题源。",
        level=4,
    )
    add_paragraph(doc, marker, "5.2.3  金币激励结算设计", level=2)
    add_paragraph(
        doc,
        marker,
        "金币激励模块的核心原则是“前端提交行为，后端统一结算”。前端只提交学习时长、答题正确性、词汇掌握状态或游戏成绩，奖励金额由后端 CoinRewardService 根据 moduleType、eventType、questionType 和游戏记录统一计算。学习时长按 10 分钟为一个结算单元，视频学习每单元 5 金币，可视化学习每单元 3 金币，油气仿真每单元 4 金币，题库、错题、OJ、作业和考试每单元 2 金币；答对选择题、简答题和掌握词汇可获得事件奖励；万题天梯跳直接按游戏金币结算，Type Warrior 按得分折算。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "金币奖励记录写入 coin_reward_records 表，表中 source_type 和 source_key 用于区分奖励来源，唯一约束 user_id + source_type + source_key 用于幂等去重。ProfileService 在生成个人主页概览时汇总 CoinRewardService.totalCoins(userId)，并叠加管理员调整值，供个人主页和兑换中心展示当前金币余额。",
        level=4,
    )
    add_figure(doc, marker)
    add_table(
        doc,
        marker,
        "表5-2  题库与金币关键接口设计",
        ["业务动作", "主要接口/服务", "设计说明"],
        [
            ["题库目录与题目分页", "GET /api/academy/question-bank/course-catalog；GET /courses/{code}", "按题库编码加载题目，支持分页和关键词检索，返回收藏状态与题目类型。"],
            ["答题与错题同步", "POST /api/academy/question-bank/mistakes/answers", "提交题目编号和用户答案，后端更新错误次数、连续答对次数和掌握状态。"],
            ["学习事件奖励", "POST /api/profile/events；ProfileService.recordLearningEvent", "记录答题正确性和词汇掌握状态，并交由 CoinRewardService 生成学习行为奖励。"],
            ["学习时长奖励", "POST /api/profile/learning-time；ProfileService.recordLearningTime", "记录题库、错题、收藏、视频、可视化和实验等模块时长，按 10 分钟单元结算金币。"],
            ["金币幂等落表", "CoinRewardRepository.insertReward；coin_reward_records", "通过 user_id、source_type、source_key 唯一约束避免刷新页面或重复上报造成重复发放。"],
        ],
        [1.35, 2.35, 2.80],
    )

    add_paragraph(doc, marker, "5.3  用户画像、教师端与后台管理模块", level=1)
    add_paragraph(doc, marker, "5.3.1  用户画像与账号角色", level=2)
    add_paragraph(
        doc,
        marker,
        "用户画像模块用于沉淀学习过程数据，前端主要由个人主页（ProfilePage.vue）展示学习时长、今日学习、答题正确率、词汇掌握、错题统计、游戏成绩、金币余额、近期动态和成就徽章。后端通过 ProfileController 和 ProfileService 提供 /api/profile/overview、/user、/avatar、/events 和 /learning-time 等接口，其中学习事件和学习时长在写入画像数据后会进一步触发金币奖励服务。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "账号与角色能力贯穿登录注册、忘记密码、引导页、个人主页和后台管理。系统通过登录页（AuthLoginPage.vue）、注册页（AuthRegisterPage.vue）、忘记密码页（AuthForgotPasswordPage.vue）和账号引导页（AuthOnboardingPage.vue）完成账号登录、注册、密码重置和初始资料补充；后台用户字段包含 roleType，用于区分学生、教师和管理员。学生主要使用学习与练习功能，教师侧重课程发布与维护，管理员负责平台数据和运营配置。",
        level=4,
    )
    add_paragraph(doc, marker, "5.3.2  教师端课程发布", level=2)
    add_paragraph(
        doc,
        marker,
        "教师端不是单独的路由模块，而是嵌入个人主页和在线学堂服务中。教师账号进入个人主页后可使用课程发布面板，调用 publishOnlineOpenCourse 提交课程名称、开课时间、学期安排、课程详情、课程概述、课程封面和课程视频。后端通过 X-Auth-User-Id 获取当前用户，ensureTeacher 校验 roleType=teacher 后才允许继续处理，并将封面和视频保存到 storage/teacher_courses/{userId} 目录。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "教师发布的课程会写入在线开放课程数据源，并标记为“教师发布”分类，便于与外部采集课程共同展示但来源可区分。教师还可以通过 /api/academy/online-open-courses/teacher/mine 查看自己发布的课程，并只能删除属于当前教师的课程，避免不同教师之间误删资源。",
        level=4,
    )
    add_paragraph(doc, marker, "5.3.3  后台管理", level=2)
    add_paragraph(
        doc,
        marker,
        "后台管理模块面向管理员账号，前端由后台管理页（AdminPage.vue）承载，后端由 AdminController 和 AdminService 提供用户管理、课程管理、课程评价管理、题库集管理和题目管理接口。管理员可以维护用户角色、教师信息和账号状态，管理在线学堂课程资源，删除不合规评论，维护题库集合与题目内容。后端在进入管理逻辑前校验管理员身份，避免普通学生或教师访问管理接口。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "金币规则和兑换商品属于平台运营配置。当前代码已经实现金币奖励记录、金币余额汇总和兑换中心展示，但兑换中心仍处于展示阶段，页面明确标注“暂未开启真实扣费兑换”。因此报告中将金币规则管理和兑换商品管理作为后台管理的扩展方向，当前实现重点是学习行为奖励结算与余额展示。",
        level=4,
    )

    add_paragraph(doc, marker, "5.4  AI宠物与学习辅助模块", level=1)
    add_paragraph(doc, marker, "5.4.1  模块描述", level=2)
    add_paragraph(
        doc,
        marker,
        "AI宠物模块由全局浮窗组件（AiPetWidget.vue）实现，不是独立页面路由。组件在应用主要页面中以可拖拽宠物形象存在，并根据路由 meta.hidePet 在作业、考试等严肃作答页面自动隐藏。前端资源包含待机、思考、开心、专注和休息等多组帧动画，组件将当前聊天状态、番茄钟状态和待办完成状态映射为不同情绪。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "AI宠物提供三个层次的学习辅助：第一，聊天问答通过 chatWithAiPet 调用 /api/ai-pet/chat，并携带 pageContext，使后端能够结合当前页面上下文生成学习建议；第二，内置导航意图识别可根据“打开题库”“进入可视化学习”等自然语言跳转到在线学堂、题库、个人主页、可视化和游戏页面；第三，待办清单和番茄钟在本地 localStorage 保存，支持创建学习任务、勾选完成、设置专注/休息时长和累计专注轮次。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "从系统设计角度看，AI宠物是学习动机增强层，与金币激励和兑换中心存在天然关联。当前实现已经具备陪伴、问答、导航、待办和专注计时功能；后续可将兑换中心中的宠物形象、装扮和称号与金币余额关联，使学习奖励能够反馈到宠物外观和互动体验上。",
        level=4,
    )

    add_paragraph(doc, marker, "5.5  可视化、实验与 OJ 模块", level=1)
    add_paragraph(doc, marker, "5.5.1  可视化学习模块", level=2)
    add_paragraph(
        doc,
        marker,
        "可视化学习模块由可视化实验中心（VisualizationHome.vue）、算法结构可视化页（DataStructureVisualization.vue）、算法演示查看页（AlgorithmDemoViewer.vue）、二维函数图像实验室（FunctionGraph2D.vue）、空间模型实验室入口页（SpaceModelGuide.vue）和三维空间模型页（SpaceModel3D.vue）组成。算法结构可视化主要加载本地演示资源；函数图像实验室基于 ECharts 绘制函数曲线；三维空间模型页基于 Three.js 和 OrbitControls 展示可交互的空间几何模型。上述页面均通过 useLearningTimeTracker 上报 moduleType=visualization 的学习时长。",
        level=4,
    )
    add_paragraph(doc, marker, "5.5.2  实验平台与 OJ 模块", level=2)
    add_paragraph(
        doc,
        marker,
        "实验平台入口页（LabPlatform.vue）负责引导用户进入 OJ 在线编程和石油气仿真平台。OJ 在线编程页（OjPlatform.vue）加载题目、展示题面、接收 C++ 代码并提交到后端判题接口，后端再与独立 judge-sandbox 服务交互完成编译运行和结果回传。页面提示需要启动 judge-sandbox 并配置 oj.sandbox-url=http://localhost:9000，体现了平台后端与外部判题沙箱之间的边界。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "石油气仿真平台（PetroleumSimulation.vue）将测井曲线仿真、抽油机展示功图和注水开发分析组织为选项卡式实验场景，并通过 moduleType=petroleum 上报实验学习时长。测井仿真旧入口（WellLogSimulation.vue）会自动重定向到石油气仿真平台，避免同一实验能力在路由层重复维护。",
        level=4,
    )

    add_paragraph(doc, marker, "5.6  学习游戏、金币激励与兑换中心模块", level=1)
    add_paragraph(doc, marker, "5.6.1  学习游戏模块", level=2)
    add_paragraph(
        doc,
        marker,
        "学习游戏模块由游戏学习平台（GamePlatform.vue）统一承载，目前包含万题天梯跳和 Type Warrior 两类游戏。万题天梯跳通过题库答题、平台跳跃和游戏内金币形成低门槛练习体验；Type Warrior 通过英文单词输入、波次、连击、击杀数和得分训练词汇反应能力。游戏结束后，前端分别调用 /api/games/ladder-jump/records 和 /api/games/type-warrior/records 保存记录。",
        level=4,
    )
    add_paragraph(
        doc,
        marker,
        "后端 GameRecordService 对游戏记录中的金币、答对/答错数量、得分、波次、连击和有效输入时长进行非负校验后入库，并立即调用 CoinRewardService。万题天梯跳按 totalCoins 发放平台金币，Type Warrior 按 score/100 折算金币。游戏成绩随后进入 ProfileService 的个人主页聚合指标，用于展示最佳成绩、平均成绩和累计金币。",
        level=4,
    )
    add_paragraph(doc, marker, "5.6.2  兑换中心模块", level=2)
    add_paragraph(
        doc,
        marker,
        "兑换中心页（ExchangeCenter.vue）通过 fetchProfileOverview 获取当前金币余额，并展示金币来源规则、AI宠物形象、装扮、主题和徽章等兑换展示项。当前页面按钮处于禁用状态，说明系统已经完成金币获取、余额展示和兑换资产展示原型，但尚未实现真实扣减金币和发放资产的交易流程。该设计为后续扩展兑换订单、用户虚拟资产表和后台商品管理预留了清晰入口。",
        level=4,
    )
    add_table(
        doc,
        marker,
        "表5-3  第五章补充模块与第四章功能架构对应关系",
        ["第四章功能模块", "第五章详细设计位置", "对应实现说明"],
        [
            ["用户与账号", "5.3 用户画像、教师端与后台管理模块", "登录、注册、密码重置、角色字段、个人资料和管理员/教师权限统一说明。"],
            ["在线学堂", "5.1 在线学堂模块", "课程聚合、详情、加入、视频学习、作业考试、精品教材和教材订单均展开描述。"],
            ["题库练习", "5.2 题库练习与金币激励模块", "题库目录、答题、错题、收藏、词汇练习和外部题源导入形成闭环。"],
            ["AI宠物", "5.4 AI宠物与学习辅助模块", "全局浮窗、聊天问答、导航意图、待办和番茄钟被纳入详细设计。"],
            ["可视化与实验", "5.5 可视化、实验与 OJ 模块", "算法、函数、三维模型、OJ 判题和石油气仿真对应第四章实验能力。"],
            ["学习游戏/金币/兑换", "5.6 学习游戏、金币激励与兑换中心模块", "游戏记录、金币结算、个人主页余额和兑换中心展示共同支撑激励闭环。"],
            ["后台管理", "5.3 用户画像、教师端与后台管理模块", "管理员维护用户、课程、评论、题库和题目，教师发布课程独立补充。"],
        ],
        [1.45, 2.05, 3.00],
    )


def main() -> None:
    if not INPUT.exists():
        raise FileNotFoundError(INPUT)
    if not FIGURE.exists():
        raise FileNotFoundError(FIGURE)
    doc = Document(str(INPUT))
    remove_chapter5(doc)
    build_chapter5(doc)
    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
